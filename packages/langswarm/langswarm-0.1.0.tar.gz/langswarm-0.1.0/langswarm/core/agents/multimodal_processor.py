"""
LangSwarm V2 Multimodal Processor Implementation

Base multimodal processor that provides common functionality for processing
images, videos, audio, documents and performing cross-modal reasoning.
"""

import asyncio
import base64
import logging
import time
from abc import ABC, abstractmethod
from datetime import datetime
from typing import Any, Dict, List, Optional
import json

try:
    import cv2
    import numpy as np
    from PIL import Image, ImageEnhance, ImageFilter
    import pytesseract
except ImportError:
    cv2 = None
    np = None
    Image = None
    ImageEnhance = None
    ImageFilter = None
    pytesseract = None

try:
    import PyPDF2
    import docx
    from pdfplumber import PDF
except ImportError:
    PyPDF2 = None
    docx = None
    PDF = None

try:
    import speech_recognition as sr
    from pydub import AudioSegment
except ImportError:
    sr = None
    AudioSegment = None

from .multimodal import (
    IMultimodalProcessor, MultimodalContent, MultimodalRequest, MultimodalResponse,
    MediaType, ModalityType, ProcessingMode
)

logger = logging.getLogger(__name__)


class BaseMultimodalProcessor(IMultimodalProcessor):
    """
    Base implementation of multimodal processor with common functionality.
    
    Provides default implementations for:
    - Image analysis and OCR
    - Document text extraction
    - Audio transcription (if dependencies available)
    - Video frame extraction and analysis
    - Cross-modal reasoning coordination
    """
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which optional dependencies are available"""
        self.has_cv2 = cv2 is not None
        self.has_pil = Image is not None
        self.has_ocr = pytesseract is not None
        self.has_pdf = PyPDF2 is not None
        self.has_docx = docx is not None
        self.has_audio = sr is not None and AudioSegment is not None
        
        missing_deps = []
        if not self.has_pil:
            missing_deps.append("Pillow")
        if not self.has_ocr:
            missing_deps.append("pytesseract")
        if not self.has_cv2:
            missing_deps.append("opencv-python")
        
        if missing_deps:
            logger.warning(f"Optional dependencies missing: {missing_deps}. "
                         f"Some multimodal features may be limited.")
    
    @property
    def supported_modalities(self) -> List[ModalityType]:
        """List of supported modalities"""
        modalities = [ModalityType.TEXT]
        
        if self.has_pil:
            modalities.append(ModalityType.IMAGE)
        if self.has_cv2:
            modalities.append(ModalityType.VIDEO)
        if self.has_audio:
            modalities.append(ModalityType.AUDIO)
        if self.has_pdf or self.has_docx:
            modalities.append(ModalityType.DOCUMENT)
        
        return modalities
    
    @property
    def supported_media_types(self) -> List[MediaType]:
        """List of supported media types"""
        media_types = [MediaType.TEXT_PLAIN]
        
        if self.has_pil:
            media_types.extend([
                MediaType.IMAGE_JPEG, MediaType.IMAGE_PNG, MediaType.IMAGE_GIF,
                MediaType.IMAGE_WEBP, MediaType.IMAGE_BMP
            ])
        
        if self.has_cv2:
            media_types.extend([
                MediaType.VIDEO_MP4, MediaType.VIDEO_AVI, MediaType.VIDEO_MOV
            ])
        
        if self.has_audio:
            media_types.extend([
                MediaType.AUDIO_MP3, MediaType.AUDIO_WAV, MediaType.AUDIO_AAC
            ])
        
        if self.has_pdf:
            media_types.append(MediaType.DOCUMENT_PDF)
        if self.has_docx:
            media_types.extend([
                MediaType.DOCUMENT_DOC, MediaType.DOCUMENT_DOCX
            ])
        
        return media_types
    
    @property
    def supported_processing_modes(self) -> List[ProcessingMode]:
        """List of supported processing modes"""
        return [
            ProcessingMode.ANALYZE,
            ProcessingMode.EXTRACT,
            ProcessingMode.TRANSFORM,
            ProcessingMode.COMPARE,
            ProcessingMode.SEARCH
        ]
    
    async def process_content(
        self,
        request: MultimodalRequest
    ) -> MultimodalResponse:
        """Process multimodal content based on request"""
        start_time = time.time()
        response = MultimodalResponse(request_id=request.request_id)
        
        try:
            # Process each piece of content
            for content in request.content:
                await self._process_single_content(content, request, response)
            
            # Perform cross-modal reasoning if requested and multiple content pieces
            if request.cross_modal_context and len(request.content) > 1:
                cross_modal_results = await self._perform_cross_modal_analysis(
                    request.content, request.instructions
                )
                response.cross_modal_insights = cross_modal_results
            
            response.processing_time = time.time() - start_time
            response.success = True
            
        except Exception as e:
            response.success = False
            response.error = str(e)
            logger.error(f"Error processing multimodal content: {e}")
        
        return response
    
    async def _process_single_content(
        self,
        content: MultimodalContent,
        request: MultimodalRequest,
        response: MultimodalResponse
    ) -> None:
        """Process a single piece of content"""
        if content.modality == ModalityType.IMAGE:
            analysis = await self.analyze_image(content, request.instructions)
            response.add_analysis(content.content_id, analysis)
            
        elif content.modality == ModalityType.VIDEO:
            analysis = await self.analyze_video(content, request.instructions)
            response.add_analysis(content.content_id, analysis)
            
        elif content.modality == ModalityType.AUDIO:
            analysis = await self.analyze_audio(content, request.instructions)
            response.add_analysis(content.content_id, analysis)
            
        elif content.modality == ModalityType.DOCUMENT:
            analysis = await self.analyze_document(content, request.instructions)
            response.add_analysis(content.content_id, analysis)
        
        # Extract text if requested
        if request.processing_mode in [ProcessingMode.EXTRACT, ProcessingMode.ANALYZE]:
            extracted_text = await self.extract_text(content)
            if extracted_text:
                response.add_extracted_content(content.content_id, extracted_text)
    
    async def analyze_image(
        self,
        image: MultimodalContent,
        instructions: Optional[str] = None
    ) -> Dict[str, Any]:
        """Analyze image content using available tools"""
        if not self.has_pil:
            return {"error": "PIL not available for image analysis"}
        
        try:
            content_bytes = image.get_content_bytes()
            if not content_bytes:
                return {"error": "No image data available"}
            
            # Load image
            pil_image = Image.open(io.BytesIO(content_bytes))
            
            # Basic image analysis
            analysis = {
                "format": pil_image.format,
                "mode": pil_image.mode,
                "size": pil_image.size,
                "width": pil_image.width,
                "height": pil_image.height,
                "has_transparency": pil_image.mode in ['RGBA', 'LA'] or 'transparency' in pil_image.info,
                "estimated_colors": len(pil_image.getcolors(maxcolors=256)) if pil_image.getcolors(maxcolors=256) else "many"
            }
            
            # OCR if available and requested
            if self.has_ocr and (not instructions or "text" in instructions.lower()):
                try:
                    extracted_text = pytesseract.image_to_string(pil_image)
                    if extracted_text.strip():
                        analysis["extracted_text"] = extracted_text.strip()
                        analysis["text_confidence"] = "medium"  # Would need more sophisticated analysis
                except Exception as e:
                    analysis["ocr_error"] = str(e)
            
            # Basic image features
            if self.has_cv2:
                cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
                
                # Detect edges
                gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
                edges = cv2.Canny(gray, 100, 200)
                edge_density = np.sum(edges > 0) / (edges.shape[0] * edges.shape[1])
                analysis["edge_density"] = float(edge_density)
                
                # Basic color analysis
                mean_color = np.mean(cv_image, axis=(0, 1))
                analysis["average_color"] = {
                    "b": int(mean_color[0]),
                    "g": int(mean_color[1]), 
                    "r": int(mean_color[2])
                }
            
            return analysis
            
        except Exception as e:
            return {"error": f"Failed to analyze image: {str(e)}"}
    
    async def analyze_video(
        self,
        video: MultimodalContent,
        instructions: Optional[str] = None
    ) -> Dict[str, Any]:
        """Analyze video content"""
        if not self.has_cv2:
            return {"error": "OpenCV not available for video analysis"}
        
        try:
            # For now, extract frames and analyze them as images
            # In a full implementation, this would include motion analysis, scene detection, etc.
            
            video_path = video.file_path
            if not video_path:
                return {"error": "Video file path required for analysis"}
            
            cap = cv2.VideoCapture(video_path)
            if not cap.isOpened():
                return {"error": "Failed to open video file"}
            
            # Get video properties
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = frame_count / fps if fps > 0 else 0
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            analysis = {
                "duration": duration,
                "fps": fps,
                "frame_count": frame_count,
                "resolution": {"width": width, "height": height},
                "frames_analyzed": 0
            }
            
            # Analyze a few key frames
            frame_indices = [0, frame_count // 4, frame_count // 2, 3 * frame_count // 4, frame_count - 1]
            frame_analyses = []
            
            for i, frame_idx in enumerate(frame_indices):
                if frame_idx >= frame_count:
                    continue
                    
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
                ret, frame = cap.read()
                
                if ret:
                    # Convert frame to PIL Image for analysis
                    pil_frame = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
                    
                    # Create MultimodalContent for frame
                    frame_content = MultimodalContent(
                        modality=ModalityType.IMAGE,
                        media_type=MediaType.IMAGE_JPEG
                    )
                    
                    # Convert frame to bytes
                    import io
                    frame_bytes = io.BytesIO()
                    pil_frame.save(frame_bytes, format='JPEG')
                    frame_content.binary_content = frame_bytes.getvalue()
                    
                    # Analyze frame
                    frame_analysis = await self.analyze_image(frame_content, instructions)
                    frame_analysis["timestamp"] = frame_idx / fps
                    frame_analyses.append(frame_analysis)
                    analysis["frames_analyzed"] += 1
            
            cap.release()
            analysis["frame_analyses"] = frame_analyses
            
            return analysis
            
        except Exception as e:
            return {"error": f"Failed to analyze video: {str(e)}"}
    
    async def analyze_audio(
        self,
        audio: MultimodalContent,
        instructions: Optional[str] = None
    ) -> Dict[str, Any]:
        """Analyze audio content"""
        if not self.has_audio:
            return {"error": "Audio processing libraries not available"}
        
        try:
            audio_path = audio.file_path
            if not audio_path:
                return {"error": "Audio file path required for analysis"}
            
            # Load audio with pydub
            audio_segment = AudioSegment.from_file(audio_path)
            
            analysis = {
                "duration": len(audio_segment) / 1000.0,  # Convert to seconds
                "sample_rate": audio_segment.frame_rate,
                "channels": audio_segment.channels,
                "frame_width": audio_segment.frame_width,
                "max_dBFS": audio_segment.max_dBFS,
                "rms": audio_segment.rms
            }
            
            # Speech recognition if requested
            if not instructions or "transcrib" in instructions.lower() or "speech" in instructions.lower():
                try:
                    # Convert to wav for speech recognition
                    wav_audio = audio_segment.export(format="wav")
                    
                    # Use speech recognition
                    recognizer = sr.Recognizer()
                    with sr.AudioFile(wav_audio) as source:
                        audio_data = recognizer.record(source)
                        try:
                            text = recognizer.recognize_google(audio_data)
                            analysis["transcription"] = text
                            analysis["transcription_confidence"] = "medium"
                        except sr.UnknownValueError:
                            analysis["transcription"] = ""
                            analysis["transcription_note"] = "Could not understand audio"
                        except sr.RequestError as e:
                            analysis["transcription_error"] = f"Recognition service error: {e}"
                
                except Exception as e:
                    analysis["transcription_error"] = str(e)
            
            return analysis
            
        except Exception as e:
            return {"error": f"Failed to analyze audio: {str(e)}"}
    
    async def analyze_document(
        self,
        document: MultimodalContent,
        instructions: Optional[str] = None
    ) -> Dict[str, Any]:
        """Analyze document content"""
        try:
            doc_path = document.file_path
            if not doc_path:
                return {"error": "Document file path required for analysis"}
            
            analysis = {
                "file_type": document.media_type.value if document.media_type else "unknown"
            }
            
            # Handle different document types
            if document.media_type == MediaType.DOCUMENT_PDF and self.has_pdf:
                analysis.update(await self._analyze_pdf(doc_path))
            elif document.media_type in [MediaType.DOCUMENT_DOC, MediaType.DOCUMENT_DOCX] and self.has_docx:
                analysis.update(await self._analyze_docx(doc_path))
            elif document.media_type == MediaType.TEXT_PLAIN:
                analysis.update(await self._analyze_text(doc_path))
            else:
                analysis["error"] = f"Unsupported document type: {document.media_type}"
            
            return analysis
            
        except Exception as e:
            return {"error": f"Failed to analyze document: {str(e)}"}
    
    async def _analyze_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Analyze PDF document"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                analysis = {
                    "page_count": len(pdf_reader.pages),
                    "has_metadata": bool(pdf_reader.metadata),
                    "metadata": dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                }
                
                # Extract text from first few pages
                text_content = ""
                for i, page in enumerate(pdf_reader.pages[:5]):  # First 5 pages
                    text_content += page.extract_text() + "\n"
                
                analysis["sample_text"] = text_content[:1000]  # First 1000 chars
                analysis["estimated_words"] = len(text_content.split())
                
                return analysis
                
        except Exception as e:
            return {"error": f"PDF analysis failed: {str(e)}"}
    
    async def _analyze_docx(self, docx_path: str) -> Dict[str, Any]:
        """Analyze DOCX document"""
        try:
            doc = docx.Document(docx_path)
            
            analysis = {
                "paragraph_count": len(doc.paragraphs),
                "has_tables": len(doc.tables) > 0,
                "table_count": len(doc.tables)
            }
            
            # Extract text
            text_content = ""
            for paragraph in doc.paragraphs[:10]:  # First 10 paragraphs
                text_content += paragraph.text + "\n"
            
            analysis["sample_text"] = text_content[:1000]
            analysis["estimated_words"] = len(text_content.split())
            
            return analysis
            
        except Exception as e:
            return {"error": f"DOCX analysis failed: {str(e)}"}
    
    async def _analyze_text(self, text_path: str) -> Dict[str, Any]:
        """Analyze plain text document"""
        try:
            with open(text_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            lines = content.split('\n')
            words = content.split()
            
            analysis = {
                "character_count": len(content),
                "word_count": len(words),
                "line_count": len(lines),
                "average_line_length": sum(len(line) for line in lines) / len(lines) if lines else 0,
                "sample_text": content[:1000]
            }
            
            return analysis
            
        except Exception as e:
            return {"error": f"Text analysis failed: {str(e)}"}
    
    async def extract_text(self, content: MultimodalContent) -> str:
        """Extract text from any supported content type"""
        try:
            if content.modality == ModalityType.TEXT:
                return content.get_content_text() or ""
            
            elif content.modality == ModalityType.IMAGE and self.has_ocr:
                content_bytes = content.get_content_bytes()
                if content_bytes:
                    pil_image = Image.open(io.BytesIO(content_bytes))
                    return pytesseract.image_to_string(pil_image).strip()
            
            elif content.modality == ModalityType.DOCUMENT:
                if content.media_type == MediaType.DOCUMENT_PDF and self.has_pdf:
                    return await self._extract_pdf_text(content.file_path)
                elif content.media_type in [MediaType.DOCUMENT_DOC, MediaType.DOCUMENT_DOCX] and self.has_docx:
                    return await self._extract_docx_text(content.file_path)
                elif content.media_type == MediaType.TEXT_PLAIN:
                    with open(content.file_path, 'r', encoding='utf-8') as f:
                        return f.read()
            
            elif content.modality == ModalityType.AUDIO and self.has_audio:
                # This would be implemented with speech recognition
                analysis = await self.analyze_audio(content)
                return analysis.get("transcription", "")
            
            return ""
            
        except Exception as e:
            logger.error(f"Failed to extract text: {e}")
            return ""
    
    async def _extract_pdf_text(self, pdf_path: str) -> str:
        """Extract all text from PDF"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return ""
    
    async def _extract_docx_text(self, docx_path: str) -> str:
        """Extract all text from DOCX"""
        try:
            doc = docx.Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            return ""
    
    async def cross_modal_reasoning(
        self,
        contents: List[MultimodalContent],
        question: str
    ) -> Dict[str, Any]:
        """Perform cross-modal reasoning across multiple content types"""
        # This is a basic implementation - in practice, this would involve
        # sophisticated AI models that can reason across modalities
        
        reasoning_results = {
            "question": question,
            "content_types": [content.modality.value for content in contents],
            "analysis_method": "basic_correlation",
            "findings": []
        }
        
        # Extract text from all content
        all_text = []
        for content in contents:
            text = await self.extract_text(content)
            if text:
                all_text.append({
                    "content_id": content.content_id,
                    "modality": content.modality.value,
                    "text": text[:500]  # First 500 chars
                })
        
        reasoning_results["extracted_texts"] = all_text
        
        # Basic cross-modal insights
        if len(contents) > 1:
            modalities = set(content.modality for content in contents)
            reasoning_results["findings"].append(
                f"Content spans {len(modalities)} different modalities: {[m.value for m in modalities]}"
            )
            
            if ModalityType.IMAGE in modalities and ModalityType.TEXT in modalities:
                reasoning_results["findings"].append(
                    "Image and text content detected - potential for visual-textual correlation analysis"
                )
        
        return reasoning_results
    
    async def _perform_cross_modal_analysis(
        self,
        contents: List[MultimodalContent],
        instructions: Optional[str]
    ) -> Dict[str, Any]:
        """Perform comprehensive cross-modal analysis"""
        if not instructions:
            instructions = "Analyze relationships between the provided content"
        
        return await self.cross_modal_reasoning(contents, instructions)


# Import statement fix
import io
