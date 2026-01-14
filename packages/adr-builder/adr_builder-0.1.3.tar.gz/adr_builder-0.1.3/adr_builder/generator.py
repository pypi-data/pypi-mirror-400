from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

from .fs import ensure_dir, format_index, next_index, slugify
from .models import CriteriaModel


@lru_cache(maxsize=8)
def _load_env(template_dir: str) -> Environment:
    """Load Jinja2 environment with caching for performance."""
    env = Environment(
        loader=FileSystemLoader(template_dir),
        autoescape=select_autoescape(enabled_extensions=("html", "xml")),
        trim_blocks=True,
        lstrip_blocks=True,
    )
    return env


def get_default_template_dir() -> Path:
    return Path(__file__).parent / "templates"


@dataclass
class AdrOutputInfo:
    """Holds computed output information for ADR generation."""
    index: int
    slug: str
    adr_dir: Path

    def md_path(self) -> Path:
        return self.adr_dir / f"{format_index(self.index)}-{self.slug}.md"

    def docx_path(self) -> Path:
        return self.adr_dir / f"{format_index(self.index)}-{self.slug}.docx"


def prepare_adr_output(criteria: CriteriaModel, docs_dir: Path) -> AdrOutputInfo:
    """
    Prepare output info for ADR generation. Call this once before generating
    multiple formats to ensure consistent index and slug across all outputs.
    """
    adr_dir = docs_dir / "adr"
    ensure_dir(adr_dir)
    index = next_index(adr_dir)
    slug = slugify(criteria.title)
    return AdrOutputInfo(index=index, slug=slug, adr_dir=adr_dir)


def render_madr(criteria: CriteriaModel, template_path: Path | None = None) -> str:
    template_dir = template_path.parent if template_path else get_default_template_dir()
    template_name = template_path.name if template_path else "madr.md.j2"
    env = _load_env(str(template_dir))  # str for lru_cache hashability
    tmpl = env.get_template(template_name)

    render_data = criteria.model_dump()
    if not render_data.get("date"):
        render_data["date"] = date.today().isoformat()
    # Back-compat naming expected by the template
    render_data["deciders"] = ", ".join(criteria.authors) if criteria.authors else ""

    return tmpl.render(**render_data)


def write_adr(
    criteria: CriteriaModel,
    docs_dir: Path,
    template: Path | None = None,
    output_info: AdrOutputInfo | None = None,
) -> Path:
    """
    Write ADR as Markdown file.

    Args:
        criteria: The ADR criteria model
        docs_dir: Base docs directory (adr/ will be appended)
        template: Optional custom Jinja2 template path
        output_info: Pre-computed output info (use to ensure consistent indexing)
    """
    if output_info is None:
        output_info = prepare_adr_output(criteria, docs_dir)

    output_path = output_info.md_path()
    content = render_madr(criteria, template)
    output_path.write_text(content, encoding="utf-8")
    return output_path


def write_adr_docx(
    criteria: CriteriaModel,
    docs_dir: Path,
    output_info: AdrOutputInfo | None = None,
) -> Path:
    """
    Write ADR as Word document.

    Args:
        criteria: The ADR criteria model
        docs_dir: Base docs directory (adr/ will be appended)
        output_info: Pre-computed output info (use to ensure consistent indexing)
    """
    if output_info is None:
        output_info = prepare_adr_output(criteria, docs_dir)

    output_path = output_info.docx_path()

    doc = Document()

    # Title
    title_para = doc.add_paragraph(criteria.title)
    title_para.style = doc.styles['Title']
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Metadata
    meta = doc.add_paragraph()
    meta.style = doc.styles['Normal']
    meta_runs = [
        f"Status: {criteria.status}",
        f"Date: {criteria.date or date.today().isoformat()}",
    ]
    if criteria.authors:
        meta_runs.append(f"Deciders: {', '.join(criteria.authors)}")
    if criteria.tags:
        meta_runs.append(f"Tags: {', '.join(criteria.tags)}")
    meta.add_run("\n".join(meta_runs))

    # Section helper
    def add_heading(text: str) -> None:
        h = doc.add_heading(level=1)
        run = h.add_run(text)
        run.font.size = Pt(14)
        run.bold = True

    def add_bullets(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(str(item), style='List Bullet')

    # Context
    add_heading("Context and Problem Statement")
    if criteria.context and criteria.context.background:
        doc.add_paragraph(criteria.context.background)
    else:
        doc.add_paragraph("Describe the problem and context here.")
    if criteria.context and criteria.context.constraints:
        doc.add_paragraph("Constraints:")
        add_bullets(criteria.context.constraints)
    if criteria.context and criteria.context.drivers:
        doc.add_paragraph("Decision Drivers:")
        add_bullets(criteria.context.drivers)

    # Considered Options
    add_heading("Considered Options")
    if criteria.options:
        for idx, opt in enumerate(criteria.options, start=1):
            doc.add_paragraph(f"{idx}. {opt.name}")
    else:
        doc.add_paragraph("1. Option A")
        doc.add_paragraph("2. Option B")

    # Decision Outcome
    add_heading("Decision Outcome")
    if criteria.decision:
        doc.add_paragraph(f"Chosen option: \"{criteria.decision.chosen}\"")
        if criteria.decision.rationale:
            doc.add_paragraph(f"Rationale: {criteria.decision.rationale}")
    else:
        doc.add_paragraph("Chosen option: \"<fill in>\"")

    # Pros and Cons
    add_heading("Pros and Cons of the Options")
    if criteria.options:
        for opt in criteria.options:
            doc.add_paragraph(f"- {opt.name}")
            if opt.pros:
                doc.add_paragraph(f"Good: {', '.join(opt.pros)}", style='List Bullet')
            if opt.cons:
                doc.add_paragraph(f"Bad: {', '.join(opt.cons)}", style='List Bullet')
            if opt.risks:
                doc.add_paragraph(f"Risks: {', '.join(opt.risks)}", style='List Bullet')
    else:
        doc.add_paragraph("- Option A")
        doc.add_paragraph("- Option B")

    # Consequences
    add_heading("Consequences")
    if criteria.consequences:
        if criteria.consequences.positive:
            doc.add_paragraph(f"Positive: {', '.join(criteria.consequences.positive)}")
        if criteria.consequences.negative:
            doc.add_paragraph(f"Negative: {', '.join(criteria.consequences.negative)}")
    else:
        doc.add_paragraph("- Positive:")
        doc.add_paragraph("- Negative:")

    # References
    add_heading("References")
    if criteria.references and criteria.references.links:
        add_bullets(criteria.references.links)
    else:
        doc.add_paragraph("- <links or issue URLs>")

    doc.save(str(output_path))
    return output_path


