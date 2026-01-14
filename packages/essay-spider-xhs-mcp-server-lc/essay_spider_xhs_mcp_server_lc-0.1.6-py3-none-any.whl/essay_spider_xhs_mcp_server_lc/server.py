from typing import Any, Dict, List, Generator
import httpx
from mcp.server.fastmcp import FastMCP
import os
from openai import OpenAI
from pathlib import Path
from datetime import datetime
from docx import Document

# Initialize FastMCP server
mcp = FastMCP("essay_spider_xhs_mcp_server_lc")

# Constants
NWS_API_BASE = "http://bitouessay.com:44451"

xhs_cookies = os.getenv("COOKIES", "")
word_save_dir = os.getenv("WORD_SAVE_DIR", "")


async def make_nws_request(url: str, data: dict) -> dict[str, Any] | None:
    """Make a request to the NWS API with proper error handling."""
    headers = {
        "Accept": "application/json"
    }
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(url, params=data, headers=headers, timeout=30.0)
            response.raise_for_status()
            return response.json()
        except Exception:
            return None


def list_note_format_alert(note: dict) -> str:
    #     """Format an alert feature into a readable string."""
    return f"""
用户: {note.get('note_type', '')}
帖子链接: {note.get('url', '')}
帖子类型: {note.get('severity', '暂无')}
帖子标题: {note.get('display_title', '暂无')}
帖子信息列表: {note.get('note_info_list', '暂无')}
"""


def item_note_format_alert(note: dict) -> str:
    return f"""
    帖子链接: {note.get('note_url', '')}
    帖子类型: {note.get('note_type', '暂无')}
    帖子标题: {note.get('title', '暂无')}
    帖子描述: {note.get('desc', '暂无')}
    帖子喜欢数：{note.get('liked_count', '0')}
    帖子评论数：{note.get('comment_count', '0')}
    帖子收藏数：{note.get('collected_count', '0')}
    """



@mcp.tool()
async def spider_note(note_url: str) -> str | dict[str, Any] | None:
    """获取小红书一篇帖子内容

    Args:
        :param cookies: 小红书Cookies
        :param note_url: 小红书帖子链接
    """
    if not xhs_cookies or len(xhs_cookies) < 10:  # 简单验证
        raise ValueError("无效的cookies格式，请提供有效的小红书cookies")
    url = f"{NWS_API_BASE}/notes/item"
    data = {'note_url': note_url, 'cookies_str': xhs_cookies}
    result = await make_nws_request(url, data)
    if not result or "info" not in result:
        return "爬取失败，请检查cookies或者小红书帖子是否正确"

    if not result["info"]:
        return "爬取失败，请检查cookies或者小红书帖子是否正确"
    # 将图片数组组合成字符串
    media = ''
    if result['info']['note_type'] == '图集':
        media = ";\n".join(result["info"]["image_list"])
        media = '帖子图片列表：' + media
    elif result['info']['note_type'] == '视频':
        media = '帖子视频需要单独获取'
    return item_note_format_alert(result['info']) + media


@mcp.tool()
async def spider_user_notes(user_url: str) -> str | list[str]:
    """获取用户下的所有帖子

    Args:
        :param user_url: 用户主页链接
        :param cookies: 小红书Cookies
    """
    if not xhs_cookies or len(xhs_cookies) < 10:  # 简单验证
        raise ValueError("无效的cookies格式，请提供有效的小红书cookies")
    url = f"{NWS_API_BASE}/user/item"
    data = {'user_url': user_url, 'cookies_str': xhs_cookies}
    result = await make_nws_request(url, data)
    if not result or "list" not in result:
        return "爬取失败，请检查cookies或者小红书帖子是否正确"

    if not result["list"]:
        return "爬取失败，请检查cookies或者小红书帖子是否正确"
    return [list_note_format_alert(note) for note in result['list']]


@mcp.tool(name="get_video_src", description="通过对小红书的链接获取到里面的视频链接，note_url: 小红书链接")
async def get_video_src(note_url: str) -> str | dict[str, Any] | None:
    url = f"{NWS_API_BASE}/notes/video"
    data = {'note_url': note_url}
    result = await make_nws_request(url, data)
    if not result or "video" not in result:
        return "爬取失败，请检查小红书帖子链接是否正确"

    if not result["video"]:
        return "爬取失败，请检查小红书帖子链接是否正确"
    return result


@mcp.tool()
async def spider_notes_with_comments(note_urls: list[str]) -> list[str]:
    """获取帖子内容和所有一级评论并直接生成 Word 文档，返回文件路径列表"""
    if not xhs_cookies or len(xhs_cookies) < 10:
        raise ValueError("无效的cookies格式，请提供有效的小红书cookies")
    if not word_save_dir:
        raise ValueError("请先在环境变量WORD_SAVE_DIR中配置文件保存路径")
    base_path = Path(word_save_dir)
    base_path.mkdir(parents=True, exist_ok=True)
    url = f"{NWS_API_BASE}/notes/item_with_comments"
    results: list[str] = []
    for note_url in note_urls:
        data = {"note_url": note_url, "cookies_str": xhs_cookies}
        result = await make_nws_request(url, data)
        if not result or "info" not in result or "comments" not in result:
            results.append("爬取失败，请检查cookies或者小红书帖子是否正确")
        else:
            filepath = _save_single_note_to_word(result, base_path)
            results.append(filepath)
    return results


def _save_single_note_to_word(single_note: dict[str, Any], base_path: Path) -> str:
    info = single_note.get("info", {})
    comments = single_note.get("comments", [])
    title = info.get("title", "无标题")
    filename = f"{title}.docx"
    filepath = base_path / filename
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(f"帖子链接: {info.get('note_url', '')}")
    document.add_paragraph(f"用户昵称: {info.get('nickname', '')}")
    document.add_paragraph(f"发布时间: {info.get('upload_time', '')}")
    document.add_paragraph(f"IP属地: {info.get('ip_location', '')}")
    document.add_paragraph("")
    document.add_paragraph("帖子描述:")
    document.add_paragraph(info.get("desc", ""))
    images = info.get("image_list") or []
    if images:
        document.add_paragraph("")
        document.add_paragraph("图片列表:")
        for image in images:
            document.add_paragraph(image)
    document.add_paragraph("")
    document.add_paragraph("评论列表:")
    for index, comment in enumerate(comments, start=1):
        document.add_paragraph(
            f"{index}. {comment.get('nickname', '')} ({comment.get('upload_time', '')} {comment.get('ip_location', '')})"
        )
        document.add_paragraph(f"内容: {comment.get('content', '')}")
        pictures = comment.get("pictures") or []
        if pictures:
            document.add_paragraph("图片:")
            for picture in pictures:
                document.add_paragraph(picture)
        document.add_paragraph(f"点赞数: {comment.get('like_count', '0')}")
        document.add_paragraph("")
    document.save(str(filepath))
    return str(filepath)


def run():
    mcp.run()


if __name__ == "__main__":
    run()
