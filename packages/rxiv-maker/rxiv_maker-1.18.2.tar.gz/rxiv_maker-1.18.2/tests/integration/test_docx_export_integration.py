"""Integration tests for DOCX export."""

from pathlib import Path

import pytest
from docx import Document

from rxiv_maker.exporters.docx_exporter import DocxExporter


class TestDocxExportIntegration:
    """Integration tests for complete DOCX export workflow."""

    @pytest.fixture
    def sample_manuscript_path(self):
        """Return path to test fixtures."""
        fixtures_dir = Path(__file__).parent.parent / "fixtures" / "docx_export" / "sample_manuscript"
        return str(fixtures_dir)

    def test_full_export_workflow(self, sample_manuscript_path, tmp_path):
        """Test complete export process from manuscript to DOCX."""
        exporter = DocxExporter(
            manuscript_path=sample_manuscript_path,
            resolve_dois=False,
            include_footnotes=True,
        )

        result_path = exporter.export()

        # Verify file was created
        assert result_path.exists()
        assert result_path.suffix == ".docx"

        # Verify DOCX is valid
        doc = Document(str(result_path))
        assert len(doc.paragraphs) > 0

        # Verify content (first H1 "Abstract" is skipped as title)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Introduction" in all_text
        assert "Methods" in all_text
        assert "Results" in all_text
        # Check for abstract content (not the heading which is skipped)
        assert "sample manuscript for testing DOCX export" in all_text

    def test_export_with_citations(self, sample_manuscript_path, tmp_path):
        """Test that citations are properly converted to numbers."""
        exporter = DocxExporter(
            manuscript_path=sample_manuscript_path,
        )

        result_path = exporter.export()
        doc = Document(str(result_path))
        all_text = "\n".join([p.text for p in doc.paragraphs])

        # Check for numbered citations
        assert "[1]" in all_text
        assert "[2]" in all_text
        assert "[3]" in all_text

        # Should not contain @ citations
        assert "@smith2021" not in all_text
        assert "@jones2022" not in all_text

    def test_export_includes_supplementary(self, sample_manuscript_path, tmp_path):
        """Test that supplementary information is included."""
        exporter = DocxExporter(
            manuscript_path=sample_manuscript_path,
        )

        result_path = exporter.export()
        doc = Document(str(result_path))
        all_text = "\n".join([p.text for p in doc.paragraphs])

        # Check for supplementary content
        assert "Supplementary" in all_text

    def test_export_preserves_formatting(self, sample_manuscript_path, tmp_path):
        """Test that inline formatting is preserved."""
        exporter = DocxExporter(
            manuscript_path=sample_manuscript_path,
        )

        result_path = exporter.export()
        doc = Document(str(result_path))

        # Check for formatted runs
        has_bold = False
        has_italic = False

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    has_bold = True
                if run.italic:
                    has_italic = True

        assert has_bold, "Expected to find bold text"
        assert has_italic, "Expected to find italic text"

    def test_export_with_default_output_path(self, sample_manuscript_path, tmp_path):
        """Test export with default output path."""
        exporter = DocxExporter(
            manuscript_path=sample_manuscript_path,
            # No output_path specified - uses custom filename from metadata
        )

        result_path = exporter.export()

        # Should create file with custom name based on metadata
        # Format: YEAR__lastname_et_al__rxiv.docx
        assert result_path.exists()
        assert result_path.suffix == ".docx"
        assert "rxiv" in result_path.name  # Custom filename pattern
        # File should be in manuscript directory
        assert result_path.parent == Path(sample_manuscript_path)

    def test_export_validates_manuscript(self, tmp_path):
        """Test that missing manuscript files are detected."""
        # Create empty directory
        empty_dir = tmp_path / "empty"
        empty_dir.mkdir()

        exporter = DocxExporter(
            manuscript_path=str(empty_dir),
        )

        # Should raise FileNotFoundError for missing 01_MAIN.md
        with pytest.raises(FileNotFoundError, match="01_MAIN.md"):
            exporter.export()

    def test_export_requires_bibliography(self, tmp_path):
        """Test that missing bibliography file is detected."""
        # Create directory with only main file
        test_dir = tmp_path / "no_bib"
        test_dir.mkdir()

        main_file = test_dir / "01_MAIN.md"
        main_file.write_text("# Test\n\nContent here.")

        exporter = DocxExporter(
            manuscript_path=str(test_dir),
        )

        # Should raise FileNotFoundError for missing bibliography
        with pytest.raises(FileNotFoundError, match="03_REFERENCES.bib"):
            exporter.export()

    def test_export_reports_statistics(self, sample_manuscript_path, tmp_path, caplog):
        """Test that export reports citation statistics."""
        exporter = DocxExporter(
            manuscript_path=sample_manuscript_path,
        )

        exporter.export()

        # Check log messages
        log_text = caplog.text
        assert "unique citations" in log_text
        assert "DOIs found" in log_text

    def test_export_without_footnotes(self, sample_manuscript_path, tmp_path):
        """Test export with footnotes disabled."""
        exporter = DocxExporter(
            manuscript_path=sample_manuscript_path,
            include_footnotes=False,
        )

        result_path = exporter.export()

        # Should still create valid document
        assert result_path.exists()
        doc = Document(str(result_path))
        assert len(doc.paragraphs) > 0

    def test_export_handles_missing_citation_keys(self, tmp_path):
        """Test handling of citation keys not in bibliography."""
        # Create test manuscript with invalid citation
        test_dir = tmp_path / "invalid_cite"
        test_dir.mkdir()

        main_file = test_dir / "01_MAIN.md"
        main_file.write_text("# Test\n\nSome text with @invalid_citation.\n")

        bib_file = test_dir / "03_REFERENCES.bib"
        bib_file.write_text(
            """
@article{valid2023,
  author = {Author, A.},
  title = {Valid Entry},
  year = {2023}
}
"""
        )

        exporter = DocxExporter(
            manuscript_path=str(test_dir),
        )

        # Should complete without error (but log warning)
        result_path = exporter.export()
        assert result_path.exists()

    def test_export_handles_yaml_header(self, tmp_path):
        """Test that YAML frontmatter is properly removed."""
        test_dir = tmp_path / "with_yaml"
        test_dir.mkdir()

        main_file = test_dir / "01_MAIN.md"
        main_file.write_text(
            """---
title: Test Title
author: Test Author
---

# Content

This is the actual content.
"""
        )

        bib_file = test_dir / "03_REFERENCES.bib"
        bib_file.write_text(
            """
@article{test2023,
  author = {Test, A.},
  title = {Test},
  year = {2023}
}
"""
        )

        exporter = DocxExporter(
            manuscript_path=str(test_dir),
        )

        result_path = exporter.export()
        doc = Document(str(result_path))
        all_text = "\n".join([p.text for p in doc.paragraphs])

        # Should not contain YAML header
        assert "---" not in all_text
        assert "title: Test Title" not in all_text

        # Should contain actual content (first H1 "Content" is skipped as title)
        assert "This is the actual content" in all_text
