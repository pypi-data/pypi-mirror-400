"""Assemble multiple RTF files into a single RTF or DOCX file."""

import os
from collections.abc import Sequence
from copy import deepcopy
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:  # pragma: no cover
    from docx.document import Document as DocxDocument
    from docx.section import Section

# from .input import RTFPage  # Unused


def assemble_rtf(
    input_files: list[str],
    output_file: str,
) -> None:
    """Combine multiple RTF files into a single RTF file.

    Args:
        input_files: List of paths to RTF files to combine.
        output_file: Path to the output RTF file.
    """
    if not input_files:
        return

    # Check if files exist
    missing_files = [f for f in input_files if not os.path.exists(f)]
    if missing_files:
        raise FileNotFoundError(f"Missing files: {', '.join(missing_files)}")

    # Read all files
    rtf_contents = []
    for f in input_files:
        with open(f, encoding="utf-8") as file:
            rtf_contents.append(file.readlines())

    if not rtf_contents:
        return

    # Process first file
    # We keep everything from the first file except the last closing brace '}'

    # Remove last line if it contains only '}' or remove the last '}' char
    # r2rtf simply removes the last line: end[-n] <- end[-n] - 1

    # Helper to find start index based on fcharset
    def find_start_index(lines):
        last_idx = 0
        found = False
        for i, line in enumerate(lines):
            if "fcharset" in line:
                last_idx = i
                found = True

        if found:
            return last_idx + 2
        return 0

    new_page_cmd = r"\page" + "\n"

    processed_parts = []

    for i, lines in enumerate(rtf_contents):
        start_idx = 0
        if i > 0:
            # For subsequent files, skip header
            start_idx = find_start_index(lines)

        end_idx = len(lines)
        if i < len(rtf_contents) - 1 and lines[-1].strip() == "}":
            # Remove last line (closing brace) for all but last file
            end_idx -= 1

        part = lines[start_idx:end_idx]
        processed_parts.extend(part)

        if i < len(rtf_contents) - 1:
            processed_parts.append(new_page_cmd)

    # Write output
    with open(output_file, "w", encoding="utf-8") as outfile:
        outfile.writelines(processed_parts)


def assemble_docx(
    input_files: list[str],
    output_file: str,
    landscape: bool | list[bool] = False,
) -> None:
    """Combine multiple RTF files into a single DOCX file.

    Args:
        input_files: List of paths to input RTF files.
        output_file: Path to the output DOCX file.
        landscape: Whether the output should be landscape. Can be a single bool
            (applies to all) or a list of bools (one per file). Defaults to False.
    """
    try:
        import docx  # type: ignore
        from docx.enum.section import WD_ORIENT  # type: ignore
    except ImportError as e:
        raise ImportError(
            "python-docx is required for assemble_docx. "
            "Install it with: pip install 'rtflite[docx]'"
        ) from e

    if not input_files:
        raise ValueError("Input files list cannot be empty")

    # Check input files exist
    missing_files = [f for f in input_files if not os.path.exists(f)]
    if missing_files:
        raise FileNotFoundError(f"Missing files: {', '.join(missing_files)}")

    # Handle landscape argument
    if isinstance(landscape, bool):
        landscape_list = [landscape] * len(input_files)
    else:
        if len(landscape) != len(input_files):
            raise ValueError("Length of landscape list must match input files")
        landscape_list = landscape

    # Create new document
    doc = docx.Document()

    for i, (input_file, is_landscape) in enumerate(
        zip(input_files, landscape_list, strict=True)
    ):
        # Set orientation for the current section
        section = doc.sections[-1]
        if is_landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            w, h = section.page_width, section.page_height
            if w is not None and h is not None and w < h:  # If currently portrait
                section.page_width = h
                section.page_height = w
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            w, h = section.page_width, section.page_height
            if w is not None and h is not None and w > h:  # If currently landscape
                section.page_width = h
                section.page_height = w

        # Absolute path needed for fields
        abs_path = os.path.abspath(input_file)

        # Escape backslashes for the field code
        path_str = abs_path.replace("\\", "\\\\")

        # Create INCLUDETEXT field
        field_code = f'INCLUDETEXT "{path_str}"'

        # Add "Table X" caption
        p = doc.add_paragraph()
        p.add_run("Table ")
        _add_field(p, r"SEQ Table \* ARABIC")
        p.add_run("\n")  # Linebreak

        # Add the INCLUDETEXT field
        _add_field(p, field_code)

        # Handle section breaks
        if i < len(input_files) - 1:
            doc.add_section()

    doc.save(output_file)


def _add_field(paragraph, field_code):
    """Add a complex field to a paragraph."""
    # This is low-level XML manipulation for python-docx to add fields
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml.shared import OxmlElement  # type: ignore

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_code
    r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "separate")
    r.append(fldChar)

    # Add placeholder text so the field is visible/clickable
    if "SEQ" in field_code:
        run = paragraph.add_run("1")
    else:
        run = paragraph.add_run("Error! Reference source not found.")

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)


def concatenate_docx(
    input_files: Sequence[str | os.PathLike[str]],
    output_file: str | os.PathLike[str],
    landscape: bool | Sequence[bool] = False,
) -> None:
    """Concatenate DOCX files without relying on Word field toggles.

    This helper is useful when `RTFDocument.write_docx` already produced DOCX
    files and you need to stitch them together into a single document that can
    be distributed without refreshing fields in Microsoft Word.

    Args:
        input_files: Ordered collection of DOCX file paths to combine. The
            first document becomes the base; subsequent documents are appended
            as new sections.
        output_file: Path to the combined DOCX file.
        landscape: Whether each appended section should be landscape. Accepts
            a single boolean applied to every section or a list/tuple matching
            ``input_files``.

    Raises:
        ImportError: If ``python-docx`` is not installed.
        ValueError: If ``input_files`` is empty or the ``landscape`` list length
            does not match ``input_files``.
        FileNotFoundError: If any input file is missing.
    """
    try:
        from docx import Document  # type: ignore
        from docx.enum.section import WD_SECTION  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for concatenate_docx. "
            "Install it with: pip install 'rtflite[docx]'"
        ) from exc

    paths = [Path(path).expanduser() for path in input_files]
    if not paths:
        raise ValueError("Input files list cannot be empty")

    missing_files = [str(path) for path in paths if not path.exists()]
    if missing_files:
        raise FileNotFoundError(f"Missing files: {', '.join(missing_files)}")

    orientation_flags = _coerce_landscape_flags(landscape, len(paths))

    combined_doc = Document(str(paths[0]))
    _set_section_orientation(combined_doc.sections[0], orientation_flags[0])

    for source_path, is_landscape in zip(paths[1:], orientation_flags[1:], strict=True):
        combined_doc.add_section(WD_SECTION.NEW_PAGE)
        _set_section_orientation(combined_doc.sections[-1], is_landscape)
        _append_document_body(combined_doc, Document(str(source_path)))

    output_path = Path(output_file).expanduser()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    combined_doc.save(str(output_path))


def _coerce_landscape_flags(
    landscape: bool | Sequence[bool],
    expected_length: int,
) -> list[bool]:
    """Normalize the ``landscape`` argument to a list and validate its length."""
    if isinstance(landscape, bool):
        return [landscape] * expected_length

    flags = list(landscape)
    if len(flags) != expected_length:
        raise ValueError("Length of landscape list must match input files")

    return flags


def _set_section_orientation(section: "Section", landscape: bool) -> None:
    """Set section orientation and swap dimensions if needed."""
    from docx.enum.section import WD_ORIENT  # type: ignore

    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    width, height = section.page_width, section.page_height
    if width is None or height is None:
        return

    if (landscape and width < height) or (not landscape and width > height):
        section.page_width, section.page_height = height, width


def _append_document_body(target: "DocxDocument", source: "DocxDocument") -> None:
    """Copy body content from ``source`` into ``target`` without section props."""
    for element in list(source.element.body):
        if element.tag.endswith("}sectPr"):
            continue
        target.element.body.append(deepcopy(element))
