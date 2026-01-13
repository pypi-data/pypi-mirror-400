"""
Word文档水印处理器 - 支持.docx格式添加明水印

注意：仅支持.docx格式（Office Open XML）
      不支持老的.doc格式（Office 97-2003二进制格式）
      如需处理.doc文件，请先转换为.docx格式
"""

import os
from pathlib import Path
from typing import Optional, Union, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from visualwm.core.info import WatermarkInfo
from visualwm.core.style import WatermarkStyle, RiskLevel


class WordWatermark:
    """
    Word文档水印处理器
    
    支持在Word文档(.docx)中添加明水印，包括：
    - 页眉/页脚水印
    - 对角线文字水印
    - 背景水印图层
    
    Features:
        - 支持.docx格式（Office 2007及以上）
        - 不支持.doc格式（需先转换为.docx）
        - 多种水印位置
        - 自定义样式
        - 批量处理
    
    Note:
        老的.doc格式（Office 97-2003）是微软专有的二进制格式，
        无法用纯Python处理。如需处理.doc文件，请使用以下方式转换：
        - Microsoft Word: 另存为 .docx
        - LibreOffice: 打开后另存为 .docx
        - 在线工具: 如 CloudConvert
    """
    
    SUPPORTED_EXTENSIONS = {'.docx'}
    
    def __init__(self, default_style: Optional[WatermarkStyle] = None):
        """
        初始化Word水印处理器
        
        Args:
            default_style: 默认水印样式
        """
        self.default_style = default_style or WatermarkStyle.default()
    
    def embed(
        self,
        input_path: Union[str, Path],
        output_path: Union[str, Path],
        info: WatermarkInfo,
        style: Optional[WatermarkStyle] = None,
        watermark_type: str = "header_footer"
    ) -> str:
        """
        在Word文档中嵌入水印
        
        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径
            info: 水印信息
            style: 水印样式
            watermark_type: 水印类型 ("header_footer", "diagonal", "background")
            
        Returns:
            输出文件路径
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        style = style or self.default_style
        
        if not input_path.exists():
            raise FileNotFoundError(f"输入文件不存在: {input_path}")
        
        # 检查文件格式
        if input_path.suffix.lower() == '.doc':
            raise ValueError(
                f"不支持.doc格式（Office 97-2003二进制格式）\n"
                f"请先将文件转换为.docx格式：\n"
                f"  - Microsoft Word: 文件 > 另存为 > .docx\n"
                f"  - LibreOffice: 文件 > 另存为 > .docx\n"
                f"  - 在线工具: CloudConvert, Zamzar等"
            )
        
        if input_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {input_path.suffix}，仅支持: {self.SUPPORTED_EXTENSIONS}")
        
        # 打开文档
        doc = Document(input_path)
        
        # 根据类型添加水印
        if watermark_type == "header_footer":
            self._add_header_footer_watermark(doc, info, style)
        elif watermark_type == "diagonal":
            self._add_diagonal_watermark(doc, info, style)
        elif watermark_type == "background":
            self._add_background_watermark(doc, info, style)
        else:
            raise ValueError(f"不支持的水印类型: {watermark_type}")
        
        # 保存文档
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        return str(output_path)
    
    def embed_new(
        self,
        output_path: Union[str, Path],
        content: str,
        info: WatermarkInfo,
        style: Optional[WatermarkStyle] = None,
    ) -> str:
        """
        创建带水印的新Word文档
        
        Args:
            output_path: 输出文档路径
            content: 文档内容
            info: 水印信息
            style: 水印样式
            
        Returns:
            输出文件路径
        """
        output_path = Path(output_path)
        style = style or self.default_style
        
        # 创建新文档
        doc = Document()
        
        # 添加内容
        for paragraph in content.split('\n'):
            doc.add_paragraph(paragraph)
        
        # 添加水印
        self._add_header_footer_watermark(doc, info, style)
        
        # 保存
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        return str(output_path)
    
    def _add_header_footer_watermark(
        self,
        doc: Document,
        info: WatermarkInfo,
        style: WatermarkStyle
    ):
        """添加页眉页脚水印"""
        watermark_text = self._format_watermark_text(info, style)
        
        # 获取或创建页眉页脚
        for section in doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.clear()
            
            # 添加水印文本
            run = header_para.add_run(watermark_text)
            run.font.size = Pt(style.font_size * 0.5)  # 页眉字体稍小
            run.font.color.rgb = RGBColor(*style.color)
            run.font.bold = style.bold
            run.font.italic = style.italic
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 页脚
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            
            run = footer_para.add_run(watermark_text)
            run.font.size = Pt(style.font_size * 0.5)
            run.font.color.rgb = RGBColor(*style.color)
            run.font.bold = style.bold
            run.font.italic = style.italic
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_diagonal_watermark(
        self,
        doc: Document,
        info: WatermarkInfo,
        style: WatermarkStyle
    ):
        """添加对角线文字水印（在页眉中实现）"""
        watermark_text = self._format_watermark_text(info, style)
        
        # 为每个section添加水印
        for section in doc.sections:
            self._add_text_watermark_to_header(section, watermark_text, style)
    
    def _add_text_watermark_to_header(self, section, text: str, style: WatermarkStyle):
        """在页眉中添加文字水印"""
        from lxml import etree
        
        # 获取页眉
        header = section.header
        header.is_linked_to_previous = False
        
        # 清除现有内容
        for para in header.paragraphs:
            para.clear()
        
        # VML命名空间
        NSMAP = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'v': 'urn:schemas-microsoft-com:vml',
            'o': 'urn:schemas-microsoft-com:office:office',
            'w10': 'urn:schemas-microsoft-com:office:word',
        }
        
        # 创建水印XML
        r_elem = etree.SubElement(header.paragraphs[0]._p, '{%s}r' % NSMAP['w'])
        
        pict = etree.SubElement(r_elem, '{%s}pict' % NSMAP['w'])
        
        # 创建VML shape
        shape = etree.SubElement(pict, '{%s}shape' % NSMAP['v'], nsmap={'v': NSMAP['v'], 'o': NSMAP['o']})
        shape.set('id', 'PowerPlusWaterMarkObject')
        shape.set('type', '#_x0000_t136')
        shape.set('style', 
            f'position:absolute;'
            f'margin-left:0;margin-top:0;'
            f'width:500pt;height:150pt;'
            f'rotation:{style.rotation};'
            f'z-index:-251657216;'
            f'mso-position-horizontal:center;'
            f'mso-position-horizontal-relative:margin;'
            f'mso-position-vertical:center;'
            f'mso-position-vertical-relative:margin'
        )
        shape.set('fillcolor', f'#{style.color[0]:02x}{style.color[1]:02x}{style.color[2]:02x}')
        shape.set('stroked', 'f')
        shape.set('{%s}allowincell' % NSMAP['o'], 'f')
        
        # 添加fill
        fill = etree.SubElement(shape, '{%s}fill' % NSMAP['v'])
        fill.set('opacity', f'{style.opacity}')
        
        # 添加textpath
        textpath = etree.SubElement(shape, '{%s}textpath' % NSMAP['v'])
        textpath.set('style', f'font-family:"{style.font_name}";font-size:{style.font_size}pt')
        textpath.set('string', text)
    
    def _add_background_watermark(
        self,
        doc: Document,
        info: WatermarkInfo,
        style: WatermarkStyle
    ):
        """添加背景水印（每页重复）"""
        # 使用页眉页脚实现背景效果
        self._add_header_footer_watermark(doc, info, style)
        
        # 在文档开头添加水印提示
        watermark_text = self._format_watermark_text(info, style)
        
        # 在每个段落前插入水印
        # 注意：这会影响文档结构，仅用于演示
        first_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        
        # 添加水印段落
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = watermark_para.add_run(f"【{style.prefix or '文档水印'}】{watermark_text}")
        run.font.size = Pt(style.font_size)
        run.font.color.rgb = RGBColor(*style.color)
        run.font.bold = style.bold
        
        # 移动到开头
        first_para._element.addprevious(watermark_para._element)
    
    def _format_watermark_text(self, info: WatermarkInfo, style: WatermarkStyle) -> str:
        """格式化水印文本"""
        text = info.to_string(separator=" | ", compact=True)
        
        if style.prefix:
            text = f"{style.prefix} {text}"
        if style.suffix:
            text = f"{text} {style.suffix}"
        
        return text
    
    def extract(
        self,
        doc_path: Union[str, Path]
    ) -> Optional[WatermarkInfo]:
        """
        从Word文档中提取水印信息
        
        Args:
            doc_path: 文档路径
            
        Returns:
            提取的水印信息
        """
        doc = Document(doc_path)
        
        # 从页眉提取
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                text = para.text.strip()
                if text and ("用户:" in text or "工号:" in text):
                    return WatermarkInfo.from_string(text)
            
            # 从页脚提取
            footer = section.footer
            for para in footer.paragraphs:
                text = para.text.strip()
                if text and ("用户:" in text or "工号:" in text):
                    return WatermarkInfo.from_string(text)
        
        # 从正文提取
        for para in doc.paragraphs:
            text = para.text.strip()
            if "用户:" in text or "工号:" in text:
                return WatermarkInfo.from_string(text)
        
        return None
    
    def batch_embed(
        self,
        input_dir: Union[str, Path],
        output_dir: Union[str, Path],
        info: WatermarkInfo,
        style: Optional[WatermarkStyle] = None,
        recursive: bool = False
    ) -> List[str]:
        """
        批量添加水印
        
        Args:
            input_dir: 输入目录
            output_dir: 输出目录
            info: 水印信息
            style: 水印样式
            recursive: 是否递归处理
            
        Returns:
            处理的文件列表
        """
        input_dir = Path(input_dir)
        output_dir = Path(output_dir)
        
        processed = []
        
        pattern = "**/*.docx" if recursive else "*.docx"
        
        for input_file in input_dir.glob(pattern):
            if input_file.name.startswith('~$'):  # 跳过临时文件
                continue
            
            relative_path = input_file.relative_to(input_dir)
            output_file = output_dir / relative_path
            
            try:
                self.embed(input_file, output_file, info, style)
                processed.append(str(output_file))
            except Exception as e:
                print(f"处理失败 {input_file}: {e}")
        
        return processed
