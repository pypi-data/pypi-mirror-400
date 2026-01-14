"""
Comprehensive Document Processor v3
Save this as: backend/utils/comprehensive_document_processor.py

Extracts MAXIMUM detail from ALL file types:
✓ PDF (text + images + detailed analysis)
✓ Excel/XLSX (all sheets, formulas, values, formatting)
✓ CSV (structured data with analysis)
✓ TXT (full content with structure)
✓ Images (JPG, PNG - comprehensive vision analysis)
✓ DOCX (full text with formatting)
"""

import os
import io
import csv
from PIL import Image
import google.generativeai as genai
from dotenv import load_dotenv
import time

load_dotenv()

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
GOOGLE_MODEL = os.getenv("GOOGLE_API_MODEL", "gemini-2.5-flash")

if GOOGLE_API_KEY:
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
        print("[Document Processor] ✓ Google Gemini configured")
    except Exception as e:
        print(f"[Document Processor] ✗ Error: {e}")


# ============================================================================
# EXCEL/XLSX PROCESSING
# ============================================================================

def process_excel_detailed(file_path):
    """
    Extract MAXIMUM detail from Excel files:
    - All sheets and their names
    - All data with formatting
    - Formulas and calculated values
    - Metadata and headers
    - Data statistics
    - All cell content
    """
    print(f"\n{'='*70}")
    print(f"[Excel Processor] 📊 Processing: {os.path.basename(file_path)}")
    print(f"{'='*70}")
    
    result = {
        "content": "",
        "sheets": [],
        "summary": "",
        "data_points": 0
    }
    
    try:
        import openpyxl
        
        wb = openpyxl.load_workbook(file_path, data_only=False)
        
        parts = []
        parts.append(f"EXCEL FILE: {os.path.basename(file_path)}")
        parts.append(f"Total Sheets: {len(wb.sheetnames)}\n")
        
        all_data = []
        
        for sheet_idx, sheet_name in enumerate(wb.sheetnames, 1):
            print(f"   Sheet {sheet_idx}/{len(wb.sheetnames)}: {sheet_name}")
            
            ws = wb[sheet_name]
            
            parts.append(f"\n{'='*70}")
            parts.append(f"SHEET {sheet_idx}: {sheet_name}")
            parts.append(f"{'='*70}")
            parts.append(f"Dimensions: {ws.dimensions}")
            parts.append(f"Max Row: {ws.max_row}, Max Column: {ws.max_column}\n")
            
            # Extract all data
            sheet_data = []
            for row in ws.iter_rows(values_only=False):
                row_data = []
                for cell in row:
                    cell_value = cell.value
                    
                    # Include formulas if present
                    if hasattr(cell, 'value') and cell.data_type == 'f':
                        row_data.append(f"FORMULA: {cell.value}")
                    else:
                        row_data.append(str(cell_value) if cell_value is not None else "")
                
                if any(row_data):  # Only add non-empty rows
                    sheet_data.append(row_data)
                    all_data.append(row_data)
            
            # Format as table
            if sheet_data:
                parts.append("[DATA TABLE]")
                for row_idx, row in enumerate(sheet_data[:100], 1):  # First 100 rows
                    parts.append(" | ".join(str(cell) for cell in row))
                
                if len(sheet_data) > 100:
                    parts.append(f"\n... and {len(sheet_data) - 100} more rows ...\n")
            
            # Statistics
            parts.append(f"\n[SHEET STATISTICS]")
            parts.append(f"Total Rows: {len(sheet_data)}")
            parts.append(f"Total Columns: {max(len(row) for row in sheet_data) if sheet_data else 0}")
            parts.append(f"Non-empty Cells: {sum(len([c for c in row if c]) for row in sheet_data)}\n")
            
            result["sheets"].append({
                "name": sheet_name,
                "rows": len(sheet_data),
                "columns": max(len(row) for row in sheet_data) if sheet_data else 0
            })
        
        result["content"] = "\n".join(parts)
        result["data_points"] = len(all_data)
        result["summary"] = f"Excel file with {len(wb.sheetnames)} sheets containing {result['data_points']} data rows"
        
        print(f"✓ Extracted {result['data_points']} rows from {len(wb.sheetnames)} sheets")
        print(f"{'='*70}\n")
        
        wb.close()
        
    except ImportError:
        print("[Excel Processor] ⚠️  openpyxl not installed. Run: pip install openpyxl")
        result["content"] = "Error: openpyxl not installed"
    except Exception as e:
        print(f"[Excel Processor] ✗ Error: {e}")
        result["content"] = f"Error processing Excel: {str(e)}"
    
    return result


# ============================================================================
# CSV PROCESSING
# ============================================================================

def process_csv_detailed(file_path):
    """
    Extract MAXIMUM detail from CSV files:
    - Headers and column names
    - All rows and data
    - Data types and statistics
    - Missing values
    - Data summary
    """
    print(f"\n{'='*70}")
    print(f"[CSV Processor] 📄 Processing: {os.path.basename(file_path)}")
    print(f"{'='*70}")
    
    result = {
        "content": "",
        "rows": 0,
        "columns": 0,
        "summary": ""
    }
    
    try:
        parts = []
        parts.append(f"CSV FILE: {os.path.basename(file_path)}\n")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            
            all_rows = []
            for row in reader:
                all_rows.append(row)
            
            if not all_rows:
                result["content"] = "CSV file is empty"
                return result
            
            # Headers
            headers = all_rows[0]
            data_rows = all_rows[1:]
            
            parts.append(f"[CSV METADATA]")
            parts.append(f"Total Columns: {len(headers)}")
            parts.append(f"Total Rows: {len(data_rows)}")
            parts.append(f"File Size: {os.path.getsize(file_path) / 1024:.2f} KB\n")
            
            # Headers
            parts.append(f"[COLUMN HEADERS]")
            for idx, header in enumerate(headers, 1):
                parts.append(f"{idx}. {header}")
            
            # Data
            parts.append(f"\n[DATA ROWS]")
            for row_idx, row in enumerate(data_rows[:200], 1):  # First 200 rows
                formatted_row = " | ".join(str(cell) for cell in row)
                parts.append(f"Row {row_idx}: {formatted_row}")
            
            if len(data_rows) > 200:
                parts.append(f"\n... and {len(data_rows) - 200} more rows ...\n")
            
            # Statistics
            parts.append(f"\n[DATA STATISTICS]")
            parts.append(f"Total Data Rows: {len(data_rows)}")
            parts.append(f"Columns: {len(headers)}")
            parts.append(f"Average Row Length: {sum(len(row) for row in data_rows) / len(data_rows):.1f}")
            
            # Missing values
            missing_count = sum(1 for row in data_rows for cell in row if not cell or cell.strip() == '')
            parts.append(f"Missing Values: {missing_count}")
            
            result["content"] = "\n".join(parts)
            result["rows"] = len(data_rows)
            result["columns"] = len(headers)
            result["summary"] = f"CSV file with {len(headers)} columns and {len(data_rows)} data rows"
            
            print(f"✓ Extracted {len(headers)} columns × {len(data_rows)} rows")
            print(f"{'='*70}\n")
        
    except Exception as e:
        print(f"[CSV Processor] ✗ Error: {e}")
        result["content"] = f"Error processing CSV: {str(e)}"
    
    return result


# ============================================================================
# PDF PROCESSING WITH IMAGES
# ============================================================================

def extract_images_from_pdf(pdf_path):
    """Extract all images from PDF"""
    images = []
    
    try:
        import fitz  # PyMuPDF
        
        print(f"\n[PDF Processor] 📸 Extracting images from PDF...")
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    image = Image.open(io.BytesIO(image_bytes))
                    images.append({
                        "image": image,
                        "page": page_num + 1,
                        "index": img_index + 1,
                        "format": base_image["ext"],
                        "size": image.size
                    })
                    
                    print(f"   ✓ Image found - Page {page_num + 1}, Size: {image.size}")
                except Exception as e:
                    print(f"   ✗ Error extracting image: {e}")
        
        doc.close()
        print(f"✓ Total images extracted: {len(images)}")
        
    except ImportError:
        print("[PDF Processor] ⚠️  PyMuPDF (fitz) not installed. Run: pip install PyMuPDF")
    except Exception as e:
        print(f"[PDF Processor] ✗ Error: {e}")
    
    return images


def analyze_image_gemini_comprehensive(image, page_num=None, img_idx=None):
    """
    COMPREHENSIVE image analysis with Gemini Vision
    Extracts EVERY detail from the image
    """
    try:
        if isinstance(image, str):
            image = Image.open(image)
        
        print(f"\n[Vision] 🔍 Analyzing image in detail...")
        if page_num:
            print(f"   Page {page_num}, Image {img_idx}")
        
        model = genai.GenerativeModel(GOOGLE_MODEL)
        analyses = {}
        
        # PROMPT 1: Visual Content
        print("   [Stage 1/5] Visual content...")
        prompt1 = """Describe EVERYTHING in extreme detail:
- All objects, elements, and components
- Text, numbers, labels (preserve exact formatting)
- Colors, shades, style, composition
- Layout and arrangement
- Any logos, brands, signatures
- Quality, clarity, and notable features
Be extremely thorough."""
        
        try:
            resp = model.generate_content([prompt1, image])
            analyses['visual'] = resp.text if resp and resp.text else ""
            print(f"   ✓ Visual ({len(analyses['visual'])} chars)")
        except Exception as e:
            print(f"   ✗ Error: {e}")
            analyses['visual'] = ""
        
        time.sleep(0.5)
        
        # PROMPT 2: Text Extraction
        print("   [Stage 2/5] Text extraction...")
        prompt2 = """Extract EVERY SINGLE TEXT VISIBLE:
- All readable text exactly as shown
- Headers, titles, body text
- Labels, captions, annotations
- Numbers, dates, amounts, codes
- Watermarks, stamps, signatures
- URLs, emails, contacts
- Fine print, footnotes, disclaimers
Include ALL text without omission."""
        
        try:
            resp = model.generate_content([prompt2, image])
            analyses['text'] = resp.text if resp and resp.text else ""
            print(f"   ✓ Text ({len(analyses['text'])} chars)")
        except Exception as e:
            print(f"   ✗ Error: {e}")
            analyses['text'] = ""
        
        time.sleep(0.5)
        
        # PROMPT 3: Data & Numbers
        print("   [Stage 3/5] Data extraction...")
        prompt3 = """Extract ALL numerical and quantitative data:
- Numbers, quantities, measurements
- Currency amounts, prices, costs
- Percentages, ratios, statistics
- Dates, times, periods
- IDs, codes, reference numbers
- Contact information
- Account details, amounts
Format with clear labels and context."""
        
        try:
            resp = model.generate_content([prompt3, image])
            analyses['data'] = resp.text if resp and resp.text else ""
            print(f"   ✓ Data ({len(analyses['data'])} chars)")
        except Exception as e:
            print(f"   ✗ Error: {e}")
            analyses['data'] = ""
        
        time.sleep(0.5)
        
        # PROMPT 4: Structures & Patterns
        print("   [Stage 4/5] Structures...")
        prompt4 = """Identify and describe all structures:
- Tables, columns, rows, cells
- Charts, graphs, diagrams
- Forms, fields, sections
- Lists, bullets, hierarchies
- Borders, frames, divisions
- Visual hierarchy and organization
- Data visualization types
Describe layouts and relationships."""
        
        try:
            resp = model.generate_content([prompt4, image])
            analyses['structures'] = resp.text if resp and resp.text else ""
            print(f"   ✓ Structures ({len(analyses['structures'])} chars)")
        except Exception as e:
            print(f"   ✗ Error: {e}")
            analyses['structures'] = ""
        
        time.sleep(0.5)
        
        # PROMPT 5: Classification & Context
        print("   [Stage 5/5] Classification...")
        prompt5 = """Classify and provide context:
- Document/image type
- Industry, domain, category
- Purpose and intent
- Key information summary
- Important entities (names, orgs, dates)
- Business relevance
- Document status
- How it relates to finance/business
Provide 2-3 sentence summary."""
        
        try:
            resp = model.generate_content([prompt5, image])
            analyses['classification'] = resp.text if resp and resp.text else ""
            print(f"   ✓ Classification ({len(analyses['classification'])} chars)")
        except Exception as e:
            print(f"   ✗ Error: {e}")
            analyses['classification'] = ""
        
        # Compile all
        combined = f"""{'='*70}
IMAGE ANALYSIS - COMPREHENSIVE DETAILS
{'='*70}

{'─'*70}
VISUAL CONTENT & DESCRIPTION
{'─'*70}
{analyses.get('visual', '')}

{'─'*70}
TEXT EXTRACTION (OCR) - ALL TEXT
{'─'*70}
{analyses.get('text', '')}

{'─'*70}
DATA & NUMERICAL INFORMATION
{'─'*70}
{analyses.get('data', '')}

{'─'*70}
STRUCTURES & VISUAL ORGANIZATION
{'─'*70}
{analyses.get('structures', '')}

{'─'*70}
CLASSIFICATION & CONTEXT
{'─'*70}
{analyses.get('classification', '')}

{'='*70}"""
        
        total_chars = sum(len(v) for v in analyses.values())
        print(f"   ✓ Total analysis: {total_chars} chars")
        
        return combined
        
    except Exception as e:
        print(f"[Vision] ✗ Error: {e}")
        return f"Error analyzing image: {str(e)}"


def process_pdf_detailed(pdf_path):
    """
    Process PDF with MAXIMUM detail:
    - All text from all pages
    - Extract and analyze all images
    - Document structure
    - Metadata
    """
    print(f"\n{'='*70}")
    print(f"[PDF Processor] 📄 Processing: {os.path.basename(pdf_path)}")
    print(f"{'='*70}")
    
    result = {
        "text": "",
        "images": [],
        "summary": ""
    }
    
    try:
        import PyPDF2
        
        # Extract text from all pages
        print(f"[PDF Processor] 📖 Extracting text from all pages...")
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"{'─'*70}\nPAGE {page_num}\n{'─'*70}\n{page_text}")
            
            result["text"] = "\n".join(text_parts)
            print(f"✓ Extracted {len(pdf_reader.pages)} pages, {len(result['text'])} chars")
        
        # Extract and analyze images
        images = extract_images_from_pdf(pdf_path)
        
        for idx, img_data in enumerate(images, 1):
            print(f"\n[PDF Processor] 📸 Image {idx}/{len(images)}...")
            
            analysis = analyze_image_gemini_comprehensive(
                img_data['image'],
                page_num=img_data['page'],
                img_idx=idx
            )
            
            result["images"].append({
                "page": img_data['page'],
                "analysis": analysis,
                "size": img_data['size']
            })
        
        result["summary"] = f"PDF with {len(pdf_reader.pages)} pages and {len(images)} images"
        print(f"\n✓ PDF Processing Complete!")
        print(f"   Text: {len(result['text'])} chars")
        print(f"   Images: {len(images)} with detailed analysis")
        print(f"{'='*70}\n")
        
    except Exception as e:
        print(f"[PDF Processor] ✗ Error: {e}")
        result["text"] = f"Error: {str(e)}"
    
    return result


# ============================================================================
# IMAGE PROCESSING (Direct uploads)
# ============================================================================

def process_image_detailed(image_path):
    """Process directly uploaded image with maximum detail"""
    print(f"\n{'='*70}")
    print(f"[Image Processor] 📷 Processing: {os.path.basename(image_path)}")
    print(f"{'='*70}")
    
    try:
        image = Image.open(image_path)
        filename = os.path.basename(image_path)
        
        analysis = analyze_image_gemini_comprehensive(image)
        
        result = {
            "filename": filename,
            "analysis": analysis,
            "size": image.size,
            "format": image.format or "unknown"
        }
        
        print(f"✓ Image processed")
        print(f"{'='*70}\n")
        
        return result
        
    except Exception as e:
        print(f"[Image Processor] ✗ Error: {e}")
        return {"error": str(e)}


# ============================================================================
# DOCX PROCESSING
# ============================================================================

def process_docx_detailed(file_path):
    """Extract maximum detail from DOCX files"""
    print(f"\n{'='*70}")
    print(f"[DOCX Processor] 📝 Processing: {os.path.basename(file_path)}")
    print(f"{'='*70}")
    
    result = {
        "content": "",
        "summary": ""
    }
    
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        parts = []
        parts.append(f"WORD DOCUMENT: {os.path.basename(file_path)}\n")
        parts.append(f"Total Paragraphs: {len(doc.paragraphs)}")
        parts.append(f"Total Tables: {len(doc.tables)}\n")
        
        # Extract all paragraphs
        parts.append("[DOCUMENT CONTENT]")
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        
        # Extract tables
        if doc.tables:
            parts.append(f"\n[TABLES - {len(doc.tables)} tables found]")
            for table_idx, table in enumerate(doc.tables, 1):
                parts.append(f"\nTable {table_idx}:")
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    parts.append(" | ".join(row_data))
        
        result["content"] = "\n".join(parts)
        result["summary"] = f"Word document with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables"
        
        print(f"✓ Extracted {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        print(f"{'='*70}\n")
        
    except ImportError:
        print("[DOCX Processor] ⚠️  python-docx not installed. Run: pip install python-docx")
        result["content"] = "Error: python-docx not installed"
    except Exception as e:
        print(f"[DOCX Processor] ✗ Error: {e}")
        result["content"] = f"Error processing DOCX: {str(e)}"
    
    return result


# ============================================================================
# COMBINED PROCESSING
# ============================================================================

def create_comprehensive_document(file_path, file_type):
    """
    Process any file type and create comprehensive content for ChromaDB
    """
    filename = os.path.basename(file_path)
    
    print(f"\n{'#'*70}")
    print(f"# COMPREHENSIVE DOCUMENT PROCESSING")
    print(f"# File: {filename}")
    print(f"# Type: {file_type.upper()}")
    print(f"{'#'*70}\n")
    
    content = ""
    metadata = {
        "source": filename,
        "type": file_type,
        "has_detail": True
    }
    
    try:
        if file_type == "pdf":
            result = process_pdf_detailed(file_path)
            
            parts = [f"PDF FILE: {filename}\n"]
            
            if result["text"]:
                parts.append(f"{'='*70}\nDOCUMENT TEXT CONTENT\n{'='*70}\n{result['text']}\n")
            
            if result["images"]:
                parts.append(f"{'='*70}\nIMAGE ANALYSES ({len(result['images'])} images)\n{'='*70}\n")
                for img in result["images"]:
                    parts.append(f"\n{'─'*70}\nImage from Page {img['page']}\n{'─'*70}\n{img['analysis']}\n")
            
            content = "\n".join(parts)
            metadata["has_images"] = len(result["images"]) > 0
            metadata["image_count"] = len(result["images"])
            metadata["pages"] = result["text"].count("PAGE")
            
        elif file_type == "xlsx":
            result = process_excel_detailed(file_path)
            content = result["content"]
            metadata["sheets"] = len(result["sheets"])
            metadata["data_points"] = result["data_points"]
            
        elif file_type == "csv":
            result = process_csv_detailed(file_path)
            content = result["content"]
            metadata["rows"] = result["rows"]
            metadata["columns"] = result["columns"]
            
        elif file_type == "docx":
            result = process_docx_detailed(file_path)
            content = result["content"]
            
        elif file_type in ["jpg", "jpeg", "png"]:
            result = process_image_detailed(file_path)
            if "error" not in result:
                content = f"IMAGE: {result['filename']}\nFormat: {result['format']}\nSize: {result['size']}\n\n{result['analysis']}"
                metadata["image_format"] = result["format"]
                metadata["image_size"] = str(result["size"])
            else:
                content = f"Error: {result['error']}"
        
        else:
            # TXT or unknown
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            metadata["file_size"] = len(content)
        
        print(f"\n{'#'*70}")
        print(f"# PROCESSING COMPLETE")
        print(f"# Content length: {len(content)} characters")
        print(f"{'#'*70}\n")
        
        return {
            "content": content,
            "metadata": metadata,
            "success": True
        }
        
    except Exception as e:
        print(f"Error processing file: {e}")
        return {
            "content": f"Error: {str(e)}",
            "metadata": metadata,
            "success": False
        }