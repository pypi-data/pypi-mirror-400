"""DOCX writer for rxiv-maker export.

This module handles the actual generation of DOCX files using python-docx,
writing structured content with formatting, citations, and references.
"""

import base64
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree

from ..core.logging_config import get_logger
from ..utils.author_affiliation_processor import AuthorAffiliationProcessor
from ..utils.docx_helpers import convert_pdf_to_image

logger = get_logger()


class DocxWriter:
    """Writes structured content to DOCX files using python-docx."""

    # Color mapping for different reference types
    XREF_COLORS = {
        "fig": WD_COLOR_INDEX.BRIGHT_GREEN,  # Figures (bright green - lighter)
        "sfig": WD_COLOR_INDEX.TURQUOISE,  # Supplementary figures (turquoise - lighter cyan)
        "stable": WD_COLOR_INDEX.TURQUOISE,  # Supplementary tables (turquoise - lighter cyan)
        "table": WD_COLOR_INDEX.BLUE,  # Main tables
        "eq": WD_COLOR_INDEX.PINK,  # Equations (pink - lighter than violet, easier to read)
        "snote": WD_COLOR_INDEX.TURQUOISE,  # Supplementary notes (turquoise - lighter cyan)
        "cite": WD_COLOR_INDEX.YELLOW,  # Citations (yellow)
    }

    @staticmethod
    def get_xref_color(xref_type: str):
        """Get highlight color for a cross-reference type.

        Args:
            xref_type: Type of cross-reference (fig, sfig, stable, table, eq, snote, cite)

        Returns:
            WD_COLOR_INDEX color for the xref type, or YELLOW as default
        """
        return DocxWriter.XREF_COLORS.get(xref_type, WD_COLOR_INDEX.YELLOW)

    def write(
        self,
        doc_structure: Dict[str, Any],
        bibliography: Dict[int, Dict],
        output_path: Path,
        include_footnotes: bool = True,
        base_path: Optional[Path] = None,
        metadata: Optional[Dict[str, Any]] = None,
        table_map: Optional[Dict[str, int]] = None,
        equation_map: Optional[Dict[str, int]] = None,
        figures_at_end: bool = False,
        hide_highlighting: bool = False,
        hide_comments: bool = False,
    ) -> Path:
        """Write DOCX file from structured content.

        Args:
            doc_structure: Structured document with sections
            bibliography: Bibliography entries mapped by number
            output_path: Path where DOCX file should be saved
            include_footnotes: Whether to add DOI footnotes
            base_path: Base path for resolving relative figure paths
            metadata: Document metadata (title, authors, affiliations)
            table_map: Mapping from table labels to numbers (for supplementary tables)
            equation_map: Mapping from equation labels to numbers
            figures_at_end: Place main figures at end before SI/bibliography
            hide_highlighting: Disable colored highlighting on references and citations
            hide_comments: Exclude all comments (block and inline) from output

        Returns:
            Path to created DOCX file
        """
        self.base_path = base_path or Path.cwd()
        self.bibliography = bibliography
        self.include_footnotes = include_footnotes
        self.table_map = table_map or {}
        self.equation_map = equation_map or {}
        self.hide_highlighting = hide_highlighting
        self.hide_comments = hide_comments
        doc = Document()

        # Set default font to Arial for entire document
        self._set_default_font(doc, "Arial")

        # Add title and author information if metadata provided
        if metadata:
            self._add_title_page(doc, metadata)

        # Build figure numbering map (figures stay inline in text)
        figure_map = {}  # Maps label to number
        figure_counter = 0

        for section in doc_structure["sections"]:
            if section["type"] == "figure":
                figure_counter += 1
                label = section.get("label", "")
                if label:
                    figure_map[label] = figure_counter

        # Store figure map for use in text processing
        self.figure_map = figure_map

        # Collect main figures if figures_at_end is True
        collected_main_figures = []

        # Process each section
        figure_counter = 0
        sfigure_counter = 0
        for section in doc_structure["sections"]:
            if section["type"] == "figure":
                is_supplementary = section.get("is_supplementary", False)
                if is_supplementary:
                    # Supplementary figures always go inline (in SI section)
                    sfigure_counter += 1
                    self._add_figure(doc, section, figure_number=sfigure_counter, is_supplementary=True)
                else:
                    # Main figures: collect if figures_at_end, otherwise add inline
                    figure_counter += 1
                    if figures_at_end:
                        collected_main_figures.append((section, figure_counter))
                    else:
                        self._add_figure(doc, section, figure_number=figure_counter, is_supplementary=False)
            else:
                self._add_section(doc, section, bibliography, include_footnotes)

        # Add collected main figures at the end (before bibliography)
        if figures_at_end and collected_main_figures:
            doc.add_page_break()
            heading = doc.add_heading("Figures", level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black text
            for section, fig_num in collected_main_figures:
                self._add_figure(doc, section, figure_number=fig_num, is_supplementary=False)

        # Add bibliography section at the end
        if include_footnotes and bibliography:
            doc.add_page_break()
            heading = doc.add_heading("Bibliography", level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black text

            # Add numbered bibliography entries
            for num in sorted(bibliography.keys()):
                bib_entry = bibliography[num]
                para = doc.add_paragraph()

                # Add citation number in bold
                num_run = para.add_run(f"[{num}] ")
                num_run.bold = True

                # Add formatted bibliography text (without DOI - added separately below)
                para.add_run(bib_entry["formatted"])

                # Add DOI as hyperlink with yellow highlighting if present (unless hide_highlighting is enabled)
                if bib_entry.get("doi"):
                    doi = bib_entry["doi"]
                    doi_url = f"https://doi.org/{doi}" if not doi.startswith("http") else doi
                    para.add_run("\nDOI: ")
                    self._add_hyperlink(para, doi_url, doi_url, highlight=not self.hide_highlighting)

                # Add spacing between entries
                para.paragraph_format.space_after = Pt(6)

        # Save document
        doc.save(str(output_path))
        return output_path

    def _set_default_font(self, doc: Document, font_name: str):
        """Set the default font for the entire document.

        Args:
            doc: Document object
            font_name: Font name to use (e.g., "Arial", "Times New Roman")
        """
        # Set font on Normal style (base style for most content)
        style = doc.styles["Normal"]
        font = style.font
        font.name = font_name
        font.size = Pt(10)  # Default body font size

        # Also set on heading styles to ensure consistency
        for i in range(1, 10):
            try:
                heading_style = doc.styles[f"Heading {i}"]
                heading_style.font.name = font_name
            except KeyError:
                # Heading style doesn't exist, skip
                pass

    def _apply_highlight(self, run, color: WD_COLOR_INDEX):
        """Apply highlight color to a run, unless highlighting is disabled.

        Args:
            run: The run object to apply highlighting to
            color: The WD_COLOR_INDEX color to apply
        """
        if not self.hide_highlighting:
            run.font.highlight_color = color

    def _add_title_page(self, doc: Document, metadata: Dict[str, Any]):
        """Add title, author and affiliation information.

        Args:
            doc: Document object
            metadata: Metadata dictionary with title and authors
        """
        # Add title first
        title = metadata.get("title", "")
        if title:
            # Title can be a string, a dict with 'long' key, or a list
            if isinstance(title, dict):
                title_text = title.get("long", title.get("short", ""))
            elif isinstance(title, list) and len(title) > 0:
                if isinstance(title[0], dict):
                    title_text = title[0].get("long", title[0].get("short", ""))
                else:
                    title_text = str(title[0])
            else:
                title_text = str(title)

            if title_text and isinstance(title_text, str):
                # Clean LaTeX formatting from title for DOCX
                title_text = self._clean_latex_from_text(title_text)

                title_para = doc.add_paragraph(title_text)
                title_para.runs[0].font.size = Pt(16)
                title_para.runs[0].bold = True
                title_para.paragraph_format.space_after = Pt(12)

        # Then add author and affiliation info
        authors = metadata.get("authors", [])
        if not authors:
            return  # Nothing more to add

        # Process author and affiliation metadata using centralized processor
        processor = AuthorAffiliationProcessor()
        processed = processor.process(metadata)

        affiliation_map = processed["affiliation_map"]
        ordered_affiliations = processed["ordered_affiliations"]
        cofirst_authors = processed["cofirst_authors"]
        corresponding_authors = processed["corresponding_authors"]

        # Add authors with superscript affiliation numbers and corresponding author markers
        if authors:
            author_para = doc.add_paragraph()
            for i, author in enumerate(authors):
                if i > 0:
                    author_para.add_run(", ")

                # Add author name
                name = author.get("name", "")
                author_para.add_run(name)

                # Add superscript affiliation numbers
                author_affils = author.get("affiliations", [])
                if author_affils:
                    affil_nums = [str(affiliation_map[a]) for a in author_affils]
                    sup_run = author_para.add_run(",".join(affil_nums))
                    sup_run.font.superscript = True

                # Add co-first author marker (dagger) if applicable
                is_cofirst = author.get("co_first_author", False)
                if is_cofirst:
                    cofirst_run = author_para.add_run("†")
                    cofirst_run.font.superscript = True

                # Add corresponding author marker (asterisk) if applicable
                is_corresponding = author.get("corresponding_author", False)
                if is_corresponding:
                    corr_run = author_para.add_run("*")
                    corr_run.font.superscript = True

            author_para.paragraph_format.space_after = Pt(8)

        # Add affiliations
        if ordered_affiliations:
            for affil_num, _affil_shortname, affil_text in ordered_affiliations:
                affil_para = doc.add_paragraph()

                # Add superscript number
                num_run = affil_para.add_run(str(affil_num))
                num_run.font.superscript = True
                num_run.font.size = Pt(8)

                # Add affiliation text
                affil_run = affil_para.add_run(f" {affil_text}")
                affil_run.font.size = Pt(8)
                affil_para.paragraph_format.space_after = Pt(4)

            # Extra space after last affiliation
            affil_para.paragraph_format.space_after = Pt(12)

        # Add co-first author information if any (already extracted by processor)
        if cofirst_authors:
            cofirst_para = doc.add_paragraph()
            cofirst_marker = cofirst_para.add_run("†")
            cofirst_marker.font.superscript = True
            cofirst_marker.font.size = Pt(8)

            cofirst_label = cofirst_para.add_run(" These authors contributed equally: ")
            cofirst_label.font.size = Pt(8)

            for i, author in enumerate(cofirst_authors):
                if i > 0:
                    sep_run = cofirst_para.add_run(", ")
                    sep_run.font.size = Pt(8)

                name = author.get("name", "")
                name_run = cofirst_para.add_run(name)
                name_run.font.size = Pt(8)

            cofirst_para.paragraph_format.space_after = Pt(12)

        # Add corresponding author information if any (already extracted by processor)
        if corresponding_authors:
            corr_para = doc.add_paragraph()
            corr_marker = corr_para.add_run("*")
            corr_marker.font.superscript = True
            corr_marker.font.size = Pt(8)

            corr_label = corr_para.add_run(" Correspondence: ")
            corr_label.font.size = Pt(8)

            for i, author in enumerate(corresponding_authors):
                if i > 0:
                    sep_run = corr_para.add_run("; ")
                    sep_run.font.size = Pt(8)

                name = author.get("name", "")
                email = author.get("email", "")

                # Decode email if it's base64 encoded
                if not email:
                    email64 = author.get("email64", "")
                    if email64:
                        try:
                            email = base64.b64decode(email64).decode("utf-8")
                        except Exception:
                            email = ""

                if email:
                    info_run = corr_para.add_run(f"{name} ({email})")
                else:
                    info_run = corr_para.add_run(name)
                info_run.font.size = Pt(8)

            corr_para.paragraph_format.space_after = Pt(12)

    def _add_section(
        self,
        doc: Document,
        section: Dict[str, Any],
        bibliography: Dict[int, Dict],
        include_footnotes: bool,
    ):
        """Add a section to the document.

        Args:
            doc: Document object
            section: Section data
            bibliography: Bibliography entries
            include_footnotes: Whether to add footnotes
        """
        section_type = section["type"]

        if section_type == "heading":
            self._add_heading(doc, section)
        elif section_type == "snote_title":
            self._add_snote_title(doc, section)
        elif section_type == "paragraph":
            self._add_paragraph(doc, section, bibliography, include_footnotes)
        elif section_type == "list":
            self._add_list(doc, section)
        elif section_type == "code_block":
            self._add_code_block(doc, section)
        elif section_type == "comment":
            if not self.hide_comments:
                self._add_comment(doc, section)
        elif section_type == "figure":
            self._add_figure(doc, section)
        elif section_type == "table":
            self._add_table(doc, section)
        elif section_type == "equation":
            self._add_equation(doc, section)
        elif section_type == "page_break":
            self._add_page_break(doc)

    def _add_snote_title(self, doc: Document, section: Dict[str, Any]):
        """Add supplementary note title to document.

        Args:
            doc: Document object
            section: Supplementary note title with 'label' and 'text'
        """
        _label = section.get("label", "")  # Reserved for future use
        text = section.get("text", "")

        # Add as bold heading
        para = doc.add_heading(level=3)
        para.clear()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black text

    def _add_heading(self, doc: Document, section: Dict[str, Any]):
        """Add heading to document.

        Args:
            doc: Document object
            section: Heading section data with 'level' and 'text'
        """
        level = section["level"]
        text = section["text"]
        heading = doc.add_heading(text, level=level)
        # Ensure heading text is black (not blue)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Explicitly set to black

    def _add_paragraph(
        self,
        doc: Document,
        section: Dict[str, Any],
        bibliography: Dict[int, Dict],
        include_footnotes: bool,
    ):
        """Add paragraph with formatted runs to document.

        Args:
            doc: Document object
            section: Paragraph section data with 'runs'
            bibliography: Bibliography entries
            include_footnotes: Whether to add footnotes
        """
        paragraph = doc.add_paragraph()

        # Set justified alignment for all paragraphs
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        runs_data = section["runs"]

        for run_data in runs_data:
            self._add_run(paragraph, run_data, bibliography, include_footnotes)

    def _add_run(self, paragraph, run_data: Dict[str, Any], bibliography: Dict[int, Dict], include_footnotes: bool):
        """Add a single run to a paragraph.

        Args:
            paragraph: Paragraph object
            run_data: Run data with type and formatting
            bibliography: Bibliography entries
            include_footnotes: Whether to add footnotes
        """
        if run_data["type"] == "text":
            text = run_data["text"]
            run = paragraph.add_run(text)

            # Always set font size explicitly for consistency (body text is 10pt)
            run.font.size = Pt(10)

            # Apply formatting
            if run_data.get("bold"):
                run.bold = True
            if run_data.get("italic"):
                run.italic = True
            if run_data.get("underline"):
                run.underline = True
            if run_data.get("subscript"):
                run.font.subscript = True
            if run_data.get("superscript"):
                run.font.superscript = True
            if run_data.get("code"):
                run.font.name = "Courier New"
                # Font size already set to Pt(10) above
            if run_data.get("xref"):
                # Use color based on xref type (fig, sfig, stable, eq, etc.)
                xref_type = run_data.get("xref_type", "cite")
                self._apply_highlight(run, self.get_xref_color(xref_type))
            if run_data.get("highlight_yellow"):
                self._apply_highlight(run, WD_COLOR_INDEX.YELLOW)

        elif run_data["type"] == "hyperlink":
            # Add hyperlink with yellow highlighting
            text = run_data.get("text", "")
            url = run_data.get("url", "")
            self._add_hyperlink(paragraph, url, text, highlight=True)

        elif run_data["type"] == "inline_equation":
            # Add inline equation as Office Math
            latex_content = run_data.get("latex", "")
            self._add_inline_equation(paragraph, latex_content)

        elif run_data["type"] == "inline_comment":
            # Add inline comment with gray highlighting (unless hide_comments is enabled)
            if not self.hide_comments:
                comment_text = run_data["text"]
                run = paragraph.add_run(f"[Comment: {comment_text}]")
                self._apply_highlight(run, WD_COLOR_INDEX.GRAY_25)
                run.italic = True
                run.font.size = Pt(10)

        elif run_data["type"] == "citation":
            cite_num = run_data["number"]
            # Add citation as [NN] inline with yellow highlighting
            run = paragraph.add_run(f"[{cite_num}]")
            self._apply_highlight(run, WD_COLOR_INDEX.YELLOW)
            run.font.size = Pt(10)

    def _add_list(self, doc: Document, section: Dict[str, Any]):
        """Add list to document with inline formatting.

        Args:
            doc: Document object
            section: List section data with 'list_type' and 'items'
        """
        list_type = section["list_type"]
        items = section["items"]
        style = "List Bullet" if list_type == "bullet" else "List Number"

        for item_runs in items:
            # Create paragraph with list style
            paragraph = doc.add_paragraph(style=style)

            # Add each run with its formatting
            for run_data in item_runs:
                if run_data["type"] == "text":
                    text = run_data["text"]
                    run = paragraph.add_run(text)

                    # Always set font size explicitly for consistency (body text is 10pt)
                    run.font.size = Pt(10)

                    # Apply formatting
                    if run_data.get("bold"):
                        run.bold = True
                    if run_data.get("italic"):
                        run.italic = True
                    if run_data.get("subscript"):
                        run.font.subscript = True
                    if run_data.get("superscript"):
                        run.font.superscript = True
                    if run_data.get("code"):
                        run.font.name = "Courier New"
                        # Font size already set to Pt(10) above
                    if run_data.get("xref"):
                        # Use color based on xref type
                        xref_type = run_data.get("xref_type", "cite")
                        self._apply_highlight(run, self.get_xref_color(xref_type))
                    if run_data.get("highlight_yellow"):
                        self._apply_highlight(run, WD_COLOR_INDEX.YELLOW)
                elif run_data["type"] == "hyperlink":
                    text = run_data.get("text", "")
                    url = run_data.get("url", "")
                    self._add_hyperlink(paragraph, url, text, highlight=True)
                elif run_data["type"] == "inline_equation":
                    # Add inline equation as Office Math
                    latex_content = run_data.get("latex", "")
                    self._add_inline_equation(paragraph, latex_content)
                elif run_data["type"] == "inline_comment":
                    # Add inline comment with gray highlighting (unless hide_comments is enabled)
                    if not self.hide_comments:
                        comment_text = run_data["text"]
                        run = paragraph.add_run(f"[Comment: {comment_text}]")
                        self._apply_highlight(run, WD_COLOR_INDEX.GRAY_25)
                        run.italic = True
                        run.font.size = Pt(10)
                elif run_data["type"] == "citation":
                    cite_num = run_data["number"]
                    run = paragraph.add_run(f"[{cite_num}]")
                    run.bold = True
                    run.font.size = Pt(10)
                    self._apply_highlight(run, WD_COLOR_INDEX.YELLOW)

    def _add_code_block(self, doc: Document, section: Dict[str, Any]):
        """Add code block to document.

        Args:
            doc: Document object
            section: Code block section data with 'content'
        """
        code_content = section["content"]
        paragraph = doc.add_paragraph(code_content)

        # Style as code
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        # Set paragraph formatting
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(36)  # Indent code blocks

    def _add_comment(self, doc: Document, section: Dict[str, Any]):
        """Add comment to document with gray highlighting.

        Args:
            doc: Document object
            section: Comment section data with 'text'
        """
        comment_text = section["text"]
        paragraph = doc.add_paragraph()

        # Add comment text with light gray highlighting to distinguish from colored xrefs
        run = paragraph.add_run(f"[Comment: {comment_text}]")
        self._apply_highlight(run, WD_COLOR_INDEX.GRAY_25)
        run.italic = True
        run.font.size = Pt(10)

    def _check_poppler_availability(self) -> bool:
        """Check if poppler is available for PDF conversion.

        Returns:
            True if poppler is available, False otherwise
        """
        from ..core.managers.dependency_manager import DependencyStatus, get_dependency_manager

        manager = get_dependency_manager()
        result = manager.check_dependency("pdftoppm")

        return result.status == DependencyStatus.AVAILABLE

    def _add_figure(
        self, doc: Document, section: Dict[str, Any], figure_number: int = None, is_supplementary: bool = False
    ):
        """Add figure to document with caption.

        Args:
            doc: Document object
            section: Figure section data with 'path', 'caption', 'label'
            figure_number: Figure number (1-indexed)
            is_supplementary: Whether this is a supplementary figure
        """
        figure_path = Path(section["path"])
        caption = section.get("caption", "")

        # Resolve relative path
        if not figure_path.is_absolute():
            figure_path = self.base_path / figure_path

        # Handle different image types
        img_source = None

        if not figure_path.exists():
            logger.warning(f"Figure file not found: {figure_path}")
        elif figure_path.suffix.lower() == ".pdf":
            # Check poppler availability first (cached after first check)
            if not hasattr(self, "_poppler_checked"):
                self._poppler_available = self._check_poppler_availability()
                self._poppler_checked = True

                if not self._poppler_available:
                    logger.warning(
                        "Poppler not installed - PDF figures will be shown as placeholders. "
                        "Install with: brew install poppler (macOS) or sudo apt install poppler-utils (Linux)"
                    )

            if self._poppler_available:
                # Convert PDF to image
                try:
                    from pdf2image.exceptions import PDFInfoNotInstalledError, PopplerNotInstalledError

                    img_source = convert_pdf_to_image(figure_path)
                    logger.debug(f"  PDF converted: {img_source is not None}")
                except (PopplerNotInstalledError, PDFInfoNotInstalledError) as e:
                    logger.error(f"Poppler utilities not found: {e}")
                    img_source = None
                    # Update our cached status
                    self._poppler_available = False
            else:
                img_source = None
        elif figure_path.suffix.lower() in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]:
            # Use image file directly
            img_source = str(figure_path)
        else:
            logger.warning(f"Unsupported image format: {figure_path.suffix}")

        if img_source:
            # Add image with proper sizing to fit page
            try:
                from PIL import Image as PILImage

                # Get image dimensions
                with PILImage.open(img_source) as img:
                    img_width, img_height = img.size
                    aspect_ratio = img_width / img_height

                # Calculate available width from document section settings
                # Get the current section to read actual page dimensions and margins
                section = doc.sections[-1]  # Use the most recent section

                # Page width minus left and right margins
                available_width = section.page_width - section.left_margin - section.right_margin

                # Page height minus top and bottom margins
                available_height = section.page_height - section.top_margin - section.bottom_margin

                # Convert available width to Inches for comparison
                max_width = available_width
                max_height = available_height

                # Calculate aspect ratio thresholds
                page_aspect_ratio = available_width / available_height

                # Add figure centered
                # Note: add_picture() creates a paragraph automatically, but we need to add it explicitly
                # to control alignment
                fig_para = doc.add_paragraph()
                fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Calculate optimal size maintaining aspect ratio
                if aspect_ratio > page_aspect_ratio:  # Wide image - constrain by width
                    run = fig_para.add_run()
                    run.add_picture(img_source, width=max_width)
                else:  # Tall image - constrain by height
                    run = fig_para.add_run()
                    run.add_picture(img_source, height=max_height)

                logger.debug(f"Embedded figure: {figure_path} ({img_width}x{img_height})")
            except Exception as e:
                logger.warning(f"Failed to embed figure {figure_path}: {e}")
                # Add placeholder text (centered)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[Figure: {figure_path.name}]")
                run.italic = True
        else:
            # Add placeholder if embedding failed (centered)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[Figure: {figure_path.name}]")
            run.italic = True
            logger.warning(f"Could not embed figure: {figure_path}")

        # Add caption
        if caption:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Add small space before caption to separate from figure
            caption_para.paragraph_format.space_before = Pt(3)

            # Format as "Figure number: " or "Supp. Fig. number: "
            if figure_number:
                if is_supplementary:
                    run = caption_para.add_run(f"Supp. Fig. S{figure_number}. ")
                else:
                    run = caption_para.add_run(f"Fig. {figure_number}. ")
                run.bold = True
                run.font.size = Pt(8)
            else:
                run = caption_para.add_run("Figure: ")
                run.bold = True
                run.font.size = Pt(8)

            # Parse and add caption with inline formatting
            # Import the processor to parse inline formatting
            from rxiv_maker.exporters.docx_content_processor import DocxContentProcessor

            processor = DocxContentProcessor()
            caption_runs = processor._parse_inline_formatting(caption, {})

            for run_data in caption_runs:
                if run_data["type"] == "text":
                    text = run_data["text"]
                    run = caption_para.add_run(text)
                    run.font.size = Pt(8)

                    # Apply formatting
                    if run_data.get("bold"):
                        run.bold = True
                    if run_data.get("italic"):
                        run.italic = True
                    if run_data.get("subscript"):
                        run.font.subscript = True
                    if run_data.get("superscript"):
                        run.font.superscript = True
                    if run_data.get("code"):
                        run.font.name = "Courier New"
                    if run_data.get("xref"):
                        # Use color based on xref type
                        xref_type = run_data.get("xref_type", "cite")
                        self._apply_highlight(run, self.get_xref_color(xref_type))
                    if run_data.get("highlight_yellow"):
                        self._apply_highlight(run, WD_COLOR_INDEX.YELLOW)
                elif run_data["type"] == "inline_equation":
                    # Add inline equation as Office Math
                    latex_content = run_data.get("latex", "")
                    self._add_inline_equation(caption_para, latex_content)
                elif run_data["type"] == "inline_comment":
                    # Add inline comment with gray highlighting (unless hide_comments is enabled)
                    if not self.hide_comments:
                        comment_text = run_data["text"]
                        run = caption_para.add_run(f"[Comment: {comment_text}]")
                        self._apply_highlight(run, WD_COLOR_INDEX.GRAY_25)
                        run.italic = True
                        run.font.size = Pt(8)
                elif run_data["type"] == "citation":
                    cite_num = run_data["number"]
                    run = caption_para.add_run(f"[{cite_num}]")
                    run.bold = True
                    run.font.size = Pt(8)
                    self._apply_highlight(run, WD_COLOR_INDEX.YELLOW)

            # Add spacing after figure (reduced from 12 to 6 for compactness)
            caption_para.paragraph_format.space_after = Pt(6)

    def _add_table(self, doc: Document, section: Dict[str, Any]):
        """Add table to document.

        Args:
            doc: Document object
            section: Table section data with 'headers' and 'rows'
        """
        from rxiv_maker.exporters.docx_content_processor import DocxContentProcessor

        processor = DocxContentProcessor()

        headers = section["headers"]
        rows = section["rows"]

        # Create table
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Light Grid Accent 1"

        # Add header row with inline formatting
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            # Clear default paragraph
            header_cells[i].text = ""
            paragraph = header_cells[i].paragraphs[0]

            # Parse inline formatting
            header_runs = processor._parse_inline_formatting(header, {})

            for run_data in header_runs:
                if run_data["type"] == "text":
                    text = run_data["text"]
                    run = paragraph.add_run(text)
                    # Headers are always bold by default, but respect markdown formatting
                    run.bold = True if run_data.get("bold") else True  # All headers bold
                    if run_data.get("italic"):
                        run.italic = True
                    if run_data.get("code"):
                        run.font.name = "Courier New"

        # Add data rows with inline formatting
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row_cells):
                    # Clear default paragraph
                    row_cells[col_idx].text = ""
                    paragraph = row_cells[col_idx].paragraphs[0]

                    # Parse inline formatting
                    cell_runs = processor._parse_inline_formatting(cell_data, {})

                    for run_data in cell_runs:
                        if run_data["type"] == "text":
                            text = run_data["text"]
                            run = paragraph.add_run(text)
                            if run_data.get("bold"):
                                run.bold = True
                            if run_data.get("italic"):
                                run.italic = True
                            if run_data.get("underline"):
                                run.underline = True
                            if run_data.get("subscript"):
                                run.font.subscript = True
                            if run_data.get("superscript"):
                                run.font.superscript = True
                            if run_data.get("code"):
                                run.font.name = "Courier New"
                            if run_data.get("xref"):
                                # Use color based on xref type
                                xref_type = run_data.get("xref_type", "cite")
                                self._apply_highlight(run, self.get_xref_color(xref_type))

        # Add table caption if present
        caption = section.get("caption")
        label = section.get("label")
        if caption:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Add small space before caption to separate from table
            caption_para.paragraph_format.space_before = Pt(3)

            # Determine table number from label using table_map
            if label and label.startswith("stable:"):
                # Extract label name (e.g., "stable:parameters" -> "parameters")
                label_name = label.split(":", 1)[1] if ":" in label else label
                # Look up number in table_map
                table_num = self.table_map.get(label_name)
                if table_num:
                    run = caption_para.add_run(f"Supp. Table S{table_num}. ")
                else:
                    # Fallback if label not in map
                    run = caption_para.add_run("Supp. Table: ")
                run.bold = True
                run.font.size = Pt(8)
            elif label and label.startswith("table:"):
                # Extract label name for main tables
                label_name = label.split(":", 1)[1] if ":" in label else label
                # Look up number in table_map (though main tables may not be in map)
                table_num = self.table_map.get(label_name)
                if table_num:
                    run = caption_para.add_run(f"Table {table_num}. ")
                else:
                    run = caption_para.add_run("Table: ")
                run.bold = True
                run.font.size = Pt(8)

            # Parse and add caption with inline formatting
            caption_runs = processor._parse_inline_formatting(caption, {})
            for run_data in caption_runs:
                if run_data["type"] == "text":
                    text = run_data["text"]
                    run = caption_para.add_run(text)
                    run.font.size = Pt(8)
                    if run_data.get("bold"):
                        run.bold = True
                    if run_data.get("italic"):
                        run.italic = True
                    if run_data.get("underline"):
                        run.underline = True
                    if run_data.get("subscript"):
                        run.font.subscript = True
                    if run_data.get("superscript"):
                        run.font.superscript = True
                    if run_data.get("code"):
                        run.font.name = "Courier New"
                    if run_data.get("xref"):
                        # Use color based on xref type
                        xref_type = run_data.get("xref_type", "cite")
                        self._apply_highlight(run, self.get_xref_color(xref_type))

            # Add spacing after table (reduced from 12 to 6 for compactness)
            caption_para.paragraph_format.space_after = Pt(6)

        # Add spacing after table
        doc.add_paragraph()

    def _add_inline_equation(self, paragraph, latex_content: str):
        """Add inline equation to paragraph as Office Math.

        Args:
            paragraph: Paragraph object to add equation to
            latex_content: LaTeX equation content
        """
        if not latex_content:
            return

        try:
            # Convert LaTeX to MathML
            mathml_str = latex_to_mathml(latex_content)

            # Parse MathML
            mathml_root = etree.fromstring(mathml_str.encode("utf-8"))

            # Convert MathML to OMML
            omml_root = self._mathml_to_omml(mathml_root)

            # Insert the OMML element inline into the paragraph
            paragraph._element.append(omml_root)

        except Exception as e:
            logger.debug(f"Could not convert inline equation to OMML: {e}")
            # Fallback: add as italic text
            run = paragraph.add_run(latex_content)
            run.italic = True
            run.font.size = Pt(10)

    def _add_equation(self, doc: Document, section: Dict[str, Any]):
        """Add equation to document as rendered image with numbering.

        Args:
            doc: Document object
            section: Equation section data with 'content' (LaTeX) and optional 'label'
        """
        import re

        latex_content = section.get("content", "")
        label = section.get("label", "")  # e.g., "eq:bernoulli_labeling"

        if not latex_content:
            return

        # Get equation number from label
        equation_number = None
        if label:
            # Extract label name (e.g., "eq:bernoulli_labeling" -> "bernoulli_labeling")
            label_name = label.split(":", 1)[1] if ":" in label else label
            equation_number = self.equation_map.get(label_name)

        # Create a paragraph for the equation
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Try to render equation as image
        try:
            equation_img = self._render_equation_to_image(latex_content)
            if equation_img and equation_img.exists():
                # Add equation image with two-tier sizing
                run = para.add_run()
                import re

                from docx.shared import Inches

                # Detect if equation has fractions or other vertical elements
                has_fractions = bool(re.search(r"\\frac|\\dfrac", latex_content))

                # Simple two-tier sizing:
                # - Equations with fractions: 0.32 inches (taller for readability)
                # - Simple equations without fractions: 0.22 inches (consistent with text)
                if has_fractions:
                    run.add_picture(str(equation_img), height=Inches(0.32))
                    logger.debug("Equation with fractions: height=0.32in")
                else:
                    run.add_picture(str(equation_img), height=Inches(0.22))
                    logger.debug("Simple equation: height=0.22in")

                logger.info(f"Equation rendered as image: {equation_img}")
            else:
                # Fallback to formatted text
                logger.warning("Failed to render equation as image, using formatted text")
                self._render_latex_formatted(para, latex_content)
        except Exception as e:
            logger.warning(f"Equation image rendering failed: {e}, using formatted text")
            self._render_latex_formatted(para, latex_content)

        # Add equation number on the right side if available
        if equation_number:
            # Add tab stop for right alignment
            from docx.enum.text import WD_TAB_ALIGNMENT
            from docx.shared import Inches

            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

            # Add tab and equation number
            num_run = para.add_run(f"\t({equation_number})")
            num_run.font.size = Pt(11)

        logger.info("Equation successfully added to document")

        # Add spacing after equation
        para.paragraph_format.space_after = Pt(12)

    def _render_equation_to_image(self, latex_content: str) -> Optional[Path]:
        """Render LaTeX equation to PNG image.

        Args:
            latex_content: LaTeX equation content

        Returns:
            Path to generated PNG image, or None if rendering failed
        """
        import hashlib
        import subprocess
        import tempfile
        from pathlib import Path

        # Create a hash of the equation content for caching
        eq_hash = hashlib.md5(latex_content.encode()).hexdigest()[:8]

        # Use temp directory for equation images
        temp_dir = Path(tempfile.gettempdir()) / "rxiv_equations"
        temp_dir.mkdir(exist_ok=True)

        output_png = temp_dir / f"eq_{eq_hash}.png"

        # Check if we already rendered this equation
        if output_png.exists():
            logger.debug(f"Using cached equation image: {output_png}")
            return output_png

        # Create LaTeX document for standalone equation
        latex_template = r"""\documentclass[preview,border=2pt]{standalone}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{amsfonts}
\begin{document}
\begin{math}
\displaystyle EQUATION_CONTENT
\end{math}
\end{document}
"""
        latex_doc = latex_template.replace("EQUATION_CONTENT", latex_content)

        # Write to temporary .tex file
        tex_file = temp_dir / f"eq_{eq_hash}.tex"
        with open(tex_file, "w") as f:
            f.write(latex_doc)

        try:
            # Compile LaTeX to PDF
            result = subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "-output-directory", str(temp_dir), str(tex_file)],
                capture_output=True,
                text=True,
                timeout=10,
            )

            if result.returncode != 0:
                logger.warning(f"pdflatex failed for equation: {latex_content[:50]}...")
                logger.debug(f"pdflatex output: {result.stdout}")
                return None

            # Convert PDF to PNG using pdftoppm (high resolution for quality)
            pdf_file = temp_dir / f"eq_{eq_hash}.pdf"
            if not pdf_file.exists():
                logger.warning(f"PDF not generated: {pdf_file}")
                return None

            # Use pdftoppm to convert to PNG at high DPI (300 for quality)
            result = subprocess.run(
                ["pdftoppm", "-png", "-singlefile", "-r", "300", str(pdf_file), str(temp_dir / f"eq_{eq_hash}")],
                capture_output=True,
                text=True,
                timeout=10,
            )

            if result.returncode != 0:
                logger.warning(f"pdftoppm failed: {result.stderr}")
                return None

            if output_png.exists():
                logger.debug(f"Equation rendered to image: {output_png}")
                return output_png
            else:
                logger.warning(f"PNG not generated: {output_png}")
                return None

        except subprocess.TimeoutExpired:
            logger.warning("Equation rendering timed out")
            return None
        except FileNotFoundError as e:
            logger.warning(f"LaTeX/pdftoppm not found: {e}")
            return None
        except Exception as e:
            logger.warning(f"Equation rendering error: {e}")
            return None
        finally:
            # Clean up intermediate files
            for ext in [".tex", ".pdf", ".aux", ".log"]:
                temp_file = temp_dir / f"eq_{eq_hash}{ext}"
                if temp_file.exists():
                    try:
                        temp_file.unlink()
                    except Exception:
                        pass

    def _render_latex_formatted(self, paragraph, latex_content: str):
        """Render LaTeX equation with formatted subscripts and superscripts.

        Args:
            paragraph: Paragraph object to add formatted equation to
            latex_content: LaTeX equation content
        """
        import re

        # LaTeX symbol replacements
        symbol_map = {
            r"\\sim": "~",
            r"\\approx": "≈",
            r"\\le": "≤",
            r"\\ge": "≥",
            r"\\ne": "≠",
            r"\\times": "×",
            r"\\pm": "±",
            r"\\in": "∈",
            r"\\subset": "⊂",
            r"\\cap": "∩",
            r"\\cup": "∪",
            r"\\mid": "|",
            r"\\alpha": "α",
            r"\\beta": "β",
            r"\\gamma": "γ",
            r"\\delta": "δ",
            r"\\epsilon": "ε",
            r"\\theta": "θ",
            r"\\lambda": "λ",
            r"\\sigma": "σ",
            r"\\mu": "μ",
            r"\\pi": "π",
            r"\\sum": "Σ",
            r"\\prod": "Π",
            r"\\int": "∫",
            r"\\infty": "∞",
            r"\\mathbb\{I\}": "𝕀",
            r"\\mathbb\{R\}": "ℝ",
            r"\\mathbb\{N\}": "ℕ",
        }

        # Replace LaTeX commands with symbols
        content = latex_content
        for latex_cmd, symbol in symbol_map.items():
            content = re.sub(latex_cmd, symbol, content)

        # Remove \text{} wrappers but keep content
        content = re.sub(r"\\text\{([^}]+)\}", r"\1", content)

        # Remove \left and \right sizing commands
        content = re.sub(r"\\left|\\right", "", content)

        # Remove \mathbf{} but keep content (for now, just remove formatting)
        content = re.sub(r"\\mathbf\{([^}]+)\}", r"\1", content)

        # Parse and render with subscripts/superscripts
        # Pattern: text followed by _{subscript} or ^{superscript}
        i = 0
        while i < len(content):
            if content[i] == "_" and i + 1 < len(content):
                # Subscript
                if content[i + 1] == "{":
                    # Find matching brace
                    brace_end = content.find("}", i + 2)
                    if brace_end != -1:
                        sub_text = content[i + 2 : brace_end]
                        run = paragraph.add_run(sub_text)
                        run.font.subscript = True
                        run.font.size = Pt(11)
                        i = brace_end + 1
                        continue
                else:
                    # Single character subscript
                    run = paragraph.add_run(content[i + 1])
                    run.font.subscript = True
                    run.font.size = Pt(11)
                    i += 2
                    continue

            elif content[i] == "^" and i + 1 < len(content):
                # Superscript
                if content[i + 1] == "{":
                    # Find matching brace
                    brace_end = content.find("}", i + 2)
                    if brace_end != -1:
                        sup_text = content[i + 2 : brace_end]
                        run = paragraph.add_run(sup_text)
                        run.font.superscript = True
                        run.font.size = Pt(11)
                        i = brace_end + 1
                        continue
                else:
                    # Single character superscript
                    run = paragraph.add_run(content[i + 1])
                    run.font.superscript = True
                    run.font.size = Pt(11)
                    i += 2
                    continue

            # Regular text
            # Collect text until next special character
            text_chunk = ""
            while i < len(content) and content[i] not in ["_", "^", "\\"]:
                text_chunk += content[i]
                i += 1

            if text_chunk:
                run = paragraph.add_run(text_chunk)
                run.font.size = Pt(11)
                run.font.italic = True  # Math content is typically italic

            # Handle backslash commands we didn't replace
            if i < len(content) and content[i] == "\\":
                # Skip the backslash and next word
                i += 1
                while i < len(content) and content[i].isalpha():
                    i += 1

    def _mathml_to_omml(self, mathml_elem):
        """Convert MathML to OMML (Office Math Markup Language).

        Args:
            mathml_elem: MathML element (lxml element)

        Returns:
            OMML element (OxmlElement)
        """
        # For now, use basic OMML structure
        # A full implementation would use XSLT transformation
        return self._create_basic_omml(mathml_elem)

    def _create_basic_omml(self, mathml_elem):
        """Create a basic OMML structure when XSLT transformation fails.

        Args:
            mathml_elem: MathML element

        Returns:
            Basic OMML element
        """
        # Create oMath element (namespace will be handled by python-docx)
        oMath = OxmlElement("m:oMath")

        # Extract text content
        math_text = self._extract_mathml_text(mathml_elem)

        # Create a run with the math text
        r = OxmlElement("m:r")

        # Add text element
        t = OxmlElement("m:t")
        t.text = math_text
        r.append(t)

        oMath.append(r)

        return oMath

    def _extract_mathml_text(self, elem):
        """Extract text content from MathML element.

        Args:
            elem: MathML element

        Returns:
            String with text content
        """
        text_parts = []
        if elem.text:
            text_parts.append(elem.text)
        for child in elem:
            text_parts.append(self._extract_mathml_text(child))
            if child.tail:
                text_parts.append(child.tail)
        return "".join(text_parts)

    def _get_mml2omml_xslt(self):
        """Get simplified MathML to OMML XSLT transformation.

        Returns:
            XSLT stylesheet as string
        """
        # Simplified XSLT for basic MathML to OMML conversion
        # This is a minimal implementation - a full version would handle all MathML elements
        return """<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0"
    xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
    xmlns:mml="http://www.w3.org/1998/Math/MathML"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">

    <xsl:output method="xml" indent="no" encoding="UTF-8"/>

    <xsl:template match="/">
        <m:oMath>
            <xsl:apply-templates/>
        </m:oMath>
    </xsl:template>

    <xsl:template match="mml:math">
        <xsl:apply-templates/>
    </xsl:template>

    <xsl:template match="mml:mi|mml:mn|mml:mo|mml:mtext">
        <m:r>
            <m:t><xsl:value-of select="."/></m:t>
        </m:r>
    </xsl:template>

    <xsl:template match="mml:mrow">
        <xsl:apply-templates/>
    </xsl:template>

    <xsl:template match="mml:msub">
        <m:sSub>
            <m:e><xsl:apply-templates select="*[1]"/></m:e>
            <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
        </m:sSub>
    </xsl:template>

    <xsl:template match="mml:msup">
        <m:sSup>
            <m:e><xsl:apply-templates select="*[1]"/></m:e>
            <m:sup><xsl:apply-templates select="*[2]"/></m:sup>
        </m:sSup>
    </xsl:template>

    <xsl:template match="mml:mfrac">
        <m:f>
            <m:num><xsl:apply-templates select="*[1]"/></m:num>
            <m:den><xsl:apply-templates select="*[2]"/></m:den>
        </m:f>
    </xsl:template>

    <xsl:template match="text()">
        <xsl:value-of select="."/>
    </xsl:template>

</xsl:stylesheet>"""

    def _add_page_break(self, doc: Document):
        """Add page break to document.

        Args:
            doc: Document object
        """
        doc.add_page_break()

    def _add_footnote(self, paragraph, cite_num: int, bib_entry: Dict[str, Any]):
        """Add footnote with bibliography entry and DOI.

        Args:
            paragraph: Paragraph object to attach footnote to
            cite_num: Citation number
            bib_entry: Bibliography entry with 'formatted' text and optional 'doi'
        """
        # Ensure footnotes part exists
        self._ensure_footnotes_part(paragraph.part)

        # Create footnote reference in text
        run = paragraph.add_run()

        # Add footnote reference element
        footnote_ref = OxmlElement("w:footnoteReference")
        footnote_ref.set(qn("w:id"), str(cite_num))
        run._element.append(footnote_ref)

        # Get footnotes part
        footnotes_part = paragraph.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        )

        # Create footnote element with hyperlink
        formatted_text = bib_entry["formatted"]
        doi = bib_entry.get("doi")

        footnote_elem = self._create_footnote_element(cite_num, formatted_text, doi, footnotes_part)

        # Append to footnotes
        footnotes_part.element.append(footnote_elem)

    def _ensure_footnotes_part(self, doc_part):
        """Ensure the document has a footnotes part with required separators."""
        try:
            # Check if it exists
            doc_part.part_related_by("http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
        except KeyError:
            # Create footnotes part
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.opc.packuri import PackURI
            from docx.opc.part import XmlPart

            # Create proper XML with lxml
            footnotes_xml = etree.Element(
                qn("w:footnotes"),
                nsmap={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                },
            )

            # Add separator footnote (required by Word)
            separator = etree.SubElement(footnotes_xml, qn("w:footnote"))
            separator.set(qn("w:type"), "separator")
            separator.set(qn("w:id"), "-1")
            sep_p = etree.SubElement(separator, qn("w:p"))
            sep_r = etree.SubElement(sep_p, qn("w:r"))
            _sep_t = etree.SubElement(sep_r, qn("w:separator"))  # noqa: F841

            # Add continuation separator footnote (required by Word)
            cont_sep = etree.SubElement(footnotes_xml, qn("w:footnote"))
            cont_sep.set(qn("w:type"), "continuationSeparator")
            cont_sep.set(qn("w:id"), "0")
            cont_p = etree.SubElement(cont_sep, qn("w:p"))
            cont_r = etree.SubElement(cont_p, qn("w:r"))
            _cont_t = etree.SubElement(cont_r, qn("w:continuationSeparator"))  # noqa: F841

            # Create the part
            partname = PackURI("/word/footnotes.xml")
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
            footnotes_part = XmlPart(partname, content_type, footnotes_xml, doc_part.package)

            # Add relationship
            doc_part.relate_to(footnotes_part, RT.FOOTNOTES)

    def _create_footnote_element(self, footnote_id: int, text: str, doi: str | None, footnotes_part) -> etree.Element:
        """Create a footnote XML element.

        Args:
            footnote_id: Footnote ID number
            text: Bibliography text
            doi: Optional DOI URL
            footnotes_part: Footnotes part for creating hyperlink relationships

        Returns:
            lxml Element for footnote
        """
        # Create footnote element
        footnote = etree.Element(qn("w:footnote"))
        footnote.set(qn("w:id"), str(footnote_id))

        # Create paragraph in footnote
        p = etree.SubElement(footnote, qn("w:p"))

        # Add paragraph properties with footnote style
        pPr = etree.SubElement(p, qn("w:pPr"))
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
        pStyle.set(qn("w:val"), "FootnoteText")

        # Add footnote reference mark
        r = etree.SubElement(p, qn("w:r"))
        rPr = etree.SubElement(r, qn("w:rPr"))
        rStyle = etree.SubElement(rPr, qn("w:rStyle"))
        rStyle.set(qn("w:val"), "FootnoteReference")
        etree.SubElement(r, qn("w:footnoteRef"))

        # Add bibliography text (excluding DOI since we'll add it separately as hyperlink)
        # Remove DOI from text if present
        text_without_doi = text.split("\nDOI:")[0] if "\nDOI:" in text else text

        r2 = etree.SubElement(p, qn("w:r"))
        # Add font size 8pt
        rPr2 = etree.SubElement(r2, qn("w:rPr"))
        sz = etree.SubElement(rPr2, qn("w:sz"))
        sz.set(qn("w:val"), "16")  # 16 half-points = 8pt
        szCs = etree.SubElement(rPr2, qn("w:szCs"))
        szCs.set(qn("w:val"), "16")  # For complex scripts

        t = etree.SubElement(r2, qn("w:t"))
        t.set(qn("xml:space"), "preserve")
        t.text = f" {text_without_doi}"

        # Add DOI as clickable hyperlink if present
        if doi:
            doi_url = f"https://doi.org/{doi}" if not doi.startswith("http") else doi

            # Create relationship for hyperlink in footnotes part
            r_id = footnotes_part.relate_to(
                doi_url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )

            # Create hyperlink element
            hyperlink = etree.SubElement(p, qn("w:hyperlink"))
            hyperlink.set(qn("r:id"), r_id)

            # Create run with link text
            r3 = etree.SubElement(hyperlink, qn("w:r"))

            # Add run properties for hyperlink style
            rPr3 = etree.SubElement(r3, qn("w:rPr"))
            rStyle3 = etree.SubElement(rPr3, qn("w:rStyle"))
            rStyle3.set(qn("w:val"), "Hyperlink")

            # Add font size 8pt
            sz3 = etree.SubElement(rPr3, qn("w:sz"))
            sz3.set(qn("w:val"), "16")  # 16 half-points = 8pt
            szCs3 = etree.SubElement(rPr3, qn("w:szCs"))
            szCs3.set(qn("w:val"), "16")

            # Add underline
            u = etree.SubElement(rPr3, qn("w:u"))
            u.set(qn("w:val"), "single")

            # Add color (blue)
            color = etree.SubElement(rPr3, qn("w:color"))
            color.set(qn("w:val"), "0000FF")

            # Add text (just the URL, no "DOI:" prefix)
            t2 = etree.SubElement(r3, qn("w:t"))
            t2.set(qn("xml:space"), "preserve")
            t2.text = f" {doi_url}"

        return footnote

    def _clean_latex_from_text(self, text: str) -> str:
        """Clean LaTeX formatting from text for DOCX display.

        Args:
            text: Text that may contain LaTeX formatting

        Returns:
            Text with LaTeX formatting removed or converted
        """
        import re

        # Replace colored text {\\color{...}text} with just text (handle both \ and \\)
        text = re.sub(r"\{\\color\{[^}]+\}([^}]+)\}", r"\1", text)
        text = re.sub(r"\{\bcolor\{[^}]+\}([^}]+)\}", r"\1", text)

        # Replace common Greek letters in math mode (handle both single and double backslashes)
        text = re.sub(r"\$\\chi\$", "χ", text)
        text = re.sub(r"\$\\\\chi\$", "χ", text)
        text = re.sub(r"\$\\alpha\$", "α", text)
        text = re.sub(r"\$\\\\alpha\$", "α", text)
        text = re.sub(r"\$\\beta\$", "β", text)
        text = re.sub(r"\$\\gamma\$", "γ", text)
        text = re.sub(r"\$\\delta\$", "δ", text)

        return text

    def _add_hyperlink(self, paragraph, url: str, text: str, highlight: bool = False):
        """Add hyperlink to paragraph.

        Args:
            paragraph: Paragraph object
            url: URL to link to
            text: Display text
            highlight: If True, add yellow background highlighting
        """
        # Add hyperlink using OOXML
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create run element
        new_run = OxmlElement("w:r")

        # Create run properties
        r_pr = OxmlElement("w:rPr")

        # Add underline
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

        # Add color (blue)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        r_pr.append(color)

        # Add yellow highlighting if requested
        if highlight:
            highlight_elem = OxmlElement("w:highlight")
            highlight_elem.set(qn("w:val"), "yellow")
            r_pr.append(highlight_elem)

        new_run.append(r_pr)

        # Create text element
        text_element = OxmlElement("w:t")
        text_element.text = text
        new_run.append(text_element)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
