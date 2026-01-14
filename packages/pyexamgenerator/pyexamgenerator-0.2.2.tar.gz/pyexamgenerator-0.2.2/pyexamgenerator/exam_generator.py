# pyexamgenerator: A tool for generating exams from PDF files using AI.
# Copyright (C) 2024 Daniel Sánchez-García

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns

import pandas as pd
import random
from xml.etree.ElementTree import Element, SubElement, tostring
from xml.dom import minidom
import re
import os
from typing import Optional, Tuple  # Importing Optional and Tuple for type hinting


# Define the WordML namespace used for direct XML manipulation in docx files.
WML_NAMESPACE = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

class NoAcceptableQuestionsError(Exception):
    """Excepción personalizada para cuando no se encuentran preguntas aceptables."""
    pass

class ExamGenerator:
    """
    A class to generate exams from a question bank stored in an Excel file.
    It can create multiple versions of an exam, shuffle questions and answers,
    and export to various formats like DOCX and Moodle XML.
    """

    def __init__(self):
        """
        Initializes the pyexamgenerator instance.
        """
        # Added self.df attribute to persist the DataFrame between calls.
        self.df = None

    def read_questions_from_excel(self, excel_path: str) -> Optional[pd.DataFrame]:
        """
        Reads questions from an Excel file and returns a pandas DataFrame.

        Args:
            excel_path (str): The path to the Excel file containing the questions.

        Returns:
            Optional[pd.DataFrame]: A DataFrame with the questions if the read is successful,
                                    None in case of an error.
        """
        try:
            df = pd.read_excel(excel_path)
            # Attempt to convert the 'Número de pregunta' (Question Number) column to an integer.
            if 'Número de pregunta' in df.columns:
                df['Número de pregunta'] = pd.to_numeric(df['Número de pregunta'], errors='coerce').fillna(0).astype(int)
            return df
        except Exception as e:
            print(f"Error al leer el archivo Excel: {e}")
            return None

    def generate_question_text(self, df: pd.DataFrame, renumber: bool = False, shuffle_answers: bool = False) -> Tuple[str, str, pd.DataFrame]:
        """
        Generates the formatted text of the questions (with and without solutions) from a DataFrame.
        This method is primarily used for the 'check' functionality to create a quick preview.

        Args:
            df (pd.DataFrame): The DataFrame containing the questions.
            renumber (bool, optional): If True, renumbers questions sequentially. Defaults to False.
            shuffle_answers (bool, optional): If True, shuffles the order of answers. Defaults to False.

        Returns:
            Tuple[str, str, pd.DataFrame]: A tuple containing the student's exam text, the full exam text (with solutions),
                                           and the modified DataFrame.
        """
        self.exam_text = ""
        self.full_exam_text = ""
        num_total_questions = len(df)
        use_two_digits = num_total_questions > 9

        # Create the "Texto respuesta correcta" (Correct answer text) column for shuffling.
        df['Texto respuesta correcta'] = df.apply(
            lambda row: row[f"Respuesta {row['Respuesta correcta'].upper()}"] if pd.notnull(
                row['Respuesta correcta']) else None, axis=1)
        for index, row in df.iterrows():
            pregunta_num = row['Número de pregunta'] if renumber else index + 1
            formatted_pregunta_num = f"{pregunta_num:02d}" if use_two_digits else str(pregunta_num)

            self.exam_text += f"Pregunta {formatted_pregunta_num}:\n"
            self.full_exam_text += f"Pregunta {formatted_pregunta_num}:\n"
            self.exam_text += f"{row['Pregunta']}\n"
            self.full_exam_text += f"{row['Pregunta']}\n"
            # Shuffle answer texts if shuffle_answers is True.
            if shuffle_answers:
                answers = [
                    row['Respuesta A'],
                    row['Respuesta B'],
                    row['Respuesta C'],
                    row['Respuesta D']
                ]
                random.shuffle(answers)  # Shuffle the list of answers.
                # Identify the new letter of the correct answer.
                new_correct_answer = None
                if row['Texto respuesta correcta'] == answers[0]:
                    new_correct_answer = 'a'
                elif row['Texto respuesta correcta'] == answers[1]:
                    new_correct_answer = 'b'
                elif row['Texto respuesta correcta'] == answers[2]:
                    new_correct_answer = 'c'
                elif row['Texto respuesta correcta'] == answers[3]:
                    new_correct_answer = 'd'
                # Print shuffled answers with letters in order.
                self.exam_text += f"a) {answers[0]}\n"
                self.exam_text += f"b) {answers[1]}\n"
                self.exam_text += f"c) {answers[2]}\n"
                self.exam_text += f"d) {answers[3]}\n"
                self.full_exam_text += f"a) {answers[0]}\n"
                self.full_exam_text += f"b) {answers[1]}\n"
                self.full_exam_text += f"c) {answers[2]}\n"
                self.full_exam_text += f"d) {answers[3]}\n"
                # Update the DataFrame with the new answers and correct answer.
                df.loc[index, 'Respuesta A'] = answers[0]
                df.loc[index, 'Respuesta B'] = answers[1]
                df.loc[index, 'Respuesta C'] = answers[2]
                df.loc[index, 'Respuesta D'] = answers[3]
                df.loc[index, 'Respuesta correcta'] = new_correct_answer
            else:
                # Print answers in original order.
                self.exam_text += f"a) {row['Respuesta A']}\n"
                self.exam_text += f"b) {row['Respuesta B']}\n"
                self.exam_text += f"c) {row['Respuesta C']}\n"
                self.exam_text += f"d) {row['Respuesta D']}\n"
                self.full_exam_text += f"a) {row['Respuesta A']}\n"
                self.full_exam_text += f"b) {row['Respuesta B']}\n"
                self.full_exam_text += f"c) {row['Respuesta C']}\n"
                self.full_exam_text += f"d) {row['Respuesta D']}\n"
                new_correct_answer = row['Respuesta correcta'].lower()  # Initialize here.
            self.full_exam_text += f"Respuesta correcta: {new_correct_answer}\n"
            self.full_exam_text += f"Texto relevante: {row['Texto relevante']}\n"
            self.exam_text += "\n"
            self.full_exam_text += "\n"
        df.drop('Texto respuesta correcta', axis=1, inplace=True)
        return self.exam_text, self.full_exam_text

    def _remove_empty_last_section(self, document: Document):
        """
        Attempts to remove the last section if it appears empty based on XML analysis.
        This is a workaround for a common issue in python-docx where an extra blank page is added.
        """
        body = document._body._element
        sect_prs = body.findall('.//w:sectPr', namespaces=WML_NAMESPACE)
        if sect_prs and sect_prs[-1] == body.getchildren()[-1]:  # Last element in body is sectPr.
            potential_empty = True
            preceding_elements = body.getchildren()[:-1]
            for element in preceding_elements:
                if element.tag.endswith(('}p', '}tbl')):  # Check for paragraphs or tables with namespace.
                    potential_empty = False
                    break
            if potential_empty:
                body.remove(sect_prs[-1])
                print("Se ha intentado eliminar la última sección vacía (XML revisado con namespace).")

    def _create_element(self, tag_name: str) -> OxmlElement:
        """Helper function to create an OxmlElement."""
        return OxmlElement(tag_name)

    def _create_attribute(self, element: OxmlElement, name: str, value: str):
        """Helper function to set an attribute on an OxmlElement."""
        element.set(ns.qn(name), value)

    def add_page_number(self, document: Document) -> Document:
        """
        Adds page numbers to the footer of a Word document using direct docx XML manipulation.
        The format will be "Página X de Y".

        Args:
            document (Document): The docx Document object to which page numbers will be added.

        Returns:
            Document: The Document object with the added page numbers.
        """
        section = document.sections[0]
        footer = section.footer
        if not footer.paragraphs:
            paragraph = footer.add_paragraph()
        else:
            paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.text = "Página "

        r = paragraph.add_run()

        fldChar = self._create_element('w:fldChar')
        self._create_attribute(fldChar, 'w:fldCharType', 'begin')
        r._element.append(fldChar)

        instrText = self._create_element('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = ' PAGE   \\* MERGEFORMAT '
        r._element.append(instrText)

        fldChar = self._create_element('w:fldChar')
        self._create_attribute(fldChar, 'w:fldCharType', 'separate')
        r._element.append(fldChar)

        fldChar = self._create_element('w:fldChar')
        self._create_attribute(fldChar, 'w:fldCharType', 'end')
        r._element.append(fldChar)

        r = paragraph.add_run(" de ")

        fldChar = self._create_element('w:fldChar')
        self._create_attribute(fldChar, 'w:fldCharType', 'begin')
        r._element.append(fldChar)

        instrText = self._create_element('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = ' NUMPAGES   \\* MERGEFORMAT '
        r._element.append(instrText)

        fldChar = self._create_element('w:fldChar')
        self._create_attribute(fldChar, 'w:fldCharType', 'separate')
        r._element.append(fldChar)

        fldChar = self._create_element('w:fldChar')
        self._create_attribute(fldChar, 'w:fldCharType', 'end')
        r._element.append(fldChar)

        for run in paragraph.runs:
            run.font.size = Pt(8)

        return document

    def save_questions_to_docx(self, question_text: str, file_name: Optional[str] = None):
        """
        Saves the formatted question text to a .docx file.

        Args:
            question_text (str): The text of the questions to be saved.
            file_name (Optional[str], optional): The name of the .docx file.
        """
        document = Document()
        document.add_paragraph(question_text)
        document.save(file_name)

    def generate_moodle_xml(self, df: pd.DataFrame, file_name: str, num_total_questions: int,
                            exam: Optional[str] = None, exam_type: Optional[str] = None,
                            xml_cat_additional_text: Optional[str] = None, penalty: int = -25):
        """
        Generates a Moodle XML file from a DataFrame, including category information and correctly
        handling shuffled answers.

        Args:
            df (pd.DataFrame): The DataFrame containing the questions and answers.
            file_name (str): The name of the XML file to generate.
            num_total_questions (int): The total number of questions.
            exam (Optional[str], optional): The name of the exam.
            exam_type (Optional[str], optional): The type of exam.
            xml_cat_additional_text (Optional[str], optional): Additional text for the Moodle XML category.
            penalty (int, optional): The penalty percentage for an incorrect answer (e.g., -25 for 25%).
        """
        self.quiz = Element('quiz')
        self.category = SubElement(self.quiz, 'question', type='category')
        self.category_name = SubElement(self.category, 'category')
        self.text = SubElement(self.category_name, 'text')

        # Build the category text dynamically.
        category_text = "$course$/top/Examen"
        if exam:
            category_text += f" {exam} -"
        if exam_type:
            category_text += f" {exam_type} -"
        if xml_cat_additional_text:
            category_text += f" {xml_cat_additional_text}"

        self.text.text = category_text
        use_two_digits = num_total_questions > 9

        for _, row in df.iterrows():
            correct_answer = row['Respuesta correcta'].lower()

            pregunta_num = row['Número de pregunta']
            formatted_pregunta_num = f"{pregunta_num:02d}" if use_two_digits else str(pregunta_num)

            self.question = SubElement(self.quiz, 'question', type='multichoice')
            self.name = SubElement(self.question, 'name')
            self.text = SubElement(self.name, 'text')
            self.text.text = f"Pregunta {formatted_pregunta_num}"

            if correct_answer == 'a':
                self.answer_a = SubElement(self.question, 'answer', fraction='100')
            else:
                self.answer_a = SubElement(self.question, 'answer', fraction=str(penalty))
            self.text = SubElement(self.answer_a, 'text')
            self.text.text = 'a'

            if correct_answer == 'b':
                self.answer_b = SubElement(self.question, 'answer', fraction='100')
            else:
                self.answer_b = SubElement(self.question, 'answer', fraction=str(penalty))
            self.text = SubElement(self.answer_b, 'text')
            self.text.text = 'b'

            if correct_answer == 'c':
                self.answer_c = SubElement(self.question, 'answer', fraction='100')
            else:
                self.answer_c = SubElement(self.question, 'answer', fraction=str(penalty))
            self.text = SubElement(self.answer_c, 'text')
            self.text.text = 'c'

            if correct_answer == 'd':
                self.answer_d = SubElement(self.question, 'answer', fraction='100')
            else:
                self.answer_d = SubElement(self.question, 'answer', fraction=str(penalty))
            self.text = SubElement(self.answer_d, 'text')
            self.text.text = 'd'

        self.xml_str = minidom.parseString(tostring(self.quiz)).toprettyxml(indent="  ")
        with open(file_name, 'w', encoding='utf-8') as f:
            f.write(self.xml_str)

    def generate_exam_from_excel(
            self,
            bank_excel_path: str,
            output_dir: Optional[str] = None,
            exam_names: Optional[list] = None,
            questions_per_topic: Optional[dict] = None,
            selection_method: str = "azar",
            subject: Optional[str] = None,
            exam: Optional[str] = None,
            course: Optional[str] = None,
            num_exams: int = 2,
            top_margin: float = 1,
            bottom_margin: float = 1,
            left_margin: float = 0.5,
            right_margin: float = 0.5,
            export_moodle_xml: bool = False,
            font_size: int = 9,
            xml_cat_additional_text: Optional[str] = None,
            penalty: int = -25,
            check: bool = False,
            update_excel: bool = False,
            answer_sheet_instructions: Optional[str] = None,
            verbose: bool = False
    ):
        """
        Generates multiple exams in .docx format with varied question and answer orders.
        This is the main orchestrator method of the class.

        Args:
            bank_excel_path (str): Path to the Excel file with the question bank.
            output_dir (Optional[str]): The directory to save the generated exam files.
                If None, uses the same directory as the bank_excel_path.
            exam_names (Optional[list]): A list of names for the exam versions (e.g., ['1A', '1B']).
            questions_per_topic (Optional[dict]): A dictionary specifying how many questions to select from each topic.
            selection_method (str): Method for selecting questions: 'azar' (random), 'primeras' (first N), 'menos usadas' (least used).
            subject (Optional[str]): The subject name for the exam header.
            exam (Optional[str]): The exam name (e.g., 'Parcial 1').
            course (Optional[str]): The course name or year (e.g., '24-25').
            num_exams (int): The number of different exam versions to generate if `exam_names` is not provided.
            top_margin (float): Top page margin in inches.
            bottom_margin (float): Bottom page margin in inches.
            left_margin (float): Left page margin in inches.
            right_margin (float): Right page margin in inches.
            export_moodle_xml (bool): If True, exports the exam to Moodle XML format.
            font_size (int): Font size for the text in the document.
            xml_cat_additional_text (Optional[str]): Additional text for the Moodle XML category.
            penalty (int): Penalty for incorrect answers in the Moodle XML export.
            check (bool): If True, generates a preview and waits for user confirmation before creating all exams.
            update_excel (bool): If True, updates the source Excel file with usage statistics.
            answer_sheet_instructions (Optional[str]): Text with instructions for the answer sheet.
            verbose (bool): If True, prints detailed progress messages to the console.
        """
        if self.df is None:
            self.df = self.read_questions_from_excel(bank_excel_path)

        if self.df is None:
            return

        exam_columns = [col for col in self.df.columns if re.match(r'.*_\d{2}-\d{2}', col)]
        for col in exam_columns:
            self.df[col] = self.df[col].fillna(0)

        self.acceptable_df = self.df[(self.df['Estado'] == 'Aceptable')]

        # Comprobación temprana: si no hay NINGUNA pregunta aceptable, paramos aquí.
        if self.acceptable_df.empty:
            raise NoAcceptableQuestionsError("No se encontraron preguntas con estado 'Aceptable' en el banco de preguntas.")

        if questions_per_topic:
            self.selected_questions = []
            # Iterate directly over the questions_per_topic dictionary.
            for topic, quantity in questions_per_topic.items():
                if verbose:
                    print(f"Verbose: Tema: {topic}, Cantidad solicitada: {quantity}")
                topic_questions = self.acceptable_df[self.acceptable_df['Tema'] == topic]
                # Comprobación por tema: si no hay suficientes preguntas para un tema, lanzamos un error claro.
                if len(topic_questions) < quantity:
                    raise ValueError(
                        f"No hay suficientes preguntas 'Aceptables' para el tema '{topic}'. "
                        f"Se solicitaron {quantity}, pero solo hay {len(topic_questions)} disponibles."
                    )
                if selection_method == "azar":
                    try:
                        topic_questions = topic_questions.sample(n=quantity).copy()
                    except ValueError as e:
                        print(f"Error al seleccionar {quantity} preguntas del tema '{topic}': {e}")
                        continue  # Skip to the next topic if there is an error.
                elif selection_method == "primeras":
                    topic_questions = topic_questions.head(quantity).copy()
                elif selection_method == "menos usadas":
                    topic_questions = topic_questions.sort_values(by='Veces usada en examen').head(quantity).copy()
                self.selected_questions.append(topic_questions)
            if self.selected_questions:
                self.exam_df = pd.concat(self.selected_questions).copy()
            else:
                self.exam_df = pd.DataFrame(columns=self.acceptable_df.columns) # Create an empty DataFrame with the same columns.
        else:
            self.exam_df = self.acceptable_df.copy()

        if verbose:
            print(f"Verbose: Tipo de self.exam_df antes del bucle: {type(self.exam_df)}")
            if isinstance(self.exam_df, pd.DataFrame):
                print(f"Verbose: Head de self.exam_df antes del bucle:\n{self.exam_df.head()}")
            else:
                print(f"Verbose: self.exam_df no es un DataFrame: {self.exam_df}")

        if check:
            preview_exam_df = self.exam_df.copy()
            preview_exam_df['Número de pregunta'] = range(1, len(preview_exam_df) + 1)
            preview_exam_df['Texto respuesta correcta'] = preview_exam_df.apply(
                lambda row: row[f"Respuesta {row['Respuesta correcta'].upper()}"] if pd.notnull(
                    row['Respuesta correcta']) else None, axis=1)
            preview_exam_text, preview_full_exam_text = self.generate_question_text(preview_exam_df.copy(),
                                                                                    renumber=True, shuffle_answers=True)
            preview_document = Document()
            title_paragraph = preview_document.add_paragraph(subject, style='Heading 1')
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle_paragraph = preview_document.add_paragraph(
                f"Examen: {exam} - Curso: {course} - Tipo: Previsualización",
                style='Heading 1'
            )
            subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            preview_paragraph = preview_document.add_paragraph(preview_full_exam_text)
            preview_paragraph.runs[0].font.size = Pt(font_size)
            preview_document = self.add_page_number(preview_document)
            preview_document.save(f'examen_previsualizacion_{subject}_{exam}_{course}_completo.docx')

            print(
                f"Se ha generado un examen de previsualización: examen_previsualizacion_{subject}_{exam}_{course}_completo.docx")
            review = input("¿El examen es correcto? (sí/no): ").lower()

            if review == 'sí':
                try:
                    os.remove(f'examen_previsualizacion_{subject}_{exam}_{course}_completo.docx')
                    print("Archivo de previsualización eliminado.")
                except OSError as e:
                    print("Archivo de previsualización eliminado.")
                except OSError as e:
                    print(f"Error al eliminar el archivo de previsualización: {e}")
            else:
                print("Generación de exámenes cancelada.")
                return

        if exam_names is None:
            num_exams_to_generate = num_exams
            exam_name_list = [f'Tipo {i + 1}' for i in range(num_exams_to_generate)]
        else:
            num_exams_to_generate = len(exam_names)
            exam_name_list = exam_names

        if output_dir is None:
            output_dir = os.path.dirname(bank_excel_path)
        os.makedirs(output_dir, exist_ok=True)

        for i in range(num_exams_to_generate):
            exam_type_name = exam_name_list[i]
            exam_df_shuffled = self.exam_df.sample(frac=1).reset_index(drop=True).copy()

            base_filename = f'examen_{subject}_{exam}_{course}_{exam_type_name}'

            docx_path = os.path.join(output_dir, f'{base_filename}.docx')
            full_docx_path = os.path.join(output_dir, f'{base_filename}_completo.docx')
            excel_path_out = os.path.join(output_dir, f'{base_filename}_completo.xlsx')
            xml_path = os.path.join(output_dir, f'{base_filename}.xml')

            if verbose:
                print(f"Verbose: Iteración {i}, tipo de exam_df_shuffled después de sample: {type(exam_df_shuffled)}")
                if isinstance(exam_df_shuffled, pd.DataFrame):
                    print(f"Verbose: Head de exam_df_shuffled:\n{exam_df_shuffled.head()}")
                else:
                    print(f"Verbose: exam_df_shuffled no es un DataFrame: {exam_df_shuffled}")

            exam_df_shuffled['Número de pregunta'] = range(1, len(exam_df_shuffled) + 1)

            exam_df_shuffled['Texto respuesta correcta'] = exam_df_shuffled.apply(
                lambda row: row[f"Respuesta {row['Respuesta correcta'].upper()}"] if pd.notnull(
                    row['Respuesta correcta']) else None, axis=1)

            def shuffle_row_answers(row):
                answers = [row['Respuesta A'], row['Respuesta B'], row['Respuesta C'], row['Respuesta D']]
                original_correct_answer = row['Texto respuesta correcta']
                random.shuffle(answers)
                new_correct_answer_letter = None
                if original_correct_answer == answers[0]:
                    new_correct_answer_letter = 'a'
                elif original_correct_answer == answers[1]:
                    new_correct_answer_letter = 'b'
                elif original_correct_answer == answers[2]:
                    new_correct_answer_letter = 'c'
                elif original_correct_answer == answers[3]:
                    new_correct_answer_letter = 'd'
                return pd.Series([answers[0], answers[1], answers[2], answers[3], new_correct_answer_letter])

            exam_df_shuffled[['Respuesta A', 'Respuesta B', 'Respuesta C', 'Respuesta D', 'Respuesta correcta']] = \
                exam_df_shuffled.apply(shuffle_row_answers, axis=1)

            exam_text, full_exam_text = self.generate_question_text(exam_df_shuffled.copy(), renumber=False,
                                                                    shuffle_answers=False)

            document = Document()
            title_paragraph = document.add_paragraph(subject, style='Heading 1')
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle_paragraph = document.add_paragraph(
                f"Examen: {exam} - Curso: {course} - Tipo: {exam_type_name}",
                style='Heading 1'
            )
            subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            section = document.sections[0]
            section.top_margin = Inches(top_margin)
            section.bottom_margin = Inches(bottom_margin)
            section.left_margin = Inches(left_margin)
            section.right_margin = Inches(right_margin)

            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = f"Asignatura: {subject} - Examen: {exam} - Curso: {course} - Tipo: {exam_type_name}"
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_paragraph.runs[0].font.size = Pt(font_size)

            for index, row in exam_df_shuffled.iterrows():
                paragraph = document.add_paragraph()
                pregunta_num = row['Número de pregunta']
                use_two_digits = len(exam_df_shuffled) > 9
                formatted_pregunta_num = f"{pregunta_num:02d}" if use_two_digits else str(pregunta_num)

                # Add the question number in bold.
                run_numero = paragraph.add_run(f"Pregunta {formatted_pregunta_num}: ")
                run_numero.font.size = Pt(font_size)
                run_numero.font.bold = True
                # Add the question statement in bold.
                run_enunciado = paragraph.add_run(f"{row['Pregunta']}\n")
                run_enunciado.font.size = Pt(font_size)
                run_enunciado.font.bold = True
                # Add the answer options.
                paragraph.add_run(f"a) {row['Respuesta A']}\n").font.size = Pt(font_size)
                paragraph.add_run(f"b) {row['Respuesta B']}\n").font.size = Pt(font_size)
                paragraph.add_run(f"c) {row['Respuesta C']}\n").font.size = Pt(font_size)
                paragraph.add_run(f"d) {row['Respuesta D']}\n").font.size = Pt(font_size)

            document.add_page_break()
            document.add_heading('Datos del alumno', level=1)
            document.paragraphs[-1].runs[0].font.size = Pt(font_size)
            student_table = document.add_table(rows=4, cols=2)
            student_table.style = 'Table Grid'
            student_data = ['Nombre', 'Apellidos', 'DNI/NIE', 'Firma']
            for j, data in enumerate(student_data):
                cell = student_table.cell(j, 0)
                cell.text = data
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].font.size = Pt(font_size)
            student_table.rows[3].height = Inches(1)

            document.add_heading('Hoja de respuestas', level=1)
            document.paragraphs[-1].runs[0].font.size = Pt(font_size)
            if answer_sheet_instructions:
                instructions_title = document.add_paragraph("Instrucciones de cumplimentación:")
                instructions_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                instructions_title.runs[0].font.size = Pt(font_size)
                instructions_paragraph = document.add_paragraph(answer_sheet_instructions)
                instructions_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                instructions_paragraph.runs[0].font.size = Pt(font_size)

            if verbose:
                print(f"Verbose: Tipo de exam_df_shuffled antes de crear la tabla: {type(exam_df_shuffled)}")
                if isinstance(exam_df_shuffled, pd.DataFrame):
                    print(f"Verbose: Longitud de exam_df_shuffled para número de filas: {len(exam_df_shuffled)}")
                else:
                    print(f"Verbose: exam_df_shuffled no es un DataFrame: {exam_df_shuffled}")

            table = document.add_table(rows=len(exam_df_shuffled) + 1, cols=5)
            table.style = 'Table Grid'
            headers = ['Pregunta', 'a', 'b', 'c', 'd']
            for j, header in enumerate(headers):
                cell = table.cell(0, j)
                cell.text = header
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].font.size = Pt(font_size)
            for j, row in enumerate(exam_df_shuffled.itertuples(), start=1):
                cell = table.cell(j, 0)
                cell.text = str(row._1)
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].font.size = Pt(font_size)
            for j in range(5):
                max_width = 0
                for k in range(len(exam_df_shuffled) + 1):
                    cell_text = table.cell(k, j).text
                    cell_width = len(cell_text)
                    if cell_width > max_width:
                        max_width = cell_width
                    table.columns[j].width = Inches(max_width * 0.08)
                    for row in table.rows:
                        cell = row.cells[j]
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(font_size)

            document = self.add_page_number(document)
            self._remove_empty_last_section(document)
            document.save(docx_path)

            full_document = Document()
            title_paragraph = full_document.add_paragraph(subject, style='Heading 1')
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle_paragraph = full_document.add_paragraph(
                f"Examen: {exam} - Curso: {course} - Tipo: {exam_type_name} - Completo",
                style='Heading 1'
            )
            subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            full_section = full_document.sections[0]
            full_section.top_margin = Inches(top_margin)
            full_section.bottom_margin = Inches(bottom_margin)
            full_section.left_margin = Inches(left_margin)
            full_section.right_margin = Inches(right_margin)

            full_header = full_section.header
            full_header_paragraph = full_header.paragraphs[0]
            full_header_paragraph.text = f"Asignatura: {subject} - Examen: {exam} - Curso: {course} - Tipo: {exam_type_name} - Completo"
            full_header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            full_header_paragraph.runs[0].font.size = Pt(font_size)

            for index, row in exam_df_shuffled.iterrows():
                paragraph = full_document.add_paragraph()
                pregunta_num = row['Número de pregunta']
                use_two_digits = len(exam_df_shuffled) > 9
                formatted_pregunta_num = f"{pregunta_num:02d}" if use_two_digits else str(pregunta_num)

                # Add the question number in bold.
                run_numero = paragraph.add_run(f"Pregunta {formatted_pregunta_num}: ")
                run_numero.font.size = Pt(font_size)
                run_numero.font.bold = True
                # Add the question statement in bold.
                run_enunciado = paragraph.add_run(f"{row['Pregunta']}\n")
                run_enunciado.font.size = Pt(font_size)
                run_enunciado.font.bold = True
                # Add the answer options (not bold).
                paragraph.add_run(f"a) {row['Respuesta A']}\n").font.size = Pt(font_size)
                paragraph.add_run(f"b) {row['Respuesta B']}\n").font.size = Pt(font_size)
                paragraph.add_run(f"c) {row['Respuesta C']}\n").font.size = Pt(font_size)
                paragraph.add_run(f"d) {row['Respuesta D']}\n").font.size = Pt(font_size)
                # Add the correct answer and relevant text (not bold).
                paragraph.add_run(f"Respuesta correcta: {row['Respuesta correcta'].lower()}\n").font.size = Pt(font_size)
                paragraph.add_run(f"Texto relevante: {row['Texto relevante']}\n").font.size = Pt(font_size)
                paragraph.add_run(f"Tema: {row['Tema']}\n").font.size = Pt(font_size)


            full_document = self.add_page_number(full_document)
            self._remove_empty_last_section(full_document)
            full_document.save(full_docx_path)

            exam_df_shuffled.to_excel(excel_path_out, index=False)

            if export_moodle_xml:
                self.generate_moodle_xml(
                    df=exam_df_shuffled.copy(),
                    file_name=xml_path,
                    num_total_questions=len(exam_df_shuffled),
                    exam=exam,
                    exam_type=exam_type_name,
                    xml_cat_additional_text=xml_cat_additional_text,
                    penalty=penalty
                )

        if update_excel:
            new_column_name_base = f'{exam}_{course}'
            safe_exam_name = re.sub(r'[^a-zA-Z0-9_]', '_', exam)
            safe_course_name = re.sub(r'[^a-zA-Z0-9_]', '_', course)
            usage_column_name = f'{safe_exam_name}_{safe_course_name}_uso'

            if usage_column_name not in self.df.columns:
                self.df[usage_column_name] = 0

            used_question_numbers = self.exam_df['Número de pregunta'].tolist()
            for question_number in used_question_numbers:
                original_index = self.df[self.df['Número de pregunta'] == question_number].index
                if not original_index.empty:
                    self.df.loc[original_index, usage_column_name] = 1

            usage_columns = [col for col in self.df.columns if col.endswith('_uso')]

            if not usage_columns:
                print("No se encontraron columnas de uso para sumar.")
                # Si la columna de suma no existe, la creamos.
                if 'Veces usada en examen' not in self.df.columns:
                    self.df['Veces usada en examen'] = 0
            else:
                print(f"Sumando las siguientes columnas de uso: {usage_columns}")
                self.df['Veces usada en examen'] = self.df[usage_columns].sum(axis=1)

            if 'Veces usada en examen' in self.df.columns and 'Texto relevante' in self.df.columns:
                veces_usada_col = self.df.pop('Veces usada en examen')
                try:
                    texto_relevante_idx = self.df.columns.get_loc('Texto relevante')
                    self.df.insert(texto_relevante_idx + 1, 'Veces usada en examen', veces_usada_col)
                    print("Columna 'Veces usada en examen' movida a su posición correcta.")
                except KeyError:
                    # Si 'Texto relevante' no existe, la añade al final.
                    self.df['Veces usada en examen'] = veces_usada_col
                    print("Advertencia: No se encontró la columna 'Texto relevante'. 'Veces usada en examen' se añadió al final.")

            self.df = self.df.loc[:, ~self.df.columns.str.contains('unnamed', case=False)]

            try:
                self.df.to_excel(bank_excel_path, index=False)
                print(
                    f"Archivo Excel '{bank_excel_path}' actualizado con la columna '{usage_column_name}' y 'Veces usada en examen'.")
            except Exception as e:
                print(f"Error al actualizar el archivo Excel '{bank_excel_path}': {e}")

## Example of use

# generator = pyexamgenerator()
#
# df = pd.read_excel('banco_prueba_pendiente_de_revisar.xlsx')
# temas = [i for i in df['Tema']]
# temas = list(dict.fromkeys(temas))
#
#
# question_dict = {}
# for t in temas:
#     question_dict.update({t: 2})
#
# instrucciones = '''
# -	Marque con una “x” la respuesta correcta en la tabla que se muestra a continuación.
# -	La tabla debe estar LIBRE DE ANOTACIONES. Se pueden hacer anotaciones en la hoja de preguntas, o anotaciones con lápiz en la tabla de respuestas, siempre y cuando el examen se entregue cumplimentado a bolígrafo.
# '''
#
# generator.generate_exam_from_excel(
#     bank_excel_path='banco_prueba_pendiente_de_revisar.xlsx',
#     questions_per_topic=question_dict,
#     subject='Calidad, Seguridad y Protección ambiental',
#     exam='Parcial 1',
#     course='24-25',
#     # num_exams=2,
#     exam_names=['1A', '1B'],
#     export_moodle_xml=True,
#     xml_cat_additional_text='',
#     penalty=-25,
#     # check=True,
#     update_excel=True,
#     answer_sheet_instructions=instrucciones
# )