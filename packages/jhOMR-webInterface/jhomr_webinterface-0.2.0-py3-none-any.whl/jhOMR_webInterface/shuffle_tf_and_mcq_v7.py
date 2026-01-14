import re
import random
from docx import Document
from docx.shared import Pt
from copy import deepcopy
from openpyxl import load_workbook
import os
import shutil



def copy_paragraph_with_formatting(source_para, target_doc):
    target_para = target_doc.add_paragraph()
    for run in source_para.runs:
        new_run = target_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size if run.font.size else Pt(12)
    return target_para
    
# ---------- T/F Utilities ----------
def extract_tf_questions_between_headers(doc, tf_header_pattern, mcq_header_pattern):
    tf_questions = []
    in_tf_section = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if not in_tf_section:
            if re.search(tf_header_pattern, text, re.IGNORECASE):
                in_tf_section = True
            continue
        if re.search(mcq_header_pattern, text, re.IGNORECASE):
            break
        tf_questions.append(text)
    return tf_questions

def count_matches(a, b):
    return sum(i == j for i, j in zip(a, b))

def generate_balanced_overlap_sets(original, num_sets=2, num_trials=2000):
    n = len(original)
    best_sets = []
    best_indices = []
    min_max_overlap = float('inf')
    for _ in range(num_trials):
        candidates = []
        indices_list = []
        for _ in range(num_sets):
            permuted_indices = list(range(n))
            while True:
                random.shuffle(permuted_indices)
                shuffled = [original[i] for i in permuted_indices]
                if shuffled != original:
                    break
            candidates.append(shuffled)
            indices_list.append(permuted_indices)
        all_sets = [original] + candidates
        pairwise_overlaps = []
        for i in range(len(all_sets)):
            for j in range(i + 1, len(all_sets)):
                overlap = count_matches(all_sets[i], all_sets[j])
                pairwise_overlaps.append(overlap)
        max_overlap = max(pairwise_overlaps)
        if max_overlap < min_max_overlap:
            min_max_overlap = max_overlap
            best_sets = candidates
            best_indices = indices_list
    return best_sets, best_indices, min_max_overlap

def shuffle_by_index(questions, index_map):
    return [questions[i] for i in index_map]

def write_tf_to_doc(doc, tf_questions,original_doc,tf_header_re):
    header_found = False
    for para in original_doc.paragraphs:
        if re.search(tf_header_re, para.text, re.IGNORECASE):
            copy_paragraph_with_formatting(para, doc)
            doc.add_paragraph()  # Add space after header
            header_found = True
            break
    if not header_found:
        print("Find True/False")

    #doc.add_paragraph()  # spacing
    for i, q in enumerate(tf_questions, 1):
        doc.add_paragraph(f"{i}. {q}")
    doc.add_paragraph()


# ---------- MCQ Utilities ----------
def is_option_paragraph(text):
    return re.match(r'^[a-fA-F][\).]', text.strip())

def extract_mcq_blocks(doc):
    mcq_blocks = []
    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if not text:
            i += 1
            continue
        if not is_option_paragraph(text):
            if i + 1 < len(doc.paragraphs) and is_option_paragraph(doc.paragraphs[i + 1].text.strip()):
                block = {
                    'question_idx': i,
                    'question_text': text,
                    'option_indices': [],
                    'option_paragraphs': []
                }
                j = i + 1
                while j < len(doc.paragraphs):
                    opt_text = doc.paragraphs[j].text.strip()
                    if is_option_paragraph(opt_text):
                        block['option_indices'].append(j)
                        block['option_paragraphs'].append(doc.paragraphs[j])
                        j += 1
                    else:
                        break
                mcq_blocks.append(block)
                i = j
            else:
                i += 1
        else:
            i += 1
    return mcq_blocks



def shift_list(lst, shift_by, direction='right'):
    n = len(lst)
    shift_by %= n
    if direction == 'right':
        return lst[-shift_by:] + lst[:-shift_by]
    elif direction == 'left':
        return lst[shift_by:] + lst[:shift_by]
    else:
        raise ValueError("Direction must be 'left' or 'right'")

def contains_special_phrase(option_text):
    # Define normalized target phrases
    phrases = [
        'none of them',
        'none of these',
        'not one of them',
        'all of them',
        'all of the above',
        'all of above'
        'all of these'
    ]

    # Normalize the input text: lowercase, strip, collapse spaces
    cleaned = re.sub(r'\s+', ' ', option_text.strip().lower())

    # Check if any of the phrases exist in the cleaned text
    return any(phrase in cleaned for phrase in phrases)
    
def shuffle_mcq_options(block, original_correct_label, shift_by=1, direction='right'):
    # Concatenate all option texts
    full_text = ' '.join(p.text.strip() for p in block['option_paragraphs'] if p.text.strip())
    parts = re.split(r'(?=([a-fA-F][\).]))', full_text)
    
    if not parts or len(parts) < 3:
        return [full_text], original_correct_label

    options = []
    labels = []
    i = 1
    while i < len(parts):
        label = parts[i]
        content = parts[i + 1].strip() if i + 1 < len(parts) else ''
        content_cleaned = re.sub(r'^[a-fA-F][\).]\s*', '', content)
        labels.append(label[0].lower())
        options.append(content_cleaned)
        i += 2

    try:
        original_index = labels.index(original_correct_label.lower())
    except ValueError:
        # fallback if label is not found
        original_index = 0

    preserve_last = False
    last_option = None
    if len(options) >= 1 and contains_special_phrase(options[-1]):
        preserve_last = True
        last_option = options.pop()
        if original_index == len(options):
            new_correct_label = chr(ord('a') + len(options))
            options.append(last_option)
            result = [f"{chr(ord('a') + idx)}) {opt}" for idx, opt in enumerate(options)]
            return result, new_correct_label

    if len(options) > 1:
        shifted_options = shift_list(options, shift_by, direction)
        total = len(options)
        shift_by_mod = shift_by % total
        if direction == 'right':
            new_index = (original_index + shift_by_mod) % total
        elif direction == 'left':
            new_index = (original_index - shift_by_mod) % total
    else:
        shifted_options = options
        new_index = original_index

    if preserve_last:
        shifted_options.append(last_option)

    result = [f"{chr(ord('a') + idx)}) {opt}" for idx, opt in enumerate(shifted_options)]
    new_correct_label = chr(ord('a') + new_index)
    return result, new_correct_label


def write_mcq_to_doc(doc, mcqs, original_correct_answers, shift_by, direction,original_doc,mcq_header_re, write_options_in_one_line=True):
    header_found = False
    for para in original_doc.paragraphs:
        if re.search(mcq_header_re, para.text, re.IGNORECASE):
            copy_paragraph_with_formatting(para, doc)
            doc.add_paragraph()  # Add space after header
            header_found = True
            break
    if not header_found:
        print("Answer the following questions:")

    new_correct_answers = []
    for i, block in enumerate(mcqs, 1):
        doc.add_paragraph(f"{i}. {block['question_text']}")

        shuffled_options, new_correct_label = shuffle_mcq_options(block, original_correct_answers[i - 1], shift_by, direction)
        new_correct_answers.append(new_correct_label)
        if write_options_in_one_line:
            option_line = '  '.join(shuffled_options)
            doc.add_paragraph(f"   {option_line}")
        else:
            for opt in shuffled_options:
                doc.add_paragraph(f"   {opt}")
        doc.add_paragraph()
    return new_correct_answers

def report_overlap(a, b, label):
    print(f"Overlap {label}: {count_matches(a, b)}")
    



def copy_word_file_with_overwrite(src_path, dest_path):
    if os.path.exists(dest_path):
        os.remove(dest_path)
        print(f"Removed existing file: {dest_path}")

    shutil.copyfile(src_path, dest_path)
    print(f"Copied {src_path} -> {dest_path}")

def write_to_excel(list_of_values,excel_path,course_name,col='A'):
    col_dict={'C':3,'D':4,'E':5,'F':6,'G':7,'J':10,'K':11,'L':12,'M':13,'N':14}
    col_num=col_dict[col]
    # Load the existing Excel workbook
    
    wb = load_workbook(excel_path)
    # Select the active sheet
    ws = wb.active  # Or specify: ws = wb["Sheet1"]


    # Start from row 2
    start_row = 2
    # Write values to a specific column
    for i, answer in enumerate(list_of_values):
        ws.cell(row=start_row + i, column=col_num, value=answer)

    ws.cell(row=start_row, column=1, value=course_name)
    
    # Save workbook
    wb.save(excel_path)
    print("\nValues written to Excel column:", col)



def create_set_2_and_set3_and_correct_answer_excel(base_dir,MS_word_folder_name,course_name,set_1_tf_answers,set_1_mcq_answers,correct_tf_mark,wrong_tf_mark,correct_mcq_mark,wrong_mcq_mark,return_debug=False):
    print("\n\n\n\n")
    print(base_dir)
    print("\n\n\n\n")
    errors = []
    tf_section_errors = []
    qa_section_errors = []
    results = {}
    
#    base_dir = os.path.dirname(__file__)  # Folder where the script is located
    static_dir = os.path.join(base_dir, MS_word_folder_name)

    #Creating the required .docx and .xlsx files
    input_file = os.path.join(static_dir, "Set_1_Uploaded.docx")
    
    output_file_2 = os.path.join(static_dir, "Set_2_"+course_name+".docx")
    output_file_3 = os.path.join(static_dir, "Set_3_"+course_name+".docx")
    template_file_2 = os.path.join(static_dir, "Set_2_blank_template.docx")
    template_file_3 = os.path.join(static_dir, "Set_3_blank_template.docx")

    copy_word_file_with_overwrite(template_file_2, output_file_2)
    copy_word_file_with_overwrite(template_file_3, output_file_3)
    

    src_file_excel = os.path.join(static_dir, "Answer_sheet_template.xlsx")
    excel_path = os.path.join(static_dir, "Correct_Answers_"+course_name+".xlsx")
    shutil.copyfile(src_file_excel, excel_path)

    print("Input file:\n",input_file)
    print("\n\n\n")
    original_doc = Document(input_file)
    doc2 = Document(output_file_2)
    doc3 = Document(output_file_3)
    ###################################################
    
    # ---- True/False ----
    try:
        tf_header_re = r'.*true\s*or\s*false'
        mcq_header_re = r'.*answer.*questions'


        original_tf_questions = extract_tf_questions_between_headers(original_doc, tf_header_re, mcq_header_re)
        
        print("\n\nTotal ",len(original_tf_questions), "T/F questions detected. They are\n\n",original_tf_questions)
        results['no of detected tf'] = {"Set1": len(original_tf_questions)}
        
        required_tf_answers = set_1_tf_answers[:len(original_tf_questions)]
        required_correct_tf_marks=correct_tf_mark[:len(original_tf_questions)]
        required_wrong_tf_marks=wrong_tf_mark[:len(original_tf_questions)]
        
        print("\n\nInside the util function\n")
        print(correct_tf_mark,'\n',wrong_tf_mark,'\n',correct_mcq_mark,'\n',wrong_mcq_mark)
        print(type(correct_tf_mark),'\n',type(wrong_tf_mark),'\n',type(correct_mcq_mark),'\n',type(wrong_mcq_mark))

        if len(required_tf_answers) < len(original_tf_questions) or "" in required_tf_answers:
            tf_section_errors.append(f"{len(original_tf_questions)} TF questions detected, but {len(required_tf_answers)} answers provided (some may be missing or blank).<br>")
        if len(required_correct_tf_marks) < len(original_tf_questions) or "" in correct_tf_mark:
            tf_section_errors.append(f"{len(original_tf_questions)} TF questions detected, but marks for 'Mark if Correct' for all of them are not provided. (some may be missing or blank)<br>.")
        if len(required_wrong_tf_marks) < len(original_tf_questions) or "" in wrong_tf_mark:
            tf_section_errors.append(f"{len(original_tf_questions)} TF questions detected, but marks for 'Mark if Wrong' for all of them are not provided. (some may be missing or blank).<br>")
            
        if tf_section_errors:
            numbered_error_messages = []
            for i, msg in enumerate(tf_section_errors, 1): # Start enumeration from 1
                numbered_error_messages.append(f"{i}# {msg}")
            # Join the numbered messages with <br> for HTML display
            combined_error_message = "<br>".join(numbered_error_messages)
            raise ValueError(combined_error_message)        

        write_to_excel(set_1_tf_answers,excel_path,course_name,col='C')        
        write_to_excel(correct_tf_mark,excel_path,course_name,col='F')
        write_to_excel(wrong_tf_mark,excel_path,course_name,col='G')

        # Generate sets
        (sets, indices, max_overlap) = generate_balanced_overlap_sets(set_1_tf_answers)
        set2_tf_answers, set3_tf_answers = sets
        index_map2, index_map3 = indices

        set2_tf_questions = shuffle_by_index(original_tf_questions, index_map2)
        set3_tf_questions = shuffle_by_index(original_tf_questions, index_map3)

        print("\n\nNew T/F answers Set 2:", set2_tf_answers)
        print("\n\nNew T/F answers Set 3:", set3_tf_answers)
        results['tf_answers'] = {
            "Set1": set_1_tf_answers,
            "Set2": set2_tf_answers,
            "Set3": set3_tf_answers,
                                }
        results['tf_overlap'] = {
            "Set1-Set2": count_matches(set_1_tf_answers, set2_tf_answers),
            "Set1-Set3": count_matches(set_1_tf_answers, set3_tf_answers),
            "Set2-Set3": count_matches(set2_tf_answers, set3_tf_answers),
                                }

        report_overlap(set_1_tf_answers, set2_tf_answers, "Set1-Set2")
        report_overlap(set_1_tf_answers, set3_tf_answers, "Set1-Set3")
        report_overlap(set2_tf_answers, set3_tf_answers, "Set2-Set3")

        
        write_tf_to_doc(doc2,  set2_tf_questions,original_doc,tf_header_re)
        doc2.save(output_file_2)
        print("Save T/F in doc2 success")

        write_to_excel(set2_tf_answers,excel_path,course_name,col='D')
                
        
        write_tf_to_doc(doc3,  set3_tf_questions,original_doc,tf_header_re)
        doc3.save(output_file_3)
        print("Save T/F in doc3 success")
        write_to_excel(set3_tf_answers,excel_path,course_name,col='E')
        

        print(f"Saved changes to {excel_path}")


    except Exception as e:
        errors.append(f"True/False section error:<br> {str(e)}")
        print(f"\n\nERROR: {e}\n\n")
    


    # ---- MCQs ----
    try:       
        #Shifting the MCQ answers randomly to the left or right
        shift_by_2 = random.randint(1, 3)
        direction_2 = random.choice(['left', 'right'])

        shift_by_3 = shift_by_2+1
        direction_3 = direction_2

        mcq_blocks = extract_mcq_blocks(original_doc)
        
        #print(mcq_blocks)
        results['no of detected mcq'] = {"Set1": len(mcq_blocks)}

        required_mcq_answers = set_1_mcq_answers[:len(mcq_blocks)]
        required_correct_mcq_marks=correct_mcq_mark[:len(mcq_blocks)]
        required_wrong_mcq_marks=wrong_mcq_mark[:len(mcq_blocks)]
         

        if len(required_mcq_answers) < len(mcq_blocks) or "" in required_mcq_answers:
            qa_section_errors.append(f"{len(mcq_blocks)} MCQ questions detected, but {len(required_mcq_answers)} answers provided (some may be missing or blank).<br>")
        if len(required_correct_mcq_marks) < len(mcq_blocks) or "" in required_correct_mcq_marks:
            qa_section_errors.append(f"{len(mcq_blocks)} MCQ questions detected, but marks for 'Mark if Correct' for all of them are not provided. (some may be missing or blank)<br>.")
        if len(required_wrong_mcq_marks) < len(mcq_blocks) or "" in required_wrong_mcq_marks:
            qa_section_errors.append(f"{len(mcq_blocks)} MCQ questions detected, but marks for 'Mark if Wrong' for all of them are not provided. (some may be missing or blank)<br>.")
        
        if qa_section_errors:
            numbered_error_messages = []
            for i, msg in enumerate(qa_section_errors, 1): # Start enumeration from 1
                numbered_error_messages.append(f"{i}# {msg}")
            # Join the numbered messages with <br> for HTML display
            combined_error_message = "<br>".join(numbered_error_messages)
            raise ValueError(combined_error_message)        



        
        #print(len(macq_blocks))
        new_mcq_answers_2 = write_mcq_to_doc(doc2, mcq_blocks, set_1_mcq_answers, shift_by_2, direction_2,original_doc,mcq_header_re)
        new_mcq_answers_3 = write_mcq_to_doc(doc3, mcq_blocks, set_1_mcq_answers, shift_by_3, direction_3,original_doc,mcq_header_re)
        
        print("\n\nNew MCQ answers Set 2:", new_mcq_answers_2)
        print("New MCQ answers Set 3:", new_mcq_answers_3)

        results['mcq answers'] = {"Set1": set_1_mcq_answers,
                                  "Set2": new_mcq_answers_2,
                                  "Set3": new_mcq_answers_3
                                 }
        
        write_to_excel(set_1_mcq_answers,excel_path,course_name,col='J')
        write_to_excel(correct_mcq_mark,excel_path,course_name,col='M')
        write_to_excel(wrong_mcq_mark,excel_path,course_name,col='N')
        
        doc2.save(output_file_2)
        print("Save Q/A in doc2 success")
        write_to_excel(new_mcq_answers_2,excel_path,course_name,col='K')

        doc3.save(output_file_3)
        print("Save Q/A in doc3 success")
        write_to_excel(new_mcq_answers_3,excel_path,course_name,col='L')  

        

    except Exception as e:
        errors.append(f"MCQ section error:<br> {str(e)}")
        print(f"\n\nERROR: {e}\n\n")


    return results, errors

    #############################################################################################################
    #############################################################################################################
    #############################################################################################################




