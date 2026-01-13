"""
DOCX Generator
Converts Markdown files to properly formatted Word documents.

Exit codes:
    0: Success
    1: Usage error (no file provided)
    2: Dependency missing (python-docx not installed)
    3: File not found
    4: File read error
    5: DOCX generation error
    6: Unexpected error
"""
import re
import sys
import traceback
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor
except ImportError as e:
    print(f"[DEPENDENCY_MISSING] python-docx package not installed", file=sys.stderr)
    print(f"  Original error: {e}", file=sys.stderr)
    print(f"  To fix: pip install python-docx", file=sys.stderr)
    sys.exit(2)

from config import OUTPUTS_DIR


def print_error(error_code: str, message: str, details: dict = None):
    """Print formatted error to stderr."""
    print(f"\n[{error_code}]", file=sys.stderr)
    print(f"Error: {message}", file=sys.stderr)
    if details:
        for key, value in details.items():
            if value is not None:
                print(f"  {key}: {value}", file=sys.stderr)


def setup_document_styles(doc: Document):
    """Configure document styles for consistent formatting."""
    # Normal text style
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(24)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(18)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def is_table_row(line: str) -> bool:
    """Check if line is a markdown table row."""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|')


def is_separator_row(line: str) -> bool:
    """Check if line is table separator (|---|---|)."""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    # Match only separator pattern: | :?-+:? | with multiple cells
    return bool(re.match(r'^\|[\s:]*-+[\s:]*(\|[\s:]*-+[\s:]*)+\|$', stripped))


def parse_table_cells(line: str) -> list[str]:
    """Extract cell contents from table row, handling escaped pipes."""
    stripped = line.strip()
    # Remove leading/trailing pipes
    inner = stripped[1:-1]
    # Split by unescaped pipes (not preceded by backslash)
    cells = re.split(r'(?<!\\)\|', inner)
    # Unescape pipes and strip whitespace
    return [cell.replace('\\|', '|').strip() for cell in cells]


def add_table_to_doc(doc: Document, rows: list[str]):
    """Add markdown table as Word table with inline formatting support."""
    if not rows:
        return

    # First row is header
    header_cells = parse_table_cells(rows[0])
    num_cols = len(header_cells)

    # Filter out separator row and get data rows
    data_rows = [r for r in rows[1:] if not is_separator_row(r)]
    num_rows = 1 + len(data_rows)

    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Fill header row (bold)
    header_row = table.rows[0]
    for i, cell_text in enumerate(header_cells):
        cell = header_row.cells[i]
        # Clear default paragraph and apply inline formatting
        cell.paragraphs[0].clear()
        parse_inline_formatting(cell.paragraphs[0], cell_text)
        # Make all runs bold for header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Fill data rows with inline formatting
    for row_idx, row_line in enumerate(data_rows):
        cells = parse_table_cells(row_line)
        table_row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(cells):
            if col_idx < num_cols:
                cell = table_row.cells[col_idx]
                cell.paragraphs[0].clear()
                parse_inline_formatting(cell.paragraphs[0], cell_text)

    # Add spacing after table
    doc.add_paragraph()


def parse_inline_formatting(paragraph, text: str):
    """Parse and apply inline markdown formatting (bold, italic, links)."""
    # Pattern to match bold, italic, and combined
    patterns = [
        (r'\*\*\*([^*]+)\*\*\*', 'bold_italic'),  # ***text***
        (r'\*\*([^*]+)\*\*', 'bold'),              # **text**
        (r'\*([^*]+)\*', 'italic'),                # *text*
        (r'__([^_]+)__', 'bold'),                  # __text__
        (r'_([^_]+)_', 'italic'),                  # _text_
        (r'\[([^\]]+)\]\([^)]+\)', 'link'),        # [text](url)
    ]

    # Simple approach: handle bold first, then regular text
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Handle links - extract just the text
            link_pattern = r'\[([^\]]+)\]\([^)]+\)'
            link_free = re.sub(link_pattern, r'\1', part)
            paragraph.add_run(link_free)


def md_to_docx(md_path: str, output_path: str = None) -> str:
    """
    Convert markdown file to DOCX format.

    Args:
        md_path: Path to markdown file
        output_path: Optional custom output path

    Returns:
        Path to generated DOCX file
    """
    md_path = Path(md_path)

    if not md_path.exists():
        # Try in outputs directory
        md_path = OUTPUTS_DIR / md_path.name
        if not md_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {md_path}")

    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    # Create document
    doc = Document()
    setup_document_styles(doc)

    i = 0
    in_code_block = False
    code_content = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block - add as formatted paragraph
                if code_content:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip metadata blocks (> lines at start)
        if stripped.startswith('>'):
            i += 1
            continue

        # Skip horizontal rules
        if stripped in ['---', '***', '___']:
            i += 1
            continue

        # H1 - Title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading = doc.add_heading(stripped[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # H2 - Main sections
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        # H3 - Subsections
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        # H4 - Sub-subsections
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Table detection
        if is_table_row(stripped):
            table_rows = []
            # Collect all consecutive table rows (optimized: reuse stripped)
            while i < len(lines):
                line_stripped = lines[i].strip()
                if not is_table_row(line_stripped):
                    break
                table_rows.append(lines[i])
                i += 1
            add_table_to_doc(doc, table_rows)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)
        i += 1

    # Determine output path
    if output_path:
        docx_path = Path(output_path)
    else:
        docx_path = md_path.with_suffix('.docx')

    # Ensure output directory exists
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    # Save document
    doc.save(str(docx_path))

    return str(docx_path)


def count_words_in_docx(docx_path: str) -> int:
    """Count words in a DOCX file."""
    doc = Document(docx_path)
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    return word_count


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print_error(
            "USAGE_NO_FILE",
            "No markdown file provided",
            {
                "usage": "python docx-generator.py <markdown_file> [output_file]",
                "example": "python docx-generator.py outputs/article-seo-tips.md"
            }
        )
        sys.exit(1)

    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        result = md_to_docx(md_file, output_file)
        word_count = count_words_in_docx(result)
        print(f"DOCX generated successfully: {result}")
        print(f"Word count: {word_count}")

    except FileNotFoundError as e:
        print_error(
            "FILE_NOT_FOUND",
            str(e),
            {
                "input_file": md_file,
                "outputs_dir": str(OUTPUTS_DIR),
                "suggestion": "Check the file path and ensure the file exists"
            }
        )
        sys.exit(3)

    except PermissionError as e:
        # Determine if it's a read or write error based on context
        if "Cannot read" in str(e) or md_file in str(e):
            print_error(
                "FILE_READ_PERMISSION",
                f"Cannot read markdown file: {md_file}",
                {
                    "original_error": str(e),
                    "suggestion": "Check file read permissions"
                }
            )
            sys.exit(4)
        else:
            print_error(
                "FILE_WRITE_PERMISSION",
                f"Cannot write DOCX file",
                {
                    "original_error": str(e),
                    "suggestion": "Check directory write permissions"
                }
            )
            sys.exit(5)

    except UnicodeDecodeError as e:
        print_error(
            "FILE_ENCODING_ERROR",
            f"Cannot decode markdown file as UTF-8: {md_file}",
            {
                "original_error": str(e),
                "suggestion": "Ensure the file is saved as UTF-8 encoding"
            }
        )
        sys.exit(4)

    except OSError as e:
        if e.errno == 28:  # ENOSPC - No space left on device
            print_error(
                "FILE_DISK_FULL",
                "Cannot write DOCX file - disk full",
                {
                    "original_error": str(e),
                    "suggestion": "Free up disk space and try again"
                }
            )
        else:
            print_error(
                "FILE_IO_ERROR",
                f"File I/O error: {type(e).__name__}",
                {
                    "errno": e.errno,
                    "original_error": str(e)
                }
            )
        sys.exit(5)

    except KeyboardInterrupt:
        print("\n\nOperation cancelled by user.", file=sys.stderr)
        sys.exit(130)

    except Exception as e:
        # Unexpected error - provide full traceback for debugging
        print_error(
            f"UNEXPECTED_{type(e).__name__.upper()}",
            str(e),
            {
                "error_type": type(e).__name__,
                "module": type(e).__module__,
                "traceback": traceback.format_exc()
            }
        )
        print("\nThis is an unexpected error. Please report this issue.", file=sys.stderr)
        sys.exit(6)
