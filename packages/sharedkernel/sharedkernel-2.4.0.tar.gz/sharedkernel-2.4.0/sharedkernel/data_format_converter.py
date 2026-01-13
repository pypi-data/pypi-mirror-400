from abc import ABC, abstractmethod
import markdown

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString
import mammoth


class DataFormatConverter(ABC):
    @abstractmethod
    def convert(input_path, output_path):
        pass

    def _get_out_title(self, title: str, file_format: str = ".docx") -> str:
        return ".".join(title.split(".")[:-1]) + file_format


class Md2Docx(DataFormatConverter):

    def __add_hyperlink(self, paragraph, url, text, color="0000FF", underline=True):
        r_id = paragraph.part.relate_to(url, "hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run, rPr = OxmlElement("w:r"), OxmlElement("w:rPr")
        if color:
            color_element = OxmlElement("w:color")
            color_element.set(qn("w:val"), color)
            rPr.append(color_element)
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
        new_run.append(rPr)

        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def __set_right_to_left(self, paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.right_to_left = True

    @staticmethod
    def convert(input_path, output_path=None):
        md2docx = Md2Docx()

        with open(input_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        output_file = output_path or md2docx._get_out_title(input_path)

        html_content = markdown.markdown(md_content)
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        for element in soup:
            if element.name == "h1":
                paragraph = doc.add_paragraph(element.get_text(), style="Heading 1")
            elif element.name == "h2":
                paragraph = doc.add_paragraph(element.get_text(), style="Heading 2")
            elif element.name == "p":
                paragraph = doc.add_paragraph()
                for part in element.contents:
                    if isinstance(part, NavigableString):
                        paragraph.add_run(str(part))
                    elif part.name == "a":
                        md2docx.__add_hyperlink(
                            paragraph, part["href"], part.get_text()
                        )
                    elif part.name == "img":
                        md2docx.__add_hyperlink(
                            paragraph, part["src"], part.get("alt", "Image")
                        )
            md2docx.__set_right_to_left(paragraph)

        doc.save(output_file)
        return output_file


class Docx2Md(DataFormatConverter):
    def _get_out_title(self, title: str, file_format: str = ".txt") -> str:
        return super()._get_out_title(title, file_format)

    @staticmethod
    def convert(input_path: str, output_path: str = None):
        docx2md = Docx2Md()
        with open(input_path, "rb") as docx:
            result = mammoth.convert_to_markdown(
                docx,
                convert_image=mammoth.images.img_element(
                    lambda image: {"src": "images/" + image.content_type.split("/")[-1]}
                ),
            )

        output_path = output_path or docx2md._get_out_title(input_path)

        with open(output_path, "w", encoding="utf-8") as md_file:
            md_file.write(result.value)

        if result.messages:
            for message in result.messages:
                print("Warning:", message)
