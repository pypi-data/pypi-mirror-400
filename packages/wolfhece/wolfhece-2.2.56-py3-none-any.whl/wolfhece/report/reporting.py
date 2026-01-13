
"""
Author: HECE - University of Liege, Pierre Archambeau
Date: 2024

Copyright (c) 2024 University of Liege. All rights reserved.

This script and its content are protected by copyright law. Unauthorized
copying or distribution of this file, via any medium, is strictly prohibited.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import RGBColor
from pathlib import Path
from typing import Union, List, Tuple,Literal
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import pandas as pd
from tempfile import NamedTemporaryFile, TemporaryDirectory
import logging
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from datetime import datetime
import os
import socket
import hashlib
import re

from gettext import gettext as _

class RapidReport:
    """ 
    Class for creating a report 'quickly'.

    It can be used in Jupyter notebooks or in scripts to create a simple report in Word format.

    Word document is created with the following structure:

    - Main page with title, author, date and hash of the document
    - Summary (automatically generated)
    - Title
    - Paragraph
    - Figure (numbered automatically with caption)
    - Bullet list
    - Table

    It is not a full-fledged reporting tool with advanced functionnalities but a simple way to create a report quickly 'on-the-fly'.

    
    Example:

    ```
    rapport = RapidReport('Rapport de Projet', 'Alice')

    rapport.add_title('Titre Principal', level=0)
    rapport.add_paragraph('Ceci est un **paragraphe** introductif avec des mots en *italique* et en **gras**.')

    rapport += "Tentative d'ajout de figure vie un lien incorrect.\nPassage à la ligne"
    rapport.add_figure('/path/to/image.png', 'Légende de la figure.')

    rapport.add_bullet_list(['Premier élément', 'Deuxième élément', 'Troisième élément'])

    rapport.add_table_from_listoflists([['Nom', 'Âge'], ['Alice', '25'], ['Bob', '30']])

    rapport.save('rapport.docx')
    ```

    """

    def __init__(self, main_title:str, author:str):
        
        self._main_title = main_title
        self._author = author
        self._date = None

        self._content = []
        self._document = None

        self._filename = None

        self._idx_figure = 0

        self._styles={}

        self._has_first_page = False
 
    def _define_default_styles(self):
        
        # Définir le style de titre
        self._title_style = self._document.styles.add_style('TitleStyle', 1)
        self._title_style.font.name = 'Arial'
        self._title_style.font.size = Pt(20)
        self._title_style.font.bold = True
        self._title_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
 
        # Définir le style de légende
        self._caption_style = self._document.styles.add_style('CaptionStyle', 1)
        self._caption_style.font.name = 'Arial'
        self._caption_style.font.size = Pt(9)
        self._caption_style.font.italic = True
        self._caption_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        self._caption_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
 
        # Définir le style de corps de texte
        self._body_text_style = self._document.styles.add_style('BodyTextStyle', 1)
        self._body_text_style.font.name = 'Arial'
        self._body_text_style.font.size = Pt(11)
        self._body_text_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        self._body_text_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Définir le style de liste à puce
        self._bullet_list_style = self._document.styles.add_style('BulletListStyle', 1)
        self._bullet_list_style.font.name = 'Arial'
        self._bullet_list_style.font.size = Pt(9)
        self._bullet_list_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        self._bullet_list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self._bullet_list_style.paragraph_format.left_indent = Inches(0.25)

        self._table_grid_style = self._document.styles.add_style('TableGrid', 3)
        self._table_grid_style.font.name = 'Arial'
        self._table_grid_style.font.size = Pt(9)
        self._table_grid_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        self._table_grid_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self._figure_style = self._document.styles.add_style('FigureStyle', 1)
        self._figure_style.font.name = 'Arial'
        self._figure_style.font.size = Pt(9)
        self._figure_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        self._figure_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self._styles['TitleStyle'] = self._title_style
        self._styles['CaptionStyle'] = self._caption_style
        self._styles['BodyTextStyle'] = self._body_text_style
        self._styles['BulletListStyle'] = self._bullet_list_style
        self._styles['TableGrid'] = self._table_grid_style
        self._styles['FigureStyle'] = self._figure_style

    def set_font(self, fontname:str='Arial', fontsize:int=12):
        """ Définir la police et la taille de la police pour les styles de texte. """

        for style in self._styles.values():
            style.font.name = fontname
            style.font.size = Pt(fontsize)

    def fill_first_page(self, main_title:str, author:str):
        """ 
        Remplir la première page du document. 
        
        Ajouter le titre, l'auteur et la date.
        
        """
        if self._has_first_page:
            return

        # Récupérer le nom de l'utilisateur
        user_name = os.getlogin()

        # Récupérer le nom de l'ordinateur
        computer_name = socket.gethostname()

        logo_path = Path(__file__).parent / 'wolf_report.png'

        self._main_title = main_title
        self._author = author
        self._date = datetime.now().strftime('%d/%m/%Y')

        self._insert_title(self._main_title, level=0, index=0)
        self._insert_figure(logo_path,caption=None, width=2.0, index=1)

        self._insert_paragraph('Ce document a été généré automatiquement par le paquet Python "wolfhece".', index =2)
        self._insert_paragraph(' ', index=3)
        self._insert_paragraph(f'Auteur : {self._author}', index=4)
        self._insert_paragraph(f'Date : {self._date}', index=5)
        self._insert_paragraph(' ', index=6)
        self._insert_paragraph(f'Utilisateur : {user_name}', index=7)
        self._insert_paragraph(f'Ordinateur : {computer_name}', index=8)
        self._insert_paragraph(' ', index=9)

        chain_hash = hashlib.md5(self._main_title.encode() + 
                                 self._author.encode() +
                                 user_name.encode() + 
                                 computer_name.encode()+
                                 self._date.encode()).hexdigest()
        
        self._insert_paragraph('Hash du document : ' + chain_hash, index=10)

        self._insert_new_page(index=11)

        self._insert_paragraph('summary', index=12)

        self._has_first_page = True

    def _insert_title(self, title:str, level:int=1, index:int = 0):
        """ Insère un titre dans le document. """

        self._content.insert(index, ('title', title, level))
    
    def _insert_paragraph(self, paragraph_text:str, style:str='BodyTextStyle', index:int = 0):
        """ Insère un paragraphe dans le document. """

        self._content.insert(index, ('paragraph', paragraph_text, style))

    def _insert_figure(self, image_path:Union[str, Path, Image.Image, Figure], caption:str, width:float=7.0, index:int = 0):
        """ Insère une figure dans le document. """

        self._content.insert(index, ('figure', image_path, caption, width, self._idx_figure))

    def add_title(self, title:str, level:int=1):
        """ Ajoute un titre au document. """
        
        self._content.append(('title', title, level))

    def _list_titles(self, level:int=None):
        """ Renvoie la liste des titres du document. """

        if level is None:
            return [item[1] for item in self._content if item[0] == 'title']
        else:
            return [item[1] for item in self._content if item[0] == 'title' and item[2] == level]
        
    def _list_captions(self):
        """ Renvoie la liste des légendes de figures du document. """

        return [item[2] for item in self._content if item[0] == 'figure' if item[2]]

    def _list_figures(self):
        """ Renvoie la liste des figures du document. """

        return [item[1] for item in self._content if item[0] == 'figure' if item[1] and item[2]]

    def _list_index(self):
        """ Renvoie la liste des index de figures du document. """

        return [item[3] for item in self._content if item[0] == 'figure' if item[3]]
    
    def fig_exists(self, fig_name:str):
        """ Vérifie si une figure existe dans le document. """

        return fig_name in self._list_figures()
    
    def get_fig_index(self, fig_name_caption:str):
        """ Renvoie la légende d'une figure. """

        list_figures = self._list_figures()
        list_captions = self._list_captions()

        if fig_name_caption in list_figures:
            idx = self._list_figures().index(fig_name_caption)+1
        elif fig_name_caption in list_captions:
            idx = self._list_captions().index(fig_name_caption)+1
        else:
            idx = None

        return idx
    
    def _add_summary(self):
        """ Ajoute un sommaire au document. """

        titles = self._list_titles()
        
        self._document.add_heading(_('Summary'), level=1).style = 'TitleStyle'
    
        for cur_title in titles:
            p = self._document.add_paragraph(cur_title, style='BodyTextStyle')
            run = p.add_run()
            run.add_tab()
            run.bold = True
            p.style = 'BodyTextStyle'
        
        self._document.add_heading(_('List of figures'), level=1).style = 'TitleStyle'
        figures = self._list_captions()
        for i, cur_figure in enumerate(figures):
            p = self._document.add_paragraph(f'Fig. {i+1} : {cur_figure}', style='BodyTextStyle')
            run = p.add_run()
            run.add_tab()
            run.bold = True
            p.style = 'BodyTextStyle'

        self._document.add_page_break()
 
    def add_paragraph(self, paragraph_text:str, style:str='BodyTextStyle'):
        """ Ajoute un paragraphe au document. """

        self._content.append(('paragraph', paragraph_text, style))
 
    def add(self, paragraph_text:str, style:str='BodyTextStyle'):
        """ Ajoute un paragraphe au document. """

        self.add_paragraph(paragraph_text, style=style)

    def __add__(self, paragraph_text:str):
        """ Surcharge de l'opérateur + pour ajouter un paragraphe. """

        self.add_paragraph(paragraph_text)

        return self

    def add_figure(self, image_path:Union[str, Path, Image.Image, Figure], caption:str, width:float=7.0):
        """ Ajoute une figure au document avec une légende. """

        if caption:
            self._idx_figure += 1
        
        self._content.append(('figure', image_path, caption, width, self._idx_figure))

    def add_bullet_list(self, bullet_list: List[str], style:str='BulletListStyle'):
        """ Ajoute une liste à puce au document. """

        for item in bullet_list:
            self.add_paragraph('- ' + item, style=style)

    def add_new_page(self):
        """ Ajoute une nouvelle page au document. """

        self._content.append(('newpage', '', None))

    def _insert_new_page(self, index:int = 0):
        """ Insère une nouvelle page au document. """

        self._content.insert(index, ('newpage', '', None))
    
    def add_table_from_listoflists(self, data:List[List[str]], style:str='TableGrid'):
        """ 
        Ajoute un tableau au document. 
        
        :param data: Liste de listes contenant les données du tableau. Chaque liste est une ligne du tableau.

        """

        self._content.append(('table', data, style))

    def add_table_from_dict(self, data:dict, style:str='TableGrid'):
        """ 
        Ajoute un tableau au document. 
        
        :param data: Dictionnaire contenant les données du tableau. Les clés sont les en-têtes de colonnes.

        """

        table_data = [list(data.keys())]
        table_data += [list(data.values())]
        self.add_table_from_listoflists(table_data, style=style)

    def add_table_as_picture(self, data:Union[List[List[str]], dict, pd.DataFrame, Figure], caption:str=None):
        """ Ajoute un tableau au document sous forme d'image. """

        def fig2img(fig):
            """Convert a Matplotlib figure to a PIL Image and return it"""
            import io
            buf = io.BytesIO()
            
            fig.savefig(buf, bbox_inches='tight')
            buf.seek(0)
            img = Image.open(buf)
            return img        

        if isinstance(data, Figure):
            tmp_image = fig2img(data)
            self.add_figure(tmp_image, caption)
            return
        
        if isinstance(data, dict):
            data = pd.DataFrame(data)
        elif isinstance(data, list):
            data = pd.DataFrame(data)

        fig, ax = plt.subplots()

        ax.axis('off')
        ax.table(cellText=data.values, 
                 colLabels=data.columns, 
                 loc='center', 
                 cellLoc='center', 
                 colColours=['#f3f3f3']*len(data.columns))
        
        fig.tight_layout()
        
        tmp_image = fig2img(fig)

        self.add_figure(tmp_image, caption, width=4.0)

    def _apply_text_styles(self, paragraph, text):
        """ Search for bold and italic styles in the text and apply them."""

        text = text.replace('\n\n', 'DOUBLE_NEWLINE')
        text = text.replace('\n', ' ')
        text = text.replace('DOUBLE_NEWLINE', '\n')

        def split_bold(text):
            return text.split('**')
        
        def split_italic(text):
            return text.split('*')
        
        splitted_bold = split_bold(text)

        bold = False
        for cur_text in splitted_bold:
            if cur_text != '':
                italic = False
                spliited_italic = split_italic(cur_text)
                for cur_text2 in spliited_italic:
                    if cur_text2 != '':
                        run = paragraph.add_run(cur_text2)
                        run.bold = bold
                        run .italic = italic
                    
                    italic = not italic
            bold = not bold
 
    def parse_content(self):
        """ Parse le contenu du document et l'ajoute au document Word. """

        # tmp_dir = TemporaryDirectory()

        for item in self._content:

            if item[0] == 'title':
                self._document.add_heading(item[1], level=item[2]).style = 'TitleStyle'
            
            elif item[0] == 'paragraph':

                if item[1] == 'summary':
                    self._add_summary()
                    continue
                else:
                    p = self._document.add_paragraph()
                    self._apply_text_styles(p, item[1])
                    p.style = item[2] if item[2] else 'BodyTextStyle'
            
            elif item[0] == 'figure':

                if isinstance(item[1], Image.Image):
                    
                    tmp_name = NamedTemporaryFile(suffix='.png').name
                    item[1].save(tmp_name)

                elif isinstance(item[1], str):
                    tmp_name = item[1]

                elif isinstance(item[1], Path):
                    tmp_name = str(item[1])

                elif isinstance(item[1], Figure):
                    item[1].tight_layout()
                    tmp_name = NamedTemporaryFile(suffix='.png').name
                    item[1].savefig(tmp_name)

                if Path(tmp_name).exists():
                    self._document.add_picture(tmp_name, width=Inches(item[3]) if item[3] else Inches(7.0))
                    self._document.paragraphs[-1].style = 'FigureStyle'
                else:
                    logging.error(f"File {tmp_name} not found.")
                    p = self._document.add_paragraph()
                    run = p.add_run(f'Error: Image not found. {tmp_name}')
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    p.style = 'BodyTextStyle'

                if item[2]:
                    caption = self._document.add_paragraph(f'Fig. {item[4]} :' + item[2])
                    caption.style = 'CaptionStyle'

            elif item[0] == 'table':
                
                data = item[1]
                style = item[2]
                table = self._document.add_table(rows=len(data), cols=len(data[0]))
                table.style = style

                for i, row in enumerate(data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell

            elif item[0] == 'newpage':
                self._document.add_page_break()
 
    def save(self, file_path:Union[str,Path]=None):
        """ Sauvegarde le document Word. """
        
        if file_path is None:
            file_path = self._filename

        if file_path is None:
            raise ValueError("Le chemin du fichier n'a pas été spécifié.")
        
        self._document = Document()
        
        self._define_default_styles()

        self.fill_first_page(self._main_title, self._author)

        self.parse_content()
        try:
            self._document.save(str(file_path))
        except Exception as e:
            logging.error(f"Error saving file: {e}")
 
if __name__ == '__main__':

    # Exemple d'utilisation
    rapport = RapidReport('Rapport de Projet', 'Alice')

    rapport.add_title('Titre Principal', level=0)
    rapport.add_paragraph('Ceci est un **paragraphe** introductif avec des mots en *italique* et en **gras**.')
    
    rapport += "Tentative d'ajout de figure vie un lien incorrect.\nPassage à la ligne"
    rapport.add_figure('/path/to/image.png', 'Légende de la figure.')
    rapport+=""" 
Commentraire sur la figure multilignes
ligne 2
ligne3"""

    rapport.add_bullet_list(['Premier élément', 'Deuxième élément', 'Troisième élément'])
    
    rapport.add_table_from_listoflists([['Nom', 'Âge'], ['Alice', '25'], ['Bob', '30']])
    rapport.add_table_from_dict({'Nom': ['Alice', 'Bob'], 'Âge': ['25', '30']})
    rapport.add_table_as_picture({'Nom': ['Alice', 'Bob'], 'Âge': ['25', '30']}, caption='Tableau de données')

    rapport.save('rapport.docx')

    assert rapport.get_fig_index('/path/to/image.png') == 1
    assert rapport.get_fig_index('Tableau de données') == 2

