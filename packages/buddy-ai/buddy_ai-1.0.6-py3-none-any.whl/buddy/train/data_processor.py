"""
Data processor for training data preparation.
Handles ALL file types by trying multiple encodings and traversing directories recursively.
Only skips files that cannot be read with any encoding.
"""

import json
import os
import csv
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Tuple
import logging
import re
from dataclasses import dataclass
import chardet

try:
    import pandas as pd
    pandas_available = True
except ImportError:
    pandas_available = False

try:
    import docx
    python_docx_available = True
except ImportError:
    python_docx_available = False

try:
    import PyPDF2
    import pdfplumber
    pdf_available = True
except ImportError:
    pdf_available = False

from buddy.utils.log import logger


@dataclass
class ProcessedData:
    """Container for processed training data."""
    texts: List[str]
    metadata: List[Dict[str, Any]]
    stats: Dict[str, Any]


class DataProcessor:
    """
    Processes ALL file types for training data by trying multiple encodings.
    Traverses directories recursively and attempts to read any file as text.
    """
    
    def __init__(self, min_text_length: int = 10, max_text_length: int = 10000):
        self.min_text_length = min_text_length
        self.max_text_length = max_text_length
        # Try these encodings in order of most common first
        self.encodings_to_try = [
            'utf-8', 'utf-8-sig', 'utf-16', 'utf-16-le', 'utf-16-be', 'utf-32',
            'latin-1', 'cp1252', 'ascii', 'iso-8859-1', 'cp437', 'cp850',
            'cp1251', 'shift_jis', 'gb2312', 'big5', 'euc-jp', 'euc-kr',
            'iso-8859-2', 'iso-8859-15', 'koi8-r', 'mac-roman'
        ]
        # Known binary file signatures to skip early (magic numbers)
        self.binary_signatures = [
            b'\x89PNG',  # PNG
            b'\xff\xd8\xff',  # JPEG
            b'GIF8',  # GIF
            b'BM',  # BMP
            b'RIFF',  # WAV/AVI
            b'\x00\x00\x01\x00',  # ICO
            b'PK\x03\x04',  # ZIP
            b'PK\x05\x06',  # ZIP
            b'PK\x07\x08',  # ZIP
            b'\x1f\x8b',  # GZIP
            b'\x42\x5a\x68',  # BZIP2
            b'\x50\x4b\x03\x04',  # ZIP
            b'\x4d\x5a',  # EXE
            b'\x7f\x45\x4c\x46',  # ELF
            b'\xfe\xed\xfa',  # Mach-O
            b'\xcf\xfa\xed\xfe',  # Mach-O
            b'\xca\xfe\xba\xbe',  # Java class
            b'\x89\x48\x44\x46',  # HDF5
        ]
        
        self.processed_files = 0
        self.skipped_files = 0
        self.total_text_length = 0
    
    def process_directory(self, directory_path: str) -> ProcessedData:
        """
        Process all files in a directory recursively using multiple encoding attempts.
        
        Args:
            directory_path: Path to directory containing training data
            
        Returns:
            ProcessedData object with processed texts and metadata
        """
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        all_texts = []
        all_metadata = []
        
        # Statistics
        processed_files = 0
        skipped_files = 0
        total_chars = 0
        encoding_stats = {}
        file_type_stats = {}
        
        logger.info(f"Processing directory: {directory_path}")
        
        # Process all files recursively
        for file_path in directory.rglob('*'):
            if file_path.is_file():
                try:
                    # Try to process the file
                    texts, metadata, encoding_used = self._process_file_with_encoding_detection(file_path)
                    if texts:
                        all_texts.extend(texts)
                        all_metadata.extend(metadata)
                        processed_files += 1
                        total_chars += sum(len(text) for text in texts)
                        
                        # Track encoding usage
                        encoding_stats[encoding_used] = encoding_stats.get(encoding_used, 0) + 1
                        
                        # Track file type usage
                        file_ext = file_path.suffix.lower() or 'no_extension'
                        file_type_stats[file_ext] = file_type_stats.get(file_ext, 0) + 1
                        
                        logger.debug(f"Processed ({encoding_used}): {file_path}")
                    else:
                        skipped_files += 1
                        logger.debug(f"Skipped (no valid content): {file_path}")
                except Exception as e:
                    skipped_files += 1
                    logger.debug(f"Skipped (error): {file_path} - {e}")
        
        stats = {
            'total_files_found': processed_files + skipped_files,
            'processed_files': processed_files,
            'skipped_files': skipped_files,
            'total_texts': len(all_texts),
            'total_characters': total_chars,
            'avg_text_length': total_chars / len(all_texts) if all_texts else 0,
            'encoding_stats': encoding_stats,
            'file_type_stats': file_type_stats
        }
        
        logger.info(f"Processing complete. Stats: {stats}")
        
        return ProcessedData(
            texts=all_texts,
            metadata=all_metadata,
            stats=stats
        )
    
    def _process_file_with_encoding_detection(self, file_path: Path) -> Tuple[List[str], List[Dict[str, Any]], str]:
        """
        Process a single file by trying multiple encoding methods.
        
        Returns:
            Tuple of (texts, metadata, encoding_used)
        """
        texts = []
        metadata = []
        encoding_used = "unknown"
        
        # Skip hidden files and system files
        if file_path.name.startswith('.'):
            return texts, metadata, encoding_used
        
        # Skip if file is too large (>100MB) to avoid memory issues
        try:
            if file_path.stat().st_size > 100 * 1024 * 1024:
                logger.debug(f"Skipping large file: {file_path}")
                return texts, metadata, encoding_used
        except:
            return texts, metadata, encoding_used
        
        # Check if file is binary by reading first few bytes
        try:
            with open(file_path, 'rb') as f:
                first_bytes = f.read(1024)
                if self._is_binary_file(first_bytes):
                    logger.debug(f"Skipping binary file: {file_path}")
                    return texts, metadata, encoding_used
        except:
            return texts, metadata, encoding_used
        
        # Try specialized extractors first
        text_content = None
        
        # PDF files
        if file_path.suffix.lower() == '.pdf' and pdf_available:
            text_content, encoding_used = self._extract_pdf_text(file_path)
        
        # Word documents
        elif file_path.suffix.lower() == '.docx' and python_docx_available:
            text_content, encoding_used = self._extract_docx_text(file_path)
        
        # If no specialized extractor worked, try encoding detection
        if text_content is None:
            text_content, encoding_used = self._extract_text_with_encoding_detection(file_path)
        
        if text_content:
            # Clean and validate text
            text_content = self._clean_text(text_content)
            
            # Check length constraints
            if len(text_content.strip()) >= self.min_text_length:
                # Split long texts into chunks
                text_chunks = self._split_text_into_chunks(text_content, self.max_text_length)
                
                for i, chunk in enumerate(text_chunks):
                    texts.append(chunk)
                    metadata.append({
                        'source_file': str(file_path),
                        'file_type': file_path.suffix.lower() or 'no_extension',
                        'encoding': encoding_used,
                        'chunk_index': i,
                        'total_chunks': len(text_chunks),
                        'file_size': file_path.stat().st_size,
                        'char_count': len(chunk),
                        'word_count': len(chunk.split())
                    })
        
        return texts, metadata, encoding_used
    
    def _is_binary_file(self, data: bytes) -> bool:
        """
        Check if file is binary by examining content.
        """
        # Check for known binary signatures
        for signature in self.binary_signatures:
            if data.startswith(signature):
                return True
        
        # Check for high percentage of null bytes or non-printable characters
        if len(data) == 0:
            return True
        
        # Count null bytes
        null_count = data.count(0)
        if null_count > len(data) * 0.3:  # More than 30% null bytes
            return True
        
        # Check for too many non-printable characters
        printable_count = sum(1 for byte in data if 32 <= byte <= 126 or byte in [9, 10, 13])
        if printable_count < len(data) * 0.7:  # Less than 70% printable
            return True
        
        return False
    
    def _extract_text_with_encoding_detection(self, file_path: Path) -> Tuple[Optional[str], str]:
        """
        Try to extract text using various encodings.
        """
        # First, try to detect encoding using chardet
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                if len(raw_data) > 0:
                    detected = chardet.detect(raw_data)
                    if detected and detected['encoding'] and detected['confidence'] > 0.7:
                        try:
                            return raw_data.decode(detected['encoding']), detected['encoding']
                        except:
                            pass
        except:
            pass
        
        # Try each encoding manually
        for encoding in self.encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                    # Validate that we got meaningful content
                    if content and len(content.strip()) > 0:
                        # Check if content seems reasonable (not too many weird characters)
                        printable_ratio = sum(1 for c in content[:1000] if c.isprintable() or c.isspace()) / min(len(content), 1000)
                        if printable_ratio > 0.8:  # At least 80% printable characters
                            return content, encoding
            except:
                continue
        
        # Try reading as raw bytes and converting to string with errors ignored
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                content = raw_data.decode('utf-8', errors='ignore')
                if content and len(content.strip()) > 0:
                    return content, 'utf-8-ignore'
        except:
            pass
        
        return None, "failed"
    
    def _extract_pdf_text(self, file_path: Path) -> Tuple[Optional[str], str]:
        """Extract text from PDF files."""
        try:
            text = ""
            
            # Try pdfplumber first (better text extraction)
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text, "pdf-pdfplumber"
            except:
                pass
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text, "pdf-pypdf2"
            except:
                pass
            
        except Exception as e:
            logger.debug(f"Error extracting PDF text from {file_path}: {e}")
        
        return None, "pdf-failed"
    
    def _extract_docx_text(self, file_path: Path) -> Tuple[Optional[str], str]:
        """Extract text from Word documents."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                return text, "docx"
        except Exception as e:
            logger.debug(f"Error extracting DOCX text from {file_path}: {e}")
        
        return None, "docx-failed"
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        if not text:
            return ""
        
        # Remove control characters except common ones
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize whitespace
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double newline
        
        # Remove lines with only whitespace
        lines = text.split('\n')
        cleaned_lines = [line.rstrip() for line in lines if line.strip()]
        
        # Rejoin with single newlines
        text = '\n'.join(cleaned_lines)
        
        return text.strip()
    
    def _split_text_into_chunks(self, text: str, max_length: int) -> List[str]:
        """Split text into chunks if it's too long."""
        if len(text) <= max_length:
            return [text]
        
        chunks = []
        words = text.split()
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            
            if current_length + word_length > max_length and current_chunk:
                # Finish current chunk
                chunks.append(' '.join(current_chunk))
                current_chunk = [word]
                current_length = len(word)
            else:
                current_chunk.append(word)
                current_length += word_length
        
        # Add the last chunk
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def save_processed_data(self, processed_data: ProcessedData, output_path: str) -> None:
        """Save processed data to JSON file."""
        output_data = {
            'texts': processed_data.texts,
            'metadata': processed_data.metadata,
            'stats': processed_data.stats,
            'processing_info': {
                'total_files': processed_data.stats.get('total_files_found', 0),
                'processed_files': processed_data.stats.get('processed_files', 0),
                'skipped_files': processed_data.stats.get('skipped_files', 0),
                'total_characters': processed_data.stats.get('total_characters', 0),
                'encoding_distribution': processed_data.stats.get('encoding_stats', {}),
                'file_type_distribution': processed_data.stats.get('file_type_stats', {})
            }
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved processed data to {output_path}")
    
    def load_processed_data(self, file_path: str) -> ProcessedData:
        """Load processed data from JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return ProcessedData(
            texts=data['texts'],
            metadata=data['metadata'],
            stats=data.get('stats', {})
        )
