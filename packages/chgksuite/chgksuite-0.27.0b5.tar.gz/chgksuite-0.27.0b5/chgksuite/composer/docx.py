import os
import re
import shutil
import sys
import tempfile
import zipfile

import docx
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt as DocxPt

import chgksuite.typotools as typotools
from chgksuite.common import DummyLogger, log_wrap, replace_escaped
from chgksuite.composer.composer_common import (
    BaseExporter,
    _parse_4s_elem,
    backtick_replace,
    parseimg,
    remove_accents_standalone,
)

WHITEN = {
    "handout": False,
    "zachet": True,
    "nezachet": True,
    "comment": True,
    "source": True,
    "author": False,
}


def replace_font_in_docx(template_path, new_font):
    """Replace Arial fonts with specified font in docx template"""
    temp_dir = tempfile.mkdtemp()
    template_name = os.path.basename(template_path)
    temp_template = os.path.join(temp_dir, template_name)
    shutil.copy2(template_path, temp_template)

    temp_zip = os.path.join(temp_dir, "template.zip")
    os.rename(temp_template, temp_zip)
    with zipfile.ZipFile(temp_zip, "r") as zip_ref:
        zip_ref.extractall(temp_dir)
    os.remove(temp_zip)

    for root, _, files in os.walk(temp_dir):
        for file in files:
            if file.endswith(".xml"):
                file_path = os.path.join(root, file)
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()

                    content = content.replace("Arial Unicode MS", new_font)
                    content = content.replace("Arial", new_font)

                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(content)
                except UnicodeError:
                    continue

    shutil.make_archive(temp_template, "zip", temp_dir)
    os.rename(temp_template + ".zip", temp_template)
    return temp_template


def replace_no_break_standalone(s, replace_spaces=True, replace_hyphens=True):
    """Standalone version of _replace_no_break"""
    return typotools.replace_no_break(s, spaces=replace_spaces, hyphens=replace_hyphens)


def get_label_standalone(
    question, field, labels, language="ru", only_question_number=False, number=None
):
    """Standalone version of get_label"""
    if field == "question" and only_question_number:
        return str(question.get("number") or number)
    if field in ("question", "tour"):
        lbl = (question.get("overrides") or {}).get(field) or labels["question_labels"][
            field
        ]
        num = question.get("number") or number
        if language in ("uz", "uz_cyr"):
            return f"{num} – {lbl}"
        elif language == "kz":
            return f"{num}-{lbl}"
        else:
            return f"{lbl} {num}"
    if field in (question.get("overrides") or {}):
        return question["overrides"][field]
    if field == "source" and isinstance(question.get("source" or ""), list):
        return labels["question_labels"]["sources"]
    return labels["question_labels"][field]


def remove_square_brackets_standalone(s, regexes):
    """Standalone version of remove_square_brackets"""
    hs = regexes["handout_short"]
    s = s.replace("\\[", "LEFTSQUAREBRACKET")
    s = s.replace("\\]", "RIGHTSQUAREBRACKET")
    # Use placeholder to preserve handout brackets during removal
    s = re.sub(f"\\[({hs}.+?)\\]", "{HANDOUT_PLACEHOLDER\\1}", s, flags=re.DOTALL)
    i = 0
    while "[" in s and "]" in s and i < 10:
        s = re.sub(" *\\[.+?\\]", "", s, flags=re.DOTALL)
        s = s.strip()
        i += 1
    if i == 10:
        sys.stderr.write(
            f"Error replacing square brackets on question: {s}, retries exceeded\n"
        )
    # Restore handout brackets - get the original matched text from the placeholder
    s = re.sub(
        r"\{HANDOUT_PLACEHOLDER(.+?)\}",
        lambda m: "[" + m.group(1) + "]",
        s,
        flags=re.DOTALL,
    )
    s = s.replace("LEFTSQUAREBRACKET", "[")
    s = s.replace("RIGHTSQUAREBRACKET", "]")
    return s


def add_hyperlink_to_docx(doc, paragraph, text, url):
    """Standalone version of add_hyperlink"""
    run = paragraph.add_run(text)
    run.style = doc.styles["Hyperlink"]
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    return hyperlink


def format_docx_element(
    doc,
    el,
    para,
    whiten,
    spoilers="none",
    logger=None,
    labels=None,
    regexes=None,
    language="ru",
    remove_accents=False,
    remove_brackets=False,
    replace_no_break_spaces=False,
    **kwargs,
):
    """
    Standalone version of docx_format that can be used outside DocxExporter.

    Args:
        doc: docx Document object
        el: Element to format
        para: Paragraph object to add content to
        whiten: Whether to apply whitening
        spoilers: Spoiler handling mode ("none", "whiten", "dots", "pagebreak")
        logger: Logger instance
        labels: Labels dictionary
        regexes: Regexes dictionary (for handout_short)
        language: Language code
        remove_accents: Whether to remove accents
        remove_brackets: Whether to remove square brackets
        replace_no_break_spaces: Whether to replace non-breaking spaces
        **kwargs: Additional arguments (tmp_dir, targetdir, etc.)
    """
    if logger is None:
        logger = DummyLogger()

    if isinstance(el, list):
        if len(el) > 1 and isinstance(el[1], list):
            format_docx_element(
                doc,
                el[0],
                para,
                whiten,
                spoilers,
                logger,
                labels,
                regexes,
                language,
                remove_accents,
                remove_brackets,
                replace_no_break_spaces,
                **kwargs,
            )
            licount = 0
            for li in el[1]:
                licount += 1
                para.add_run("\n{}. ".format(licount))
                format_docx_element(
                    doc,
                    li,
                    para,
                    whiten,
                    spoilers,
                    logger,
                    labels,
                    regexes,
                    language,
                    remove_accents,
                    remove_brackets,
                    replace_no_break_spaces,
                    **kwargs,
                )
        else:
            licount = 0
            for li in el:
                licount += 1
                para.add_run("\n{}. ".format(licount))
                format_docx_element(
                    doc,
                    li,
                    para,
                    whiten,
                    spoilers,
                    logger,
                    labels,
                    regexes,
                    language,
                    remove_accents,
                    remove_brackets,
                    replace_no_break_spaces,
                    **kwargs,
                )

    if isinstance(el, str):
        logger.debug("parsing element {}:".format(log_wrap(el)))

        if remove_accents and regexes:
            el = remove_accents_standalone(el, regexes)
        if remove_brackets and regexes:
            el = remove_square_brackets_standalone(el, regexes)
        else:
            el = replace_escaped(el)

        el = backtick_replace(el)

        for run in _parse_4s_elem(el, logger=logger):
            if run[0] == "pagebreak":
                if spoilers == "dots":
                    for _ in range(30):
                        para = doc.add_paragraph()
                        para.add_run(".")
                    para = doc.add_paragraph()
                else:
                    para = doc.add_page_break()
            elif run[0] == "linebreak":
                para.add_run("\n")
            elif run[0] == "screen":
                if remove_accents or remove_brackets:
                    text = run[1]["for_screen"]
                else:
                    text = run[1]["for_print"]
                if replace_no_break_spaces:
                    text = replace_no_break_standalone(text)
                r = para.add_run(text)
            elif run[0] == "hyperlink" and not (whiten and spoilers == "whiten"):
                r = add_hyperlink_to_docx(doc, para, run[1], run[1])
            elif run[0] == "img":
                if run[1].endswith(".shtml"):
                    r = para.add_run("(ТУТ БЫЛА ССЫЛКА НА ПРОТУХШУЮ КАРТИНКУ)\n")
                    continue
                parsed_image = parseimg(
                    run[1],
                    dimensions="inches",
                    tmp_dir=kwargs.get("tmp_dir"),
                    targetdir=kwargs.get("targetdir"),
                )
                imgfile = parsed_image["imgfile"]
                width = parsed_image["width"]
                height = parsed_image["height"]
                inline = parsed_image["inline"]
                if inline:
                    r = para.add_run("")
                else:
                    r = para.add_run("\n")

                try:
                    if inline:
                        r.add_picture(imgfile, height=Inches(1.0 / 6))
                    else:
                        r.add_picture(
                            imgfile, width=Inches(width), height=Inches(height)
                        )
                except UnrecognizedImageError:
                    sys.stderr.write(
                        f"python-docx can't recognize header for {imgfile}\n"
                    )
                if not inline:
                    r = para.add_run("\n")
                continue
            else:
                text = run[1]
                if replace_no_break_spaces:
                    text = replace_no_break_standalone(text)
                r = para.add_run(text)
                if "italic" in run[0]:
                    r.italic = True
                if "bold" in run[0]:
                    r.bold = True
                if "underline" in run[0]:
                    r.underline = True
                if run[0] == "strike":
                    r.font.strike = True
                if run[0] == "sc":
                    r.small_caps = True
                if whiten and spoilers == "whiten":
                    r.style = "Whitened"


def add_question_to_docx(
    doc,
    question_data,
    labels,
    regexes=None,
    qcount=None,
    skip_qcount=False,
    screen_mode=False,
    external_para=None,
    noparagraph=False,
    noanswers=False,
    spoilers="none",
    language="ru",
    only_question_number=False,
    add_question_label=True,
    logger=None,
    **kwargs,
):
    """
    Standalone function to add a question to a docx document.

    Args:
        doc: docx Document object
        question_data: Dictionary containing question data
        labels: Labels dictionary
        regexes: Regexes dictionary (for handout_short)
        qcount: Current question count (will be incremented if not skip_qcount)
        skip_qcount: Whether to skip incrementing question count
        screen_mode: Whether to use screen mode formatting
        external_para: External paragraph to use instead of creating new ones
        noparagraph: Whether to skip paragraph breaks
        noanswers: Whether to skip adding answers
        spoilers: Spoiler handling mode ("none", "whiten", "dots", "pagebreak")
        language: Language code
        only_question_number: Whether to show only question numbers
        logger: Logger instance
        **kwargs: Additional arguments passed to format_docx_element

    Returns:
        Updated question count
    """
    if not kwargs.get("tmp_dir"):
        kwargs["tmp_dir"] = tempfile.mkdtemp()
    if not kwargs.get("targetdir"):
        kwargs["targetdir"] = os.getcwd()
    if logger is None:
        logger = DummyLogger()

    q = question_data
    if external_para is None:
        p = doc.add_paragraph()
    else:
        p = external_para
    if add_question_label:
        p.paragraph_format.space_before = DocxPt(18)
    p.paragraph_format.keep_together = True

    # Handle question numbering
    if qcount is None:
        qcount = 1
    if "number" not in q and not skip_qcount:
        qcount += 1
    if "setcounter" in q:
        qcount = int(q["setcounter"])

    # Add question label
    if add_question_label:
        question_label = get_label_standalone(
            q,
            "question",
            labels,
            language,
            only_question_number,
            number=qcount if "number" not in q else q["number"],
        )
        p.add_run(f"{question_label}. ").bold = True

    # Add handout if present
    if "handout" in q:
        handout_label = get_label_standalone(q, "handout", labels, language)
        p.add_run(f"\n[{handout_label}: ")
        format_docx_element(
            doc,
            q["handout"],
            p,
            WHITEN["handout"],
            spoilers,
            logger,
            labels,
            regexes,
            language,
            remove_accents=screen_mode,
            remove_brackets=screen_mode,
            **kwargs,
        )
        p.add_run("\n]")

    if not noparagraph:
        p.add_run("\n")

    # Add question text
    format_docx_element(
        doc,
        q["question"],
        p,
        False,
        spoilers,
        logger,
        labels,
        regexes,
        language,
        remove_accents=screen_mode,
        remove_brackets=screen_mode,
        replace_no_break_spaces=True,
        **kwargs,
    )

    # Add answers and other fields if not disabled
    if not noanswers:
        if spoilers == "pagebreak":
            p = doc.add_page_break()
        elif spoilers == "dots":
            for _ in range(30):
                if external_para is None:
                    p = doc.add_paragraph()
                else:
                    p.add_run("\n")
                p.add_run(".")
            if external_para is None:
                p = doc.add_paragraph()
            else:
                p.add_run("\n")
        else:
            if external_para is None:
                p = doc.add_paragraph()
            else:
                p.add_run("\n")

        p.paragraph_format.keep_together = True
        p.paragraph_format.space_before = DocxPt(6)

        # Add answer
        answer_label = get_label_standalone(q, "answer", labels, language)
        p.add_run(f"{answer_label}: ").bold = True
        format_docx_element(
            doc,
            q["answer"],
            p,
            True,
            spoilers,
            logger,
            labels,
            regexes,
            language,
            remove_accents=screen_mode,
            replace_no_break_spaces=True,
            **kwargs,
        )

        # Add other fields
        for field in ["zachet", "nezachet", "comment", "source", "author"]:
            if field in q:
                if field == "source":
                    if external_para is None:
                        p = doc.add_paragraph()
                        p.paragraph_format.keep_together = True
                    else:
                        p.add_run("\n")
                else:
                    p.add_run("\n")

                field_label = get_label_standalone(q, field, labels, language)
                p.add_run(f"{field_label}: ").bold = True
                format_docx_element(
                    doc,
                    q[field],
                    p,
                    WHITEN[field],
                    spoilers,
                    logger,
                    labels,
                    regexes,
                    language,
                    remove_accents=screen_mode,
                    remove_brackets=screen_mode,
                    replace_no_break_spaces=field != "source",
                    **kwargs,
                )

    return qcount


class DocxExporter(BaseExporter):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.qcount = 0

        if self.args.font_face:
            self.args.docx_template = replace_font_in_docx(
                self.args.docx_template, self.args.font_face
            )

    def __del__(self):
        # Cleanup temp directory if it exists
        if hasattr(self, "temp_dir") and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def _docx_format(self, *args, **kwargs):
        kwargs.update(self.dir_kwargs)
        return format_docx_element(
            self.doc,
            *args,
            spoilers=self.args.spoilers,
            logger=self.logger,
            labels=self.labels,
            regexes=self.regexes,
            language=self.args.language,
            **kwargs,
        )

    def docx_format(self, el, para, whiten, **kwargs):
        # Redirect to standalone function
        return format_docx_element(
            self.doc,
            el,
            para,
            whiten,
            spoilers=self.args.spoilers,
            logger=self.logger,
            labels=self.labels,
            regexes=self.regexes,
            language=self.args.language,
            **kwargs,
        )

    def add_hyperlink(self, paragraph, text, url):
        return add_hyperlink_to_docx(self.doc, paragraph, text, url)

    def add_question(
        self, element, skip_qcount=False, screen_mode=False, external_para=None
    ):
        self.qcount = add_question_to_docx(
            self.doc,
            element[1],
            self.labels,
            self.regexes,
            self.qcount,
            skip_qcount,
            screen_mode,
            external_para,
            self.args.noparagraph,
            self.args.noanswers,
            self.args.spoilers,
            self.args.language,
            self.args.only_question_number,
            self.logger,
            **self.dir_kwargs,
        )

    def _add_question_columns(self, element):
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = True

        def set_cell_border(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            for edge in ["top", "left", "bottom", "right"]:
                border = OxmlElement("w:{}Border".format(edge))
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                tcPr.append(border)

        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)

        table.cell(0, 0).paragraphs[0].add_run("Версия для ведущего\n").bold = True
        table.cell(0, 1).paragraphs[0].add_run("Версия для экрана\n").bold = True

        self.add_question(
            element, screen_mode=False, external_para=table.cell(0, 0).paragraphs[0]
        )
        self.add_question(
            element, screen_mode=True, external_para=table.cell(0, 1).paragraphs[0]
        )

        self.doc.add_paragraph()

    def export(self, outfilename):
        self.logger.debug(self.args.docx_template)
        self.doc = Document(self.args.docx_template)
        self.logger.debug(log_wrap(self.structure))

        firsttour = True
        prev_element = None
        para = None
        page_break_before_heading = False
        for element in self.structure:
            if element[0] == "meta":
                para = self.doc.add_paragraph()
                if prev_element and prev_element[0] == "Question":
                    para.paragraph_format.space_before = DocxPt(18)
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                self.doc.add_paragraph()

            if element[0] in ["editor", "date", "heading", "section"]:
                if element[0] == "heading" and para is not None:
                    page_break_before_heading = True
                if para is None:
                    para = self.doc.paragraphs[0]
                else:
                    para = self.doc.add_paragraph()
                self._docx_format(element[1], para, False, replace_no_break_spaces=True)
                if element[0] == "heading" and page_break_before_heading:
                    para.paragraph_format.page_break_before = True
                if element[0] == "section":
                    if not firsttour:
                        para.paragraph_format.page_break_before = True
                    else:
                        firsttour = False
                if element[0] == "heading":
                    for st in self.doc.styles:
                        if st.name == "Heading 1":
                            break
                    para.style = st
                elif element[0] == "section":
                    for st in self.doc.styles:
                        if st.name == "Heading 2":
                            break
                    para.style = st
                para.paragraph_format.keep_with_next = True
                para.add_run("\n")

            if element[0] == "Question":
                if self.args.screen_mode == "add_versions_columns":
                    self._add_question_columns(element)
                elif self.args.screen_mode == "add_versions":
                    para = self.doc.add_paragraph()
                    para = self.doc.add_paragraph()
                    para.add_run("Версия для ведущего:").bold = True
                    self.add_question(element, screen_mode=False)
                    para = self.doc.add_paragraph()
                    para = self.doc.add_paragraph()
                    para.add_run("Версия для экрана:").bold = True
                    self.add_question(element, skip_qcount=True, screen_mode=True)
                elif self.args.screen_mode == "replace_all":
                    self.add_question(element, screen_mode=True)
                else:
                    self.add_question(element)
            prev_element = element

        self.doc.save(outfilename)
        self.logger.info("Output: {}".format(outfilename))


# Example usage of the extracted DOCX functions:
"""
from docx import Document
import toml
from chgksuite.composer.docx import add_question_to_docx, format_docx_element

# Load labels
with open("labels.toml", encoding="utf8") as f:
    labels = toml.load(f)

# Create a new document
doc = Document()

# Example question data
question_data = {
    "question": "What is the capital of France?",
    "answer": "Paris",
    "comment": "This is a basic geography question",
    "source": "World Geography Book"
}

# Add question to document
qcount = add_question_to_docx(
    doc=doc,
    question_data=question_data,
    labels=labels,
    qcount=0,  # Starting question count
    noanswers=False,  # Include answers
    spoilers="none",  # No spoiler handling
    language="en",
    only_question_number=False
)

# Or use the lower-level formatting function directly
paragraph = doc.add_paragraph()
format_docx_element(
    doc=doc,
    el="This is **bold text** and _italic text_",
    para=paragraph,
    whiten=False,
    spoilers="none",
    labels=labels,
    language="en"
)

# Save the document
doc.save("example_output.docx")
"""
