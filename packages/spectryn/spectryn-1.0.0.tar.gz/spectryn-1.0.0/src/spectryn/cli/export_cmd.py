"""
Export Command - Export stories to various formats.

Supported formats:
- PDF (requires weasyprint or pdfkit)
- HTML
- DOCX (requires python-docx)
- CSV
- JSON
"""

import json
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from .exit_codes import ExitCode
from .output import Console, Symbols


@dataclass
class ExportOptions:
    """Options for export."""

    include_subtasks: bool = True
    include_comments: bool = False
    include_links: bool = True
    include_metadata: bool = True
    template: str | None = None
    style: str = "default"  # default, minimal, detailed


@dataclass
class ExportResult:
    """Result of export operation."""

    success: bool = True
    output_path: str = ""
    format: str = ""
    stories_exported: int = 0
    errors: list[str] = field(default_factory=list)


def export_to_html(stories: list, epic_title: str, options: ExportOptions) -> str:
    """
    Export stories to HTML format.

    Args:
        stories: List of UserStory objects.
        epic_title: Title of the epic.
        options: Export options.

    Returns:
        HTML content string.
    """
    # Modern CSS styles
    css = """
    <style>
        * { box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 2rem;
            color: #1a1a1a;
            background: #fafafa;
        }
        h1 {
            color: #2563eb;
            border-bottom: 3px solid #2563eb;
            padding-bottom: 0.5rem;
        }
        h2 {
            color: #374151;
            margin-top: 2rem;
        }
        .story {
            background: white;
            border-radius: 8px;
            padding: 1.5rem;
            margin: 1rem 0;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            border-left: 4px solid #2563eb;
        }
        .story-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 1rem;
        }
        .story-id {
            font-weight: bold;
            color: #2563eb;
        }
        .story-title {
            font-size: 1.1rem;
            font-weight: 600;
        }
        .meta {
            display: flex;
            gap: 1rem;
            font-size: 0.9rem;
            color: #6b7280;
        }
        .badge {
            display: inline-block;
            padding: 0.2rem 0.6rem;
            border-radius: 4px;
            font-size: 0.8rem;
            font-weight: 500;
        }
        .status-done { background: #d1fae5; color: #065f46; }
        .status-progress { background: #fef3c7; color: #92400e; }
        .status-planned { background: #e5e7eb; color: #374151; }
        .priority-high { background: #fee2e2; color: #991b1b; }
        .priority-medium { background: #fef3c7; color: #92400e; }
        .priority-low { background: #dbeafe; color: #1e40af; }
        .description { margin: 1rem 0; }
        .subtasks {
            margin-top: 1rem;
            padding-top: 1rem;
            border-top: 1px solid #e5e7eb;
        }
        .subtask {
            display: flex;
            align-items: center;
            gap: 0.5rem;
            padding: 0.3rem 0;
        }
        .subtask-done { text-decoration: line-through; color: #9ca3af; }
        .ac-list { list-style: none; padding: 0; }
        .ac-item { padding: 0.3rem 0; }
        .footer {
            margin-top: 3rem;
            padding-top: 1rem;
            border-top: 1px solid #e5e7eb;
            font-size: 0.8rem;
            color: #6b7280;
            text-align: center;
        }
    </style>
    """

    html_parts = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "<meta charset='UTF-8'>",
        f"<title>{epic_title}</title>",
        css,
        "</head>",
        "<body>",
        f"<h1>🚀 {epic_title}</h1>",
    ]

    # Summary
    total_points = sum(s.story_points or 0 for s in stories)
    done_count = sum(
        1 for s in stories if s.status and s.status.value.lower() in ("done", "closed")
    )

    html_parts.append(f"<p><strong>Stories:</strong> {len(stories)} | ")
    html_parts.append(f"<strong>Total Points:</strong> {total_points} | ")
    html_parts.append(f"<strong>Completed:</strong> {done_count}</p>")

    # Stories
    html_parts.append("<h2>User Stories</h2>")

    for story in stories:
        status = story.status.value if story.status else "Planned"
        status_lower = status.lower()
        status_class = (
            "status-done"
            if status_lower in ("done", "closed")
            else ("status-progress" if "progress" in status_lower else "status-planned")
        )

        priority = story.priority.value if story.priority else "Medium"
        priority_class = (
            "priority-high"
            if priority.lower() in ("high", "critical")
            else ("priority-low" if priority.lower() == "low" else "priority-medium")
        )

        html_parts.append("<div class='story'>")
        html_parts.append("<div class='story-header'>")
        html_parts.append(f"<span class='story-id'>{story.id}</span>")
        html_parts.append("<div class='meta'>")
        html_parts.append(f"<span class='badge {status_class}'>{status}</span>")
        html_parts.append(f"<span class='badge {priority_class}'>{priority}</span>")
        if story.story_points:
            html_parts.append(f"<span>{story.story_points} pts</span>")
        html_parts.append("</div>")
        html_parts.append("</div>")

        html_parts.append(f"<div class='story-title'>{story.title}</div>")

        # Description
        if story.description:
            desc = story.description
            if hasattr(desc, "as_a") and desc.as_a:
                html_parts.append("<div class='description'>")
                html_parts.append(f"<p><strong>As a</strong> {desc.as_a}</p>")
                if hasattr(desc, "i_want") and desc.i_want:
                    html_parts.append(f"<p><strong>I want</strong> {desc.i_want}</p>")
                if hasattr(desc, "so_that") and desc.so_that:
                    html_parts.append(f"<p><strong>So that</strong> {desc.so_that}</p>")
                html_parts.append("</div>")

        # Acceptance criteria
        if story.acceptance_criteria and story.acceptance_criteria.items:
            html_parts.append("<div class='subtasks'>")
            html_parts.append("<strong>Acceptance Criteria</strong>")
            html_parts.append("<ul class='ac-list'>")
            for ac in story.acceptance_criteria.items:
                html_parts.append(f"<li class='ac-item'>☐ {ac}</li>")
            html_parts.append("</ul>")
            html_parts.append("</div>")

        # Subtasks
        if options.include_subtasks and story.subtasks:
            html_parts.append("<div class='subtasks'>")
            html_parts.append("<strong>Subtasks</strong>")
            for subtask in story.subtasks:
                done_class = "subtask-done" if subtask.is_complete else ""
                checkbox = "☑" if subtask.is_complete else "☐"
                html_parts.append(
                    f"<div class='subtask {done_class}'>{checkbox} {subtask.name}</div>"
                )
            html_parts.append("</div>")

        html_parts.append("</div>")

    # Footer
    html_parts.append("<div class='footer'>")
    html_parts.append(f"Generated by spectra on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    html_parts.append("</div>")

    html_parts.append("</body>")
    html_parts.append("</html>")

    return "\n".join(html_parts)


def export_to_csv(stories: list) -> str:
    """Export stories to CSV format."""
    import csv
    import io

    output = io.StringIO()
    writer = csv.writer(output)

    # Header
    writer.writerow(
        [
            "ID",
            "Title",
            "Status",
            "Priority",
            "Story Points",
            "Subtask Count",
            "AC Count",
            "Assignee",
        ]
    )

    # Data
    for story in stories:
        writer.writerow(
            [
                str(story.id),
                story.title,
                story.status.value if story.status else "",
                story.priority.value if story.priority else "",
                story.story_points or "",
                len(story.subtasks) if story.subtasks else 0,
                len(story.acceptance_criteria.items) if story.acceptance_criteria else 0,
                story.assignee or "",
            ]
        )

    return output.getvalue()


def export_to_json(stories: list, epic_title: str) -> str:
    """Export stories to JSON format."""
    data = {
        "epic": epic_title,
        "exported_at": datetime.now().isoformat(),
        "story_count": len(stories),
        "stories": [
            {
                "id": str(story.id),
                "title": story.title,
                "status": story.status.value if story.status else None,
                "priority": story.priority.value if story.priority else None,
                "story_points": story.story_points,
                "description": {
                    "as_a": story.description.as_a if story.description else None,
                    "i_want": story.description.i_want if story.description else None,
                    "so_that": story.description.so_that if story.description else None,
                }
                if story.description
                else None,
                "acceptance_criteria": (
                    story.acceptance_criteria.items if story.acceptance_criteria else []
                ),
                "subtasks": [
                    {"name": st.name, "complete": st.is_complete} for st in (story.subtasks or [])
                ],
            }
            for story in stories
        ],
    }

    return json.dumps(data, indent=2)


def run_export(
    console: Console,
    input_path: str,
    output_path: str | None = None,
    output_format: str = "html",
    include_subtasks: bool = True,
    include_comments: bool = False,
) -> int:
    """
    Run the export command.

    Args:
        console: Console for output.
        input_path: Path to markdown file.
        output_path: Output file path.
        output_format: Output format (html, pdf, csv, json, docx).
        include_subtasks: Include subtasks in export.
        include_comments: Include comments in export.

    Returns:
        Exit code.
    """
    from spectryn.adapters.parsers import MarkdownParser

    console.header(f"spectra Export {Symbols.DOWNLOAD}")
    console.print()

    # Check file exists
    if not Path(input_path).exists():
        console.error(f"File not found: {input_path}")
        return ExitCode.FILE_NOT_FOUND

    console.info(f"Source: {input_path}")
    console.info(f"Format: {output_format}")

    # Determine output path
    if not output_path:
        stem = Path(input_path).stem
        extensions = {
            "html": ".html",
            "pdf": ".pdf",
            "csv": ".csv",
            "json": ".json",
            "docx": ".docx",
        }
        ext = extensions.get(output_format, ".html")
        output_path = f"{stem}_export{ext}"

    console.info(f"Output: {output_path}")
    console.print()

    # Parse stories
    console.info("Parsing stories...")
    parser = MarkdownParser()

    try:
        epic = parser.parse_epic(input_path)
        stories = epic.stories
        epic_title = f"{epic.key}: {epic.title}"
    except Exception:
        stories = parser.parse_stories(input_path)
        epic_title = Path(input_path).stem

    console.info(f"Found {len(stories)} stories")

    options = ExportOptions(
        include_subtasks=include_subtasks,
        include_comments=include_comments,
    )

    # Export based on format
    try:
        if output_format == "html":
            content = export_to_html(stories, epic_title, options)
            Path(output_path).write_text(content, encoding="utf-8")

        elif output_format == "csv":
            content = export_to_csv(stories)
            Path(output_path).write_text(content, encoding="utf-8")

        elif output_format == "json":
            content = export_to_json(stories, epic_title)
            Path(output_path).write_text(content, encoding="utf-8")

        elif output_format == "pdf":
            # Generate HTML first, then convert
            html_content = export_to_html(stories, epic_title, options)

            try:
                from weasyprint import HTML

                HTML(string=html_content).write_pdf(output_path)
            except ImportError:
                console.error("PDF export requires weasyprint: pip install weasyprint")
                return ExitCode.CONFIG_ERROR

        elif output_format == "docx":
            try:
                from docx import Document  # type: ignore

                doc = Document()
                doc.add_heading(epic_title, 0)

                for story in stories:
                    doc.add_heading(f"{story.id}: {story.title}", 1)

                    # Metadata
                    p = doc.add_paragraph()
                    p.add_run("Status: ").bold = True
                    p.add_run(story.status.value if story.status else "Planned")
                    p.add_run(" | Points: ").bold = True
                    p.add_run(str(story.story_points or "TBD"))

                    # Description
                    if story.description and hasattr(story.description, "as_a"):
                        doc.add_paragraph(f"As a {story.description.as_a}")
                        if hasattr(story.description, "i_want"):
                            doc.add_paragraph(f"I want {story.description.i_want}")

                    # Subtasks
                    if options.include_subtasks and story.subtasks:
                        doc.add_heading("Subtasks", 2)
                        for st in story.subtasks:
                            checkbox = "☑" if st.is_complete else "☐"
                            doc.add_paragraph(f"{checkbox} {st.name}")

                doc.save(output_path)

            except ImportError:
                console.error("DOCX export requires python-docx: pip install python-docx")
                return ExitCode.CONFIG_ERROR

        else:
            console.error(f"Unknown format: {output_format}")
            console.info("Supported: html, pdf, csv, json, docx")
            return ExitCode.CONFIG_ERROR

        console.print()
        console.success(f"Exported to: {output_path}")
        console.info(f"Stories exported: {len(stories)}")

        return ExitCode.SUCCESS

    except Exception as e:
        console.error(f"Export failed: {e}")
        return ExitCode.ERROR
