
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn

CONFIG = {
    "output_path": "caption_output.docx",
    "case_no": "322-744263-23",
    "petitioner": "MORGAN MICHELLE MYERS",
    "respondent": "CHARLES DUSTIN MYERS",
    "children_line": "M.E.M. AND C.R.M.,",
    "motion_title": "MOTION FOR NO-EVIDENCE SUMMARY JUDGMENT",
    "court_heading": "IN THE DISTRICT COURT",
    "judicial_district": "322ND JUDICIAL DISTRICT",
    "county_line": "TARRANT COUNTY, TEXAS",
}

def _tbl_borders_top_bottom(table, weight=18):
    tblPr = table._tbl.tblPr
    borders = None
    for child in tblPr.iterchildren():
        if child.tag.endswith('tblBorders'):
            borders = child; break
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    def side(name, val, sz):
        el = None
        for c in borders.iterchildren():
            if c.tag.endswith(name):
                el = c; break
        if el is None:
            el = OxmlElement(f"w:{name}")
            borders.append(el)
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
    side('top', 'single', weight)
    side('bottom', 'single', weight)
    for n in ('left','right','insideH','insideV'):
        side(n, 'nil', 0)

def _para_rule(doc, above=False, below=False):
    p = doc.add_paragraph("")
    pPr = p._p.get_or_add_pPr()
    bdr = None
    for c in pPr.iterchildren():
        if c.tag.endswith('pBdr'):
            bdr = c; break
    if bdr is None:
        bdr = OxmlElement('w:pBdr'); pPr.append(bdr)
    def add(tag):
        el = None
        for c in bdr.iterchildren():
            if c.tag.endswith(tag):
                el = c; break
        if el is None:
            el = OxmlElement(f"w:{tag}"); bdr.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '18')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
    if above: add('top')
    if below: add('bottom')
    return p

def generate_caption_doc(
    path,
    case_no,
    petitioner,
    respondent,
    children_line,
    motion_title,
    court_heading,
    judicial_district,
    county_line,
):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    # Case number centered
    p = doc.add_paragraph(f"NO. {case_no}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # Caption table
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False  # lock widths
    # Set explicit column widths
    widths = (Inches(3.9), Inches(0.35), Inches(2.95))
    for col, w in zip(table.columns, widths):
        for cell in col.cells:
            cell.width = w
    _tbl_borders_top_bottom(table, weight=24)

    left_lines = [
        "IN THE MATTER OF",
        "THE MARRIAGE OF",
        petitioner,
        "AND",
        respondent,
        "AND IN THE INTEREST OF",
        children_line,
        "CHILDREN",
    ]
    right_lines = [
        court_heading,
        "",
        judicial_district,
        county_line,
        "", "", "", "",
    ]

    for i in range(8):
        row = table.rows[i]
        L, M, R = row.cells
        L.paragraphs[0].text = left_lines[i]
        M.paragraphs[0].text = "§"
        R.paragraphs[0].text = right_lines[i]

        # Alignment & fonts
        L.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        M.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        R.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        for c in (L, M, R):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p0 = c.paragraphs[0]
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after = Pt(0)
            for r in p0.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)

    # Rules framing the title
    _para_rule(doc, below=True)
    title = doc.add_paragraph(motion_title.upper())
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12); r.font.bold = True; r.font.all_caps = True
    _para_rule(doc, above=True)

    # First line after title
    p2 = doc.add_paragraph("TO THE HONORABLE JUDGE OF THE COURT:")
    for r in p2.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    doc.save(path)

def main():
    generate_caption_doc(
        path=CONFIG["output_path"],
        case_no=CONFIG["case_no"],
        petitioner=CONFIG["petitioner"],
        respondent=CONFIG["respondent"],
        children_line=CONFIG["children_line"],
        motion_title=CONFIG["motion_title"],
        court_heading=CONFIG["court_heading"],
        judicial_district=CONFIG["judicial_district"],
        county_line=CONFIG["county_line"],
    )

main()
