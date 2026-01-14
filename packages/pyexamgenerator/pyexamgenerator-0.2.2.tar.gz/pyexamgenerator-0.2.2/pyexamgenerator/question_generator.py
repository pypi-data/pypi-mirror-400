# pyexamgenerator: A tool for generating exams from PDF files using AI.
# Copyright (C) 2024 Daniel Sánchez-García

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

from google import genai
from google.genai import types # Opcional, útil para tipado, pero no estricto
import PyPDF2
import json
import re
import os
import pandas as pd
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict, Optional

class QuotaExceededError(Exception):
    """Excepción para cuando se excede el límite de cuota de la API de Gemini."""
    pass

class ServiceOverloadedError(Exception):
    """Excepción para cuando los servidores de Google están saturados (Error 503)."""
    pass

class QuestionGenerator:
    """
    Generates multiple-choice questions from PDF files using the Gemini model.
    """

    def __init__(self, api_key: str, model_name: str):
        """
        Initializes the QuestionGenerator class.

        Args:
            api_key (str): The API key for the Gemini model.
            model_name (str): The name of the Gemini model to use (e.g., 'models/gemini-2.5-pro').
                              This is a required argument.
        """
        self.client = genai.Client(api_key=api_key)
        self.model_name = model_name

        print(f"Cliente de Gemini inicializado para modelo: {model_name}")

        self.prompt_types = {
            "PRL": """
            Genera preguntas tipo test de opción múltiple enfocadas en Prevención de Riesgos Laborales (PRL), basadas en el siguiente texto, siguiendo estas instrucciones iniciales:
            1.  **Centrarse en leyes y reglamentos:** Las preguntas deben poder entenderse por sí mismas, haciendo referencia únicamente a leyes o reglamentos relevantes. No se puede poner: "Según el texto"
            2.  **Evita preguntas sobre números de artículos o códigos:** No incluyas preguntas que hagan referencia a números de artículos, códigos de leyes o reglamentos.
            Además, ten en cuenta las siguientes indicaciones adicionales
            """,
            "PM": """
            Genera preguntas tipo test de opción múltiple enfocadas en Gestión de Proyectos.
            """
        }
        self.existing_bank_df = None
        self.similarity_threshold = 0.8  # Default threshold
        self.current_prompt_example_content_type = "solo_enunciados"  # Default value

    def extract_pdf_text(self, pdf_path: str) -> Optional[str]:
        """
        Extracts text from a PDF file.

        Args:
            pdf_path (str): The path to the PDF file.

        Returns:
            Optional[str]: The extracted text from the PDF, or None if an error occurs.
        """
        try:
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                self.text = ""
                self.page_texts = []  # List to store the text of each page
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    self.page_texts.append(page_text)
                    self.text += page_text
                self.num_pages = len(pdf_reader.pages)  # Save the total number of pages
            return self.text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None

    def extract_topic_number_from_path(self, pdf_path: str) -> str:
        """
        Extracts the topic name from the PDF filename.

        Args:
            pdf_path (str): The path to the PDF file.

        Returns:
            str: The extracted topic name, or "Tema desconocido" if not found.
        """
        file_name = os.path.basename(pdf_path)

        # Using topic match: looking for a pattern like TEMA XX
        # topic_match = re.search(r'TEMA\s*(\d+)', file_name, re.IGNORECASE)
        # self.topic_number = topic_match.group(1) if topic_match else file_name.split('.pdf')[0]

        # Using the PDF name directly
        self.topic_number = file_name.split('.pdf')[0]

        return self.topic_number

    def _group_pages_into_chunks(self, pages_per_chunk: int) -> List[str]:
        """
        Groups the text from pages into larger chunks.

        Args:
            pages_per_chunk (int): The number of pages to include in each chunk.

        Returns:
            List[str]: A list of strings, where each string is a text chunk.
        """
        if not hasattr(self, 'page_texts') or not self.page_texts:
            return []

        text_chunks = []
        current_chunk_pages = []
        for i, page_text in enumerate(self.page_texts):
            current_chunk_pages.append(page_text)
            if (i + 1) % pages_per_chunk == 0 or (i + 1) == len(self.page_texts):
                text_chunks.append("\n".join(current_chunk_pages))
                current_chunk_pages = []
        return text_chunks

    def _build_prompt(
            self,
            working_chunk_text: str,
            num_questions_to_generate: int,
            prompt_type: str,
            custom_prompt: Optional[str],
            existing_questions_df: Optional[pd.DataFrame],
            current_topic: str,
            full_document_context_text: Optional[str] = None,
            bank_prompt_scope: Optional[str] = None,
            questions_from_current_chunk_attempts: Optional[List[Dict]] = None
    ) -> str:
        """
        Constructs the complete prompt to be sent to the Gemini model.

        Args:
            working_chunk_text (str): The specific text segment to generate questions from.
            num_questions_to_generate (int): The target number of questions to generate.
            prompt_type (str): The key for the predefined prompt type (e.g., "PRL").
            custom_prompt (Optional[str]): A user-provided custom prompt.
            existing_questions_df (Optional[pd.DataFrame]): DataFrame of existing questions for guidance.
            current_topic (str): The topic of the current PDF.
            full_document_context_text (Optional[str]): The full text of the document for context.
            bank_prompt_scope (Optional[str]): Scope for selecting guidance questions ('mismo_tema', 'todos_los_temas').
            questions_from_current_chunk_attempts (Optional[List[Dict]]): Questions already generated for this chunk.

        Returns:
            str: The fully constructed prompt string.
        """
        user_prefix = ""
        if custom_prompt:
            user_prefix = custom_prompt
        elif prompt_type in self.prompt_types:
            user_prefix = self.prompt_types[prompt_type]
        else:
            user_prefix = """
            Genera preguntas tipo test de opción múltiple basadas en el siguiente texto, siguiendo estas instrucciones:
            """

        fixed_prefix = f"""
        Genera {num_questions_to_generate} preguntas tipo test de opción múltiple con una única
        respuesta correcta basadas en el texto que adjunto al final, siguiendo estas instrucciones:
        """

        # --- Guidance from the existing question bank ---
        prompt_guidance_from_bank_str = ""
        if existing_questions_df is not None and \
                not existing_questions_df.empty and \
                bank_prompt_scope and bank_prompt_scope != "no_incluir_en_prompt":

            questions_for_guidance_df = pd.DataFrame()
            header_for_bank_guidance = ""

            if bank_prompt_scope == "mismo_tema":
                if 'Tema' in existing_questions_df.columns:
                    questions_for_guidance_df = existing_questions_df[
                        existing_questions_df['Tema'].astype(str).str.lower() == current_topic.lower()
                        ]
                if not questions_for_guidance_df.empty:
                    header_for_bank_guidance = "\n\nPara evitar redundancia con el BANCO DE PREGUNTAS EXISTENTE, considera las siguientes preguntas del MISMO TEMA (asegúrate de que las nuevas preguntas sean originales y basadas estrictamente en el FRAGMENTO DE TRABAJO proporcionado):\n"
            elif bank_prompt_scope == "todos_los_temas":
                questions_for_guidance_df = existing_questions_df
                if not questions_for_guidance_df.empty:
                    header_for_bank_guidance = "\n\nPara evitar redundancia con el BANCO DE PREGUNTAS EXISTENTE, considera las siguientes preguntas de TODO EL BANCO (asegúrate de que las nuevas preguntas sean originales y basadas estrictamente en el TEXTO COMPLETO proporcionado):\n"

            if not questions_for_guidance_df.empty:
                prompt_guidance_from_bank_str += header_for_bank_guidance
                for _, row in questions_for_guidance_df.iterrows():
                    question_text = row.get('Pregunta', '')
                    prompt_guidance_from_bank_str += f"- Pregunta del banco: {question_text}\n"
                    if self.current_prompt_example_content_type == "enunciados_y_respuestas":
                        options = [row.get(f'Respuesta {chr(65 + i)}', '') for i in range(4) if pd.notna(row.get(f'Respuesta {chr(65 + i)}', ''))]
                        options_str = "\n".join([f"  - {opt}" for opt in options if opt])
                        if options_str:
                            prompt_guidance_from_bank_str += f"{options_str}\n"
                    prompt_guidance_from_bank_str += "\n"

        # --- Guidance from questions already generated for this chunk in previous attempts ---
        prompt_guidance_from_current_attempts_str = ""
        if questions_from_current_chunk_attempts:  # This is a list of dictionaries
            prompt_guidance_from_current_attempts_str += "\n\nAdicionalmente, para evitar generar preguntas idénticas o muy similares a las YA GENERADAS PARA ESTE MISMO FRAGMENTO en intentos previos, ten en cuenta las siguientes preguntas que ya se han aceptado:\n"
            for question_dict in questions_from_current_chunk_attempts:
                question_text = question_dict.get('Pregunta', '')
                if question_text:
                    prompt_guidance_from_current_attempts_str += f"- Pregunta ya generada para este fragmento: {question_text}\n"
                    # Optionally, you could also include the options if you find it useful
                    # if self.current_prompt_example_content_type == "enunciados_y_respuestas":
                    #     options = [question_dict.get(f'Respuesta {chr(65 + i)}', '') for i in range(4)]
                    #     options_str = "\n".join([f"  - {opt}" for opt in options if opt])
                    #     if options_str:
                    #         prompt_guidance_from_current_attempts_str += f"{options_str}\n"
                    prompt_guidance_from_current_attempts_str += "\n"

        context_section = ""
        chunk_instruction_and_text_section = ""

        if full_document_context_text is not None:
            context_section = f"""
            Aquí está el **CONTEXTO COMPLETO** del documento, que te proporcionamos para que tengas una comprensión general del tema y el documento en su totalidad. Puedes referenciar este texto para obtener contexto adicional, pero las preguntas deben basarse estrictamente en el **FRAGMENTO DE TRABAJO**.
            {full_document_context_text}
            """
            chunk_instruction_and_text_section = f"""
            ---
            **IMPORTANTE**: Debes generar las preguntas **ÚNICAMENTE** basándote en el siguiente "**FRAGMENTO DE TRABAJO**". La respuesta a cada pregunta debe ser directamente derivable del contenido de este FRAGMENTO DE TRABAJO.
            ---
            Aquí está el FRAGMENTO DE TRABAJO sobre el que debes generar las preguntas:
            {working_chunk_text}
            """
        else:
            chunk_instruction_and_text_section = f"""
            Aquí está el **TEXTO COMPLETO** sobre el que debes generar las preguntas:
            {working_chunk_text}
            """

        prompt = f"""
        {user_prefix}

        {fixed_prefix}


        1. **Incluye la respuesta correcta:** Para cada pregunta, indica cuál de las opciones es la respuesta correcta (solo la letra, sin paréntesis).
        2. **Proporciona el texto relevante:** Incluye un fragmento del texto del PDF donde se encuentra la justificación de la respuesta correcta.
        3. **Formato de diccionario:** Presenta las preguntas y respuestas en formato de diccionario de Python, donde cada diccionario representa una pregunta y tiene las siguientes claves:
        * "Pregunta"
        * "Opciones" (una lista de las opciones A, B, C, D, sin la letra inicial y el paréntesis)
        * "Respuesta correcta" (la letra de la opción correcta, sin paréntesis)
        * "Texto relevante" (el texto relevante del PDF)
        Separa cada pregunta con una línea en blanco.
        Puedes incluir preguntas del tipo: señala la afirmación incorrecta entre las siguientes opciones; o señala la opción correcta entre las siguientes.
        Además, puedes incluir respuestas como todas son correctas, o todas son incorrectas.

        {prompt_guidance_from_bank_str}
        {prompt_guidance_from_current_attempts_str}

        {context_section}

        {chunk_instruction_and_text_section}
        """
        return prompt

    def _analyze_gemini_response(self, response_text: str, topic_number: str) -> List[Dict]:
        """
        Analyzes the model's response, filters it by similarity, and returns a list of dictionaries with accepted questions.
        Rejected questions are added to the instance's `self.rejected_questions_list`.

        Args:
            response_text (str): The raw text response from the Gemini model.
            topic_number (str): The topic associated with these questions.

        Returns:
            List[Dict]: A list of accepted question dictionaries.
        """
        accepted_questions = []
        # The model is asked to separate question dictionaries with a blank line.
        question_strs = response_text.split('\n\n')

        for question_str in question_strs:
            cleaned_str = question_str.strip()
            # Attempt to find a Python dictionary-like structure within the string.
            match = re.search(r'\{(.*?)\}', cleaned_str, re.DOTALL)
            if not match:
                continue

            json_like_str = "{" + match.group(1) + "}"
            try:
                # Use eval() to parse the string as a Python dictionary.
                question_data = eval(json_like_str)
                if not isinstance(question_data, dict) or "Pregunta" not in question_data:
                    continue

                # Check for similarity against the existing bank.
                is_similar, sim_score, matched_text = self._is_similar_question(question_data["Pregunta"])

                # Build the base dictionary for the generated question.
                question_dict = {
                    "Tema": f"{topic_number}",
                    "Estado": "Pendiente de revisar",
                    "Pregunta": question_data.get("Pregunta", ""),
                    "Respuesta A": question_data.get("Opciones", [""] * 4)[0],
                    "Respuesta B": question_data.get("Opciones", [""] * 4)[1],
                    "Respuesta C": question_data.get("Opciones", [""] * 4)[2],
                    "Respuesta D": question_data.get("Opciones", [""] * 4)[3],
                    "Respuesta correcta": question_data.get("Respuesta correcta", "").strip().upper(),
                    "Texto relevante": question_data.get("Texto relevante", "")
                }

                if is_similar:
                    # If similar, add it to the rejected list with extra details for reporting.
                    question_dict['Similitud'] = sim_score
                    question_dict['Pregunta Coincidente'] = matched_text
                    self.rejected_questions_list.append(question_dict)
                else:
                    # If not similar, add it to the accepted list for this batch.
                    accepted_questions.append(question_dict)

            except (SyntaxError, TypeError, NameError):
                print(f"Error al evaluar la estructura como diccionario: {json_like_str}")

        return accepted_questions

    def save_questions_to_excel(
            self,
            df: pd.DataFrame,
            filepath: str
    ) -> None:
        """Saves the questions to an Excel file."""
        if not df.empty:
            if 'Número de pregunta' not in df.columns:
                df.insert(0, 'Número de pregunta', range(1, len(df) + 1))
            df.to_excel(filepath, index=False)
            print(f"Preguntas aceptadas guardadas en: '{filepath}'")
        else:
            print("No se generaron preguntas aceptadas para guardar en Excel.")

    def save_questions_to_docx(
            self,
            df: pd.DataFrame,
            filepath: str
    ) -> None:
        """Saves the questions to a DOCX file for manual review."""
        if not df.empty:
            document = Document()
            document.add_heading('Banco de Preguntas - Pendiente de Revisar', level=1)
            for index, row in df.iterrows():
                pregunta_paragraph = document.add_paragraph()
                pregunta_run = pregunta_paragraph.add_run(f"Pregunta {index + 1}: {row['Pregunta']}")
                pregunta_run.bold = True
                document.add_paragraph(f"Tema: {row['Tema']}")
                document.add_paragraph(f"Estado: {row['Estado']}")
                document.add_paragraph(f"A) {row['Respuesta A']}")
                document.add_paragraph(f"B) {row['Respuesta B']}")
                document.add_paragraph(f"C) {row['Respuesta C']}")
                document.add_paragraph(f"D) {row['Respuesta D']}")
                document.add_paragraph(f"Respuesta correcta: {row['Respuesta correcta']}")
                document.add_paragraph(f"Texto relevante: {row['Texto relevante']}")
                document.add_paragraph()
            document.save(filepath)
        else:
            print("No hay preguntas para guardar en DOCX.")

    def save_rejected_questions_to_docx(self, df: pd.DataFrame, filepath: str) -> None:
        """
        Saves questions discarded due to similarity to a DOCX file for inspection.

        Args:
            df (pd.DataFrame): The DataFrame containing the discarded questions and similarity details.
            filepath (str): The filepath for the output file.
        """
        if df.empty:
            print("No se descartaron preguntas por similitud.")
            return

        document = Document()
        document.add_heading('Preguntas Descartadas por Similitud', level=1)
        document.add_paragraph(
            "Las siguientes preguntas fueron generadas por el modelo, pero se descartaron "
            "automáticamente porque su similitud con una pregunta existente en el banco "
            "superó el umbral establecido."
        )

        for index, row in df.iterrows():
            # Visual separator between questions
            document.add_paragraph("---")

            # Information about the discarded question
            p = document.add_paragraph()
            p.add_run('Pregunta descartada: ').bold = True
            p.add_run(f"{row['Pregunta']}")

            document.add_paragraph(f"Tema: {row['Tema']}")
            document.add_paragraph(f"A) {row['Respuesta A']}")
            document.add_paragraph(f"B) {row['Respuesta B']}")
            document.add_paragraph(f"C) {row['Respuesta C']}")
            document.add_paragraph(f"D) {row['Respuesta D']}")
            document.add_paragraph(f"Respuesta correcta: {row['Respuesta correcta']}")

            # Information about the reason for rejection
            p_reason = document.add_paragraph()
            p_reason.add_run('Motivo del descarte:').bold = True

            # Format the similarity score as a percentage
            similarity_score_pct = f"{row.get('Similitud', 0.0) * 100:.1f}%"

            p_details = document.add_paragraph()
            p_details.add_run(f"Similitud: ").bold = True
            p_details.add_run(f"{similarity_score_pct}")

            p_match = document.add_paragraph()
            p_match.add_run("Coincide con la pregunta existente: ").bold = True
            p_match.add_run(f"{row.get('Pregunta Coincidente', 'N/A')}")

        document.save(filepath)
        print(f"Se guardó un informe de {len(df)} preguntas descartadas en: '{filepath}'")

    def _normalize_text(self, text: str) -> str:
        """Normalizes text for comparison (lowercase, no punctuation)."""
        if not isinstance(text, str):
            return ""
        text = text.lower()
        text = re.sub(r'[^\w\s]', '', text)  # Remove punctuation
        return text

    def _calculate_jaccard_similarity(self, text1: str, text2: str) -> float:
        """Calculates the Jaccard similarity between two texts."""
        words1 = set(self._normalize_text(text1).split())
        words2 = set(self._normalize_text(text2).split())
        if not words1 and not words2:
            return 1.0
        if not words1 or not words2:
            return 0.0
        intersection = len(words1.intersection(words2))
        union = len(words1.union(words2))
        return intersection / union

    def _is_similar_question(self,
                             new_question_text: str
                             ) -> tuple[bool, float, Optional[str]]:
        """
        Checks if the new question is similar to any existing one in the bank,
        comparing ONLY the question statements for filtering purposes.

        Args:
            new_question_text (str): The text of the new question.

        Returns:
            tuple[bool, float, Optional[str]]: (is_similar, similarity_score, matching_question_text)
        """
        if self.existing_bank_df is None or self.existing_bank_df.empty or self.similarity_threshold > 1.0:
            # If the threshold is > 1.0 (e.g., 1.1), the filter is effectively disabled
            return (False, 0.0, None)

        normalized_new_question_text = self._normalize_text(new_question_text)

        for _, row in self.existing_bank_df.iterrows():
            existing_question_text_from_bank = row.get('Pregunta')
            if existing_question_text_from_bank and isinstance(existing_question_text_from_bank, str):
                normalized_existing_question_text = self._normalize_text(existing_question_text_from_bank)

                question_similarity = self._calculate_jaccard_similarity(
                    normalized_new_question_text,
                    normalized_existing_question_text
                )

                if question_similarity >= self.similarity_threshold:
                    # For filtering, the similarity of the statement is sufficient
                    return (True, question_similarity, existing_question_text_from_bank)

        return (False, 0.0, None)

    def generate_multiple_choice_questions(
            self,
            pdf_paths: List[str],
            prompt_type: str,
            num_questions_per_chunk_target: int = 5,
            output_filename: Optional[str] = None,
            output_dir: Optional[str] = None,
            custom_prompt: Optional[str] = None,
            generate_docx: bool = True,
            existing_bank_path: Optional[str] = None,
            process_by_pages: bool = False,
            pages_per_chunk: int = 1,
            similarity_threshold: Optional[float] = 0.8,
            bank_prompt_scope: Optional[str] = None,
            prompt_example_content_type: str = "solo_enunciados",
            print_raw_gemini_answer: bool = False,
            max_generation_attempts_per_chunk: int = 3
    ) -> pd.DataFrame:
        """
        Main function to generate multiple-choice questions from multiple PDFs,
        with retries to reach the desired number of questions per chunk.

        This method orchestrates the entire generation process, including:
        - Loading an existing question bank for similarity checks.
        - Iterating through each provided PDF.
        - Splitting PDFs into chunks if required.
        - Looping with multiple attempts per chunk to reach the target number of questions.
        - Building prompts, generating content, and analyzing responses.
        - Saving accepted and rejected questions to files.

        Args:
            pdf_paths (List[str]): List of paths to the PDF files.
            prompt_type (str): The key for the predefined prompt type.
            num_questions_per_chunk_target (int): The target number of questions per text chunk.
            output_filename (Optional[str]): The base name for the output files.
            output_dir (Optional[str]): The directory to save the output files.
                If None, uses the current working directory.
            custom_prompt (Optional[str]): A user-provided custom prompt to override the default.
            generate_docx (bool): Whether to generate a DOCX file for review.
            existing_bank_path (Optional[str]): Path to an existing Excel question bank for similarity filtering.
            process_by_pages (bool): If True, process the PDF in chunks of pages.
            pages_per_chunk (int): The number of pages per chunk if processing by pages.
            similarity_threshold (Optional[float]): The threshold for the Jaccard similarity filter.
            bank_prompt_scope (Optional[str]): Scope for selecting guidance questions from the bank.
            prompt_example_content_type (str): Content type for guidance questions ('solo_enunciados' or 'enunciados_y_respuestas').
            print_raw_gemini_answer (bool): If True, prints the raw API response for debugging.
            max_generation_attempts_per_chunk (int): The maximum number of attempts to reach the target per chunk.

        Returns:
            pd.DataFrame: A DataFrame containing all the successfully generated and filtered questions.
        """
        if output_filename is None:
            output_filename = 'preguntas_pendiente_de_revisar'

        if output_dir is None:
            output_dir = os.getcwd()

        os.makedirs(output_dir, exist_ok=True)

        self.final_df = pd.DataFrame()
        self.rejected_questions_list = []  # Accumulates all rejected questions

        if similarity_threshold is not None and 0 <= similarity_threshold <= 1.0:
            self.similarity_threshold = similarity_threshold
            print(f"Filtro de similitud activado con umbral: {self.similarity_threshold}")
        else:
            self.similarity_threshold = 1.1  # Effectively disables the filter
            print(f"Filtro de similitud DESACTIVADO (umbral: {self.similarity_threshold}).")

        self.current_prompt_example_content_type = prompt_example_content_type
        print(f"Contenido de ejemplos en prompt: {self.current_prompt_example_content_type}")

        if existing_bank_path and os.path.exists(existing_bank_path):
            try:
                self.existing_bank_df = pd.read_excel(existing_bank_path)
                print(f"Banco de preguntas existente cargado desde '{existing_bank_path}' para el filtro de similitud.")
            except Exception as e:
                print(f"Error al cargar el banco de preguntas existente: {e}")
                self.existing_bank_df = None
        else:
            self.existing_bank_df = None
            if existing_bank_path:  # Only show warning if a path was provided
                print(f"Advertencia: El archivo del banco de preguntas existente no se encontró en '{existing_bank_path}'.")

        for pdf_path in pdf_paths:
            print(f"\n--- Procesando PDF: {pdf_path} ---")
            pdf_full_text_content = self.extract_pdf_text(pdf_path)
            current_pdf_topic = self.extract_topic_number_from_path(pdf_path)

            if pdf_full_text_content:
                all_questions_for_this_pdf = []
                text_chunks_to_process = []
                context_for_chunks = None

                if process_by_pages and hasattr(self, 'page_texts') and self.page_texts:
                    text_chunks_to_process = self._group_pages_into_chunks(pages_per_chunk)
                    context_for_chunks = pdf_full_text_content
                    print(f"PDF dividido en {len(text_chunks_to_process)} fragmentos de ~{pages_per_chunk} página(s).")
                else:
                    text_chunks_to_process = [pdf_full_text_content]
                    print("Procesando el PDF completo como un solo fragmento.")

                for i, chunk_text in enumerate(text_chunks_to_process):
                    print(f"\nProcesando fragmento {i + 1}/{len(text_chunks_to_process)} del PDF...")
                    questions_accepted_for_this_chunk = []  # This list will hold question dicts
                    attempts_for_this_chunk = 0

                    while len(questions_accepted_for_this_chunk) < num_questions_per_chunk_target and \
                            attempts_for_this_chunk < max_generation_attempts_per_chunk:

                        attempts_for_this_chunk += 1
                        num_still_needed = num_questions_per_chunk_target - len(questions_accepted_for_this_chunk)

                        print(f"  Intento {attempts_for_this_chunk}/{max_generation_attempts_per_chunk} para el fragmento: Se necesitan {num_still_needed} preguntas más.")

                        prompt = self._build_prompt(
                            working_chunk_text=chunk_text,
                            num_questions_to_generate=num_still_needed,
                            prompt_type=prompt_type,
                            custom_prompt=custom_prompt,
                            existing_questions_df=self.existing_bank_df,
                            current_topic=current_pdf_topic,
                            full_document_context_text=context_for_chunks if process_by_pages else None,
                            bank_prompt_scope=bank_prompt_scope,
                            questions_from_current_chunk_attempts=questions_accepted_for_this_chunk
                        )

                        try:
                            response = self.client.models.generate_content(
                                model=self.model_name,
                                contents=prompt
                            )
                            if print_raw_gemini_answer:
                                print(f"    Respuesta cruda de Gemini (Fragmento {i + 1}, Intento {attempts_for_this_chunk}):\n    >>>\n{response.text}\n    <<<\n    ---")

                            time.sleep(1)
                            newly_accepted_this_attempt = self._analyze_gemini_response(response.text, current_pdf_topic)

                            # Filter out duplicates that might have been generated within the same chunk attempt
                            unique_newly_accepted = []
                            for q_new in newly_accepted_this_attempt:
                                is_dup_within_chunk = False
                                for q_existing_chunk_dict in questions_accepted_for_this_chunk:
                                    # Compare only the statement for duplicates within the chunk
                                    if self._normalize_text(q_new.get("Pregunta", "")) == self._normalize_text(q_existing_chunk_dict.get("Pregunta", "")):
                                        is_dup_within_chunk = True
                                        break
                                if not is_dup_within_chunk:
                                    unique_newly_accepted.append(q_new)

                            questions_accepted_for_this_chunk.extend(unique_newly_accepted)
                            print(
                                f"    Intento {attempts_for_this_chunk}: {len(unique_newly_accepted)} preguntas nuevas aceptadas para este fragmento. Total para fragmento: {len(questions_accepted_for_this_chunk)}/{num_questions_per_chunk_target}")

                        except Exception as e_gen:
                            error_str = str(e_gen)
                            # Comprobamos si el error es por límite de cuota
                            if "429" in error_str and "quota" in error_str:
                                # Construimos el mensaje de error mejorado
                                suggestion = (
                                    "\n\nSugerencias para solucionarlo:\n"
                                    "1. Espera unos minutos: La cuota gratuita se reinicia cada cierto tiempo.\n"
                                    "2. Cambia de modelo: Los modelos 'Pro' son más potentes pero consumen la cuota más rápido. "
                                    "Prueba a cambiar a un modelo 'Flash' (como 'gemini-2.5-flash'), que es más rápido y económico."
                                )
                                message = (
                                    "Has excedido el límite de solicitudes a la API de Gemini (Error 429).\n"
                                    f"{suggestion}"
                                )
                                # Lanzamos nuestra excepción personalizada con el mensaje mejorado
                                raise QuotaExceededError(message) from e_gen
                            elif "503" in error_str or "overloaded" in error_str:
                                suggestion = (
                                    "\n\nSugerencias para solucionarlo:\n"
                                    "1. Cambia de modelo: El modelo que estás intentando usar está saturado actualmente. "
                                    "Intenta usar 'gemini-2.5-flash', que suele tener mayor disponibilidad.\n"
                                    "2. Espera unos minutos e inténtalo de nuevo."
                                )
                                message = (
                                    "Los servidores de Google para este modelo están saturados (Error 503).\n"
                                    f"{suggestion}"
                                )
                                raise ServiceOverloadedError(message) from e_gen
                            else:
                                # Si es otro error, lo imprimimos para depuración pero no detenemos el proceso
                                print(f"    Error durante la generación/análisis en el intento {attempts_for_this_chunk} para el fragmento {i + 1}: {e_gen}")
                                import traceback
                                print(traceback.format_exc())

                    if len(questions_accepted_for_this_chunk) < num_questions_per_chunk_target:
                        print(
                            f"  Advertencia: No se alcanzó el objetivo de {num_questions_per_chunk_target} preguntas para el fragmento {i + 1} después de {max_generation_attempts_per_chunk} intentos. Se obtuvieron {len(questions_accepted_for_this_chunk)}.")

                    all_questions_for_this_pdf.extend(questions_accepted_for_this_chunk)

                if all_questions_for_this_pdf:
                    pdf_df = pd.DataFrame(all_questions_for_this_pdf)
                    self.final_df = pd.concat([self.final_df, pdf_df], ignore_index=True)
            else:
                print(f"No se pudo extraer texto del PDF: {pdf_path}")

        excel_output_path = os.path.join(output_dir, f'{output_filename}_pendiente_de_revisar.xlsx')
        docx_output_path = os.path.join(output_dir, f'{output_filename}_pendiente_de_revisar.docx')

        self.save_questions_to_excel(self.final_df, excel_output_path)
        if generate_docx:
            self.save_questions_to_docx(self.final_df, docx_output_path)

        if self.rejected_questions_list:
            rejected_df = pd.DataFrame(self.rejected_questions_list)
            rejected_docx_path = os.path.join(output_dir, f'{output_filename}_similares_descartadas.docx')
            self.save_rejected_questions_to_docx(rejected_df, rejected_docx_path)
        else:
            print("No se descartaron preguntas por similitud durante todo el proceso.")

        print(f"\n--- Proceso de generación completado. ---")
        print(f"Total preguntas aceptadas y guardadas: {len(self.final_df)}")
        if self.rejected_questions_list:
            print(f"Total preguntas rechazadas por similitud (acumulado de todos los intentos): {len(self.rejected_questions_list)}")

        return self.final_df

# Example of use (uncomment to test)
# if __name__ == '__main__':
#     api_key = "YOUR_GEMINI_API_KEY_HERE"  # Replace with your Gemini API key
#     generator = QuestionGenerator(api_key)
#     pdf_files = ['TEMA 05_Inspecciones de seguridad.pdf', 'TEMA 08_Equipos de proteccion individual.pdf']
#
#     # Example 1: Normal processing (the whole text at once)
#     # questions_df = generator.generate_multiple_choice_questions(
#     #     pdf_files,
#     #     prompt_type="PRL",
#     #     num_questions_per_chunk_target=10,  # Generate 10 questions per PDF
#     #     output_filename="banco_de_preguntas",
#     #     generate_docx=True
#     # )
#
#     # Example 2: Processing by individual pages
#     # questions_df_pages = generator.generate_multiple_choice_questions(
#     #     pdf_files,
#     #     prompt_type="PRL",
#     #     num_questions_per_chunk_target=2,  # Generate 2 questions per page
#     #     output_filename="banco_de_preguntas_por_paginas",
#     #     generate_docx=True,
#     #     process_by_pages=True,  # Activate processing by pages
#     #     pages_per_chunk=1,  # Process each page individually
#     #     existing_bank_path="examen_PIR.xlsx"
#     # )
#
#     # Example 3: Processing by groups of pages
#     questions_df_page_groups = generator.generate_multiple_choice_questions(
#         pdf_files,
#         prompt_type="PRL",
#         num_questions_per_chunk_target=3,  # Generate 3 questions for each group of 1 page(s)
#         pages_per_chunk=1,  # Group every 1 page to process them together
#         output_filename="banco_de_preguntas_grupos_paginas",
#         generate_docx=True,
#         process_by_pages=True,  # Activate processing by pages
#         existing_bank_path=r'D:\path\to\your\existing_bank.xlsx',
#         bank_prompt_scope='todos_los_temas',
#         prompt_example_content_type='solo_enunciados'
#     )
#
#     # You can use the DataFrame to inspect the generated questions